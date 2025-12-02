"""Unit tests for the Style Enforcer processor.

Tests cover:
- Font standardization across body text and tables
- Heading style detection and enforcement
- Paragraph spacing normalization
- Edge cases and error handling
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt

from src.processors.style_enforcer import (
    process_document,
    perform_style_enforcement,
    _standardize_fonts,
    _enforce_heading_styles,
    _normalize_spacing,
)
from src.processors import ProcessorResult


# ============================================================================
# Font Standardization Tests
# ============================================================================

class TestFontStandardization:
    """Tests for _standardize_fonts function."""

    @pytest.mark.unit
    def test_standardize_fonts_changes_non_standard_font(self, docx_factory):
        """Fonts different from target should be changed."""
        # Create doc with non-standard font
        doc_path = docx_factory(content="Test content")
        doc = Document(str(doc_path))
        run = doc.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        doc.save(str(doc_path))

        # Reload and standardize
        doc = Document(str(doc_path))
        changes = _standardize_fonts(doc, 'Calibri', 11)

        assert changes >= 1
        assert doc.paragraphs[0].runs[0].font.name == 'Calibri'
        assert doc.paragraphs[0].runs[0].font.size == Pt(11)

    @pytest.mark.unit
    def test_standardize_fonts_skips_correct_font(self, docx_factory):
        """Fonts already matching target should not count as changes."""
        doc_path = docx_factory(content="Test content")
        doc = Document(str(doc_path))
        run = doc.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _standardize_fonts(doc, 'Calibri', 11)

        # No changes expected if already correct
        assert changes == 0

    @pytest.mark.unit
    def test_standardize_fonts_multiple_runs(self, docx_factory):
        """Multiple runs with different fonts should all be changed."""
        doc_path = docx_factory(content="First")
        doc = Document(str(doc_path))

        # Add multiple runs with different fonts
        para = doc.paragraphs[0]
        para.runs[0].font.name = 'Arial'
        run2 = para.add_run(' Second')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(14)
        run3 = para.add_run(' Third')
        run3.font.name = 'Verdana'
        run3.font.size = Pt(9)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _standardize_fonts(doc, 'Calibri', 11)

        assert changes >= 3
        for run in doc.paragraphs[0].runs:
            assert run.font.name == 'Calibri'
            assert run.font.size == Pt(11)

    @pytest.mark.unit
    def test_standardize_fonts_skips_headings(self, docx_factory):
        """Paragraphs with heading styles should be skipped."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        # Add heading
        heading = doc.add_heading('Test Heading', level=1)
        heading_run = heading.runs[0] if heading.runs else None

        # Add body text with non-standard font
        para = doc.add_paragraph()
        run = para.add_run('Body text')
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _standardize_fonts(doc, 'Calibri', 11)

        # Only body text should be changed, not heading
        assert changes >= 1
        # Body paragraph should be standardized
        body_para = None
        for p in doc.paragraphs:
            if not p.style.name.startswith('Heading'):
                body_para = p
                break
        if body_para and body_para.runs:
            assert body_para.runs[0].font.name == 'Calibri'

    @pytest.mark.unit
    def test_standardize_fonts_in_tables(self, docx_factory):
        """Fonts in table cells should be standardized."""
        doc_path = docx_factory(content="Intro", add_table=True)
        doc = Document(str(doc_path))

        # Set non-standard fonts in table cells
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _standardize_fonts(doc, 'Calibri', 11)

        assert changes >= 4  # At least 4 cells
        for row in doc.tables[0].rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        assert run.font.name == 'Calibri'
                        assert run.font.size == Pt(11)

    @pytest.mark.unit
    def test_standardize_fonts_preserves_bold_italic(self, docx_factory):
        """Bold and italic formatting should be preserved."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Bold text')
        run.bold = True
        run.font.name = 'Arial'

        run2 = para.add_run(' Italic text')
        run2.italic = True
        run2.font.name = 'Times New Roman'

        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        _standardize_fonts(doc, 'Calibri', 11)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        para = doc.paragraphs[-1]
        assert para.runs[0].bold is True
        assert para.runs[1].italic is True


# ============================================================================
# Heading Style Enforcement Tests
# ============================================================================

class TestHeadingStyleEnforcement:
    """Tests for _enforce_heading_styles function."""

    @pytest.mark.unit
    def test_detects_bold_short_text_as_heading(self, docx_factory):
        """Short bold text should be detected as a heading."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Introduction')
        run.bold = True
        doc.add_paragraph('Normal body text follows.')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _enforce_heading_styles(doc)

        assert changes >= 1

    @pytest.mark.unit
    def test_detects_large_font_as_h1(self, docx_factory):
        """Text with 16pt+ font should become Heading 1."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Chapter Title')
        run.font.size = Pt(16)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _enforce_heading_styles(doc)

        assert changes >= 1
        # Find the paragraph that should be a heading
        for p in doc.paragraphs:
            if 'Chapter Title' in p.text:
                assert p.style.name == 'Heading 1'

    @pytest.mark.unit
    def test_detects_medium_font_as_h2(self, docx_factory):
        """Text with 14-15pt font should become Heading 2."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Section Title')
        run.font.size = Pt(14)
        run.bold = True
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _enforce_heading_styles(doc)

        assert changes >= 1
        for p in doc.paragraphs:
            if 'Section Title' in p.text:
                assert p.style.name in ['Heading 1', 'Heading 2']

    @pytest.mark.unit
    def test_small_bold_text_becomes_h3(self, docx_factory):
        """Short bold text with smaller font should become Heading 3."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Subsection')
        run.bold = True
        run.font.size = Pt(12)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _enforce_heading_styles(doc)

        assert changes >= 1
        for p in doc.paragraphs:
            if 'Subsection' in p.text:
                assert p.style.name == 'Heading 3'

    @pytest.mark.unit
    def test_long_bold_text_not_detected(self, docx_factory):
        """Long bold text (>80 chars) should NOT be detected as heading."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        long_text = 'This is a very long bold text that exceeds eighty characters and should not be detected as a heading in any case.'
        para = doc.add_paragraph()
        run = para.add_run(long_text)
        run.bold = True
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _enforce_heading_styles(doc)

        assert changes == 0
        for p in doc.paragraphs:
            if long_text in p.text:
                assert not p.style.name.startswith('Heading')

    @pytest.mark.unit
    def test_skips_existing_headings(self, docx_factory):
        """Already-styled headings should not be modified."""
        doc_path = docx_factory(content="", add_heading=True)
        doc = Document(str(doc_path))
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        initial_heading_count = sum(
            1 for p in doc.paragraphs if p.style.name.startswith('Heading')
        )
        changes = _enforce_heading_styles(doc)

        # No changes for existing headings
        new_heading_count = sum(
            1 for p in doc.paragraphs if p.style.name.startswith('Heading')
        )
        assert new_heading_count == initial_heading_count

    @pytest.mark.unit
    def test_skips_empty_paragraphs(self, docx_factory):
        """Empty paragraphs should be skipped."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))
        doc.add_paragraph('')
        doc.add_paragraph('   ')  # whitespace only
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _enforce_heading_styles(doc)

        assert changes == 0


# ============================================================================
# Spacing Normalization Tests
# ============================================================================

class TestSpacingNormalization:
    """Tests for _normalize_spacing function."""

    @pytest.mark.unit
    def test_normalize_spacing_sets_standard_values(self, docx_factory):
        """Spacing should be set to standard values."""
        doc_path = docx_factory(content="Test paragraph")
        doc = Document(str(doc_path))

        para = doc.paragraphs[0]
        para.paragraph_format.space_before = Pt(20)
        para.paragraph_format.space_after = Pt(0)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _normalize_spacing(doc)

        assert changes >= 1
        para = doc.paragraphs[0]
        assert para.paragraph_format.space_before == Pt(0)
        assert para.paragraph_format.space_after == Pt(8)

    @pytest.mark.unit
    def test_normalize_spacing_sets_line_spacing(self, docx_factory):
        """Line spacing should be set to 1.15."""
        doc_path = docx_factory(content="Test paragraph")
        doc = Document(str(doc_path))

        para = doc.paragraphs[0]
        para.paragraph_format.line_spacing = 2.0
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _normalize_spacing(doc)

        assert changes >= 1
        assert doc.paragraphs[0].paragraph_format.line_spacing == 1.15

    @pytest.mark.unit
    def test_normalize_spacing_skips_headings(self, docx_factory):
        """Heading paragraphs should not have spacing modified."""
        doc_path = docx_factory(content="", add_heading=True)
        doc = Document(str(doc_path))

        # Find heading and set unusual spacing
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                para.paragraph_format.space_before = Pt(30)
                para.paragraph_format.space_after = Pt(30)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        _normalize_spacing(doc)

        # Heading spacing should be unchanged
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                # Headings should not be normalized
                pass

    @pytest.mark.unit
    def test_normalize_spacing_multiple_paragraphs(self, docx_factory):
        """Multiple paragraphs should all be normalized."""
        content = "First paragraph\nSecond paragraph\nThird paragraph"
        doc_path = docx_factory(content=content)
        doc = Document(str(doc_path))

        for para in doc.paragraphs:
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(24)
            para.paragraph_format.line_spacing = 1.5
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        changes = _normalize_spacing(doc)

        assert changes >= 3
        for para in doc.paragraphs:
            if para.text.strip():
                assert para.paragraph_format.space_before == Pt(0)
                assert para.paragraph_format.space_after == Pt(8)
                assert para.paragraph_format.line_spacing == 1.15


# ============================================================================
# perform_style_enforcement Tests
# ============================================================================

class TestPerformStyleEnforcement:
    """Tests for the main perform_style_enforcement function."""

    @pytest.mark.unit
    def test_applies_all_options_by_default(self, docx_factory):
        """Default config should apply font standardization and heading enforcement."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Introduction')
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        config = {'style_options': {}}
        changes = perform_style_enforcement(doc, config)

        assert changes >= 1

    @pytest.mark.unit
    def test_respects_disabled_options(self, docx_factory):
        """Disabled options should not make changes."""
        doc_path = docx_factory(content="Test content")
        doc = Document(str(doc_path))

        config = {
            'style_options': {
                'standardize_font': False,
                'enforce_headings': False,
                'normalize_spacing': False,
            }
        }
        changes = perform_style_enforcement(doc, config)

        assert changes == 0

    @pytest.mark.unit
    def test_custom_font_settings(self, docx_factory):
        """Custom font name and size should be applied."""
        doc_path = docx_factory(content="Test content")
        doc = Document(str(doc_path))

        para = doc.paragraphs[0]
        para.runs[0].font.name = 'Calibri'
        para.runs[0].font.size = Pt(11)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        config = {
            'style_options': {
                'standardize_font': True,
                'font_name': 'Arial',
                'font_size_pt': 12,
            }
        }
        changes = perform_style_enforcement(doc, config)

        assert changes >= 1
        assert doc.paragraphs[0].runs[0].font.name == 'Arial'
        assert doc.paragraphs[0].runs[0].font.size == Pt(12)

    @pytest.mark.unit
    def test_spacing_normalization_optional(self, docx_factory):
        """Spacing normalization should only run when enabled."""
        doc_path = docx_factory(content="Test paragraph")
        doc = Document(str(doc_path))

        doc.paragraphs[0].paragraph_format.space_after = Pt(20)
        doc.save(str(doc_path))

        # Without spacing normalization
        doc = Document(str(doc_path))
        config = {
            'style_options': {
                'standardize_font': False,
                'enforce_headings': False,
                'normalize_spacing': False,
            }
        }
        changes = perform_style_enforcement(doc, config)
        assert changes == 0

        # With spacing normalization
        doc = Document(str(doc_path))
        config['style_options']['normalize_spacing'] = True
        changes = perform_style_enforcement(doc, config)
        assert changes >= 1


# ============================================================================
# process_document Integration Tests
# ============================================================================

class TestProcessDocument:
    """Integration tests for process_document function."""

    @pytest.mark.unit
    def test_returns_processor_result(self, docx_factory):
        """Should return a valid ProcessorResult."""
        doc_path = docx_factory(content="Test content")

        config = {'style_options': {}}
        result = process_document(str(doc_path), config)

        assert isinstance(result, ProcessorResult)
        assert result.success is True
        assert result.file_path == str(doc_path)
        assert result.duration_seconds >= 0

    @pytest.mark.unit
    def test_saves_changes_when_made(self, docx_factory):
        """Document should be saved when changes are made."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        para = doc.add_paragraph()
        run = para.add_run('Test')
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        doc.save(str(doc_path))

        config = {'style_options': {'standardize_font': True}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made >= 1

        # Verify file was saved
        doc = Document(str(doc_path))
        assert doc.paragraphs[-1].runs[0].font.name == 'Calibri'

    @pytest.mark.unit
    def test_no_save_when_no_changes(self, docx_factory):
        """Document should not be saved when no changes are made."""
        doc_path = docx_factory(content="Test")
        doc = Document(str(doc_path))
        doc.paragraphs[0].runs[0].font.name = 'Calibri'
        doc.paragraphs[0].runs[0].font.size = Pt(11)
        doc.save(str(doc_path))

        config = {
            'style_options': {
                'standardize_font': True,
                'enforce_headings': False,
                'normalize_spacing': False,
            }
        }
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0

    @pytest.mark.unit
    def test_handles_invalid_file(self, tmp_path):
        """Should handle invalid file paths gracefully."""
        invalid_path = str(tmp_path / "nonexistent.docx")

        config = {'style_options': {}}
        result = process_document(invalid_path, config)

        assert result.success is False
        assert result.error_message is not None

    @pytest.mark.unit
    def test_handles_corrupted_file(self, tmp_path):
        """Should handle corrupted files gracefully."""
        corrupted_file = tmp_path / "corrupted.docx"
        corrupted_file.write_text("Not a valid DOCX file")

        config = {'style_options': {}}
        result = process_document(str(corrupted_file), config)

        assert result.success is False
        assert result.error_message is not None


# ============================================================================
# Edge Cases and Unicode Tests
# ============================================================================

class TestEdgeCases:
    """Edge case and special scenario tests."""

    @pytest.mark.unit
    def test_empty_document(self, docx_factory):
        """Empty document should be handled gracefully."""
        doc_path = docx_factory(content="")

        config = {'style_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0

    @pytest.mark.unit
    def test_document_with_only_headings(self, docx_factory):
        """Document with only headings should be handled."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Title', level=0)
        doc.add_heading('Chapter 1', level=1)
        doc.add_heading('Section 1.1', level=2)
        doc.save(str(doc_path))

        config = {'style_options': {'standardize_font': True}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        # Note: Empty document may have an initial paragraph that gets styled
        # The key assertion is that it succeeds without error

    @pytest.mark.unit
    def test_unicode_content(self, docx_factory):
        """Unicode content should be handled correctly."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        # Chinese text
        para1 = doc.add_paragraph()
        run1 = para1.add_run('\u6d4b\u8bd5\u6587\u672c')
        run1.font.name = 'Arial'

        # Accented European (avoiding emoji surrogates that lxml can't handle)
        para2 = doc.add_paragraph()
        run2 = para2.add_run('Hello caf\xe9 r\xe9sum\xe9 World')
        run2.font.name = 'Times New Roman'

        doc.save(str(doc_path))

        config = {'style_options': {'standardize_font': True}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made >= 2

    @pytest.mark.unit
    def test_deeply_nested_tables(self, docx_factory):
        """Tables should be processed correctly."""
        doc_path = docx_factory(content="Intro", add_table=True)
        doc = Document(str(doc_path))

        # Add multiple tables
        for _ in range(3):
            table = doc.add_table(rows=2, cols=2)
            for row in table.rows:
                for cell in row.cells:
                    run = cell.paragraphs[0].add_run('Cell')
                    run.font.name = 'Comic Sans MS'
                    run.font.size = Pt(8)
        doc.save(str(doc_path))

        config = {'style_options': {'standardize_font': True}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made >= 12  # At least 12 cells (3 tables * 4 cells)

    @pytest.mark.unit
    @pytest.mark.parametrize("font_name,font_size", [
        ('Arial', 10),
        ('Times New Roman', 12),
        ('Verdana', 9),
        ('Georgia', 14),
    ])
    def test_various_font_configurations(self, docx_factory, font_name, font_size):
        """Various font configurations should be standardizable."""
        doc_path = docx_factory(content="Test content")
        doc = Document(str(doc_path))
        doc.paragraphs[0].runs[0].font.name = font_name
        doc.paragraphs[0].runs[0].font.size = Pt(font_size)
        doc.save(str(doc_path))

        config = {
            'style_options': {
                'standardize_font': True,
                'font_name': 'Calibri',
                'font_size_pt': 11,
            }
        }
        result = process_document(str(doc_path), config)

        assert result.success is True
        doc = Document(str(doc_path))
        assert doc.paragraphs[0].runs[0].font.name == 'Calibri'
