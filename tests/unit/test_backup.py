"""Unit tests for BackupManager."""

import shutil
import time
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from src.core.backup import BackupManager
from src.database.db_manager import DatabaseManager


@pytest.fixture
def db_manager(tmp_path):
    """Provide DatabaseManager for testing."""
    db_path = tmp_path / "test.db"
    return DatabaseManager(db_path)


@pytest.fixture
def backup_dir(tmp_path):
    """Provide temporary backup directory."""
    return tmp_path / "backups"


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample .docx file for testing."""
    doc_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has multiple paragraphs.")
    doc.save(doc_path)
    return doc_path


@pytest.fixture
def unicode_docx(tmp_path):
    """Create a .docx file with Unicode filename."""
    doc_path = tmp_path / "测试_文档_emoji_🎉.docx"
    doc = Document()
    doc.add_paragraph("Unicode test document with 中文 and emoji 😀")
    doc.save(doc_path)
    return doc_path


@pytest.fixture
def corrupt_docx(tmp_path):
    """Create a corrupt .docx file (not valid docx format)."""
    doc_path = tmp_path / "corrupt.docx"
    with open(doc_path, 'w', encoding='utf-8') as f:
        f.write("This is not a valid docx file")
    return doc_path


class TestBackupManagerInitialization:
    """Test suite for BackupManager initialization."""

    def test_initialization(self, backup_dir, db_manager):
        """Test BackupManager initializes with correct parameters."""
        manager = BackupManager(backup_dir, db_manager, retention_days=14)

        assert manager.backup_dir == backup_dir
        assert manager.db_manager == db_manager
        assert manager.retention_days == 14

    def test_backup_directory_created(self, backup_dir, db_manager):
        """Test backup directory is created automatically."""
        # Directory should not exist initially
        assert not backup_dir.exists()

        # Creating BackupManager should create directory
        manager = BackupManager(backup_dir, db_manager)

        assert backup_dir.exists()
        assert backup_dir.is_dir()

    def test_nested_backup_directory(self, tmp_path, db_manager):
        """Test nested backup directory creation."""
        nested_dir = tmp_path / "level1" / "level2" / "backups"

        manager = BackupManager(nested_dir, db_manager)

        assert nested_dir.exists()


class TestBackupCreation:
    """Test suite for backup creation functionality."""

    def test_create_backup_success(self, sample_docx, backup_dir, db_manager):
        """Test successful backup creation."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)

        # Verify backup created
        assert backup_path is not None
        assert backup_path.exists()

        # Verify backup in daily subdirectory
        today = datetime.now().strftime('%Y-%m-%d')
        assert backup_path.parent.name == today

        # Verify backup filename format
        assert sample_docx.stem in backup_path.name
        assert ".backup.docx" in backup_path.name

    def test_backup_naming_format(self, sample_docx, backup_dir, db_manager):
        """Test backup filename follows correct format."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)

        # Format: {original_name}.{YYYYMMDD_HHMMSS}.backup.docx
        filename = backup_path.name
        parts = filename.split('.')

        assert len(parts) >= 4  # name, timestamp, backup, docx
        assert parts[0] == sample_docx.stem
        assert parts[-2] == "backup"
        assert parts[-1] == "docx"

        # Verify timestamp format (YYYYMMDD_HHMMSS)
        timestamp = parts[1]
        assert len(timestamp) == 15  # YYYYMMDD_HHMMSS
        assert timestamp[8] == '_'

    def test_backup_content_preserved(self, sample_docx, backup_dir, db_manager):
        """Test backup preserves original document content."""
        manager = BackupManager(backup_dir, db_manager)

        # Read original content
        original_doc = Document(sample_docx)
        original_text = [p.text for p in original_doc.paragraphs]

        # Create backup
        backup_path = manager.create_backup(sample_docx)

        # Read backup content
        backup_doc = Document(backup_path)
        backup_text = [p.text for p in backup_doc.paragraphs]

        # Verify content matches
        assert original_text == backup_text

    def test_backup_database_record(self, sample_docx, backup_dir, db_manager):
        """Test backup is recorded in database."""
        manager = BackupManager(backup_dir, db_manager)

        # Create backup without job_id
        backup_path = manager.create_backup(sample_docx)

        # Verify database record
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()

        assert row is not None
        assert row["original_path"] == str(sample_docx.resolve())
        assert row["backup_path"] == str(backup_path.resolve())
        assert row["job_id"] is None  # No job_id provided
        assert row["restored"] == 0
        assert row["size_bytes"] > 0

    def test_backup_with_unicode_filename(self, unicode_docx, backup_dir, db_manager):
        """Test backup with Unicode characters in filename."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(unicode_docx)

        # Verify backup created
        assert backup_path is not None
        assert backup_path.exists()

        # Verify can read backup
        doc = Document(backup_path)
        assert len(doc.paragraphs) > 0

    def test_backup_multiple_files(self, tmp_path, backup_dir, db_manager):
        """Test creating multiple backups."""
        manager = BackupManager(backup_dir, db_manager)

        # Create multiple test files
        files = []
        for i in range(5):
            doc_path = tmp_path / f"doc_{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.save(doc_path)
            files.append(doc_path)

        # Create backups
        backup_paths = []
        for file_path in files:
            backup_path = manager.create_backup(file_path)
            backup_paths.append(backup_path)
            # Small delay to ensure unique timestamps
            time.sleep(0.01)

        # Verify all backups created
        assert all(bp is not None for bp in backup_paths)
        assert all(bp.exists() for bp in backup_paths)

        # Verify all different
        assert len(set(backup_paths)) == 5

    def test_backup_nonexistent_file(self, backup_dir, db_manager):
        """Test backup fails gracefully for nonexistent file."""
        manager = BackupManager(backup_dir, db_manager)

        nonexistent = Path("nonexistent.docx")
        backup_path = manager.create_backup(nonexistent)

        assert backup_path is None

    def test_backup_with_job_id(self, sample_docx, backup_dir, db_manager):
        """Test backup associated with job ID."""
        manager = BackupManager(backup_dir, db_manager)

        # Create a valid job first (foreign key requirement)
        job_id = db_manager.generate_id()
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO jobs (id, created_at, operation, config, status)
                VALUES (?, ?, ?, ?, ?)
            """, (job_id, datetime.now().isoformat(), "test", "{}", "pending"))
            conn.commit()

        # Create backup with job_id
        backup_path = manager.create_backup(sample_docx, job_id=job_id)

        # Verify job_id in database
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT job_id FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()

        assert row["job_id"] == job_id

    def test_backup_without_job_id(self, sample_docx, backup_dir, db_manager):
        """Test backup can be created without job ID."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)

        # Verify job_id is None in database
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT job_id FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()

        assert row["job_id"] is None


class TestBackupValidation:
    """Test suite for backup validation."""

    def test_validate_valid_backup(self, sample_docx, backup_dir, db_manager):
        """Test validation succeeds for valid backup."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)

        # Validation should succeed (implicit in successful create_backup)
        assert backup_path is not None
        assert backup_path.exists()

    def test_validate_corrupt_backup(self, corrupt_docx, backup_dir, db_manager):
        """Test validation fails for corrupt file."""
        manager = BackupManager(backup_dir, db_manager)

        # Validation should detect corrupt file
        is_valid = manager._validate_backup(corrupt_docx)

        assert is_valid is False

    def test_create_backup_removes_invalid(self, tmp_path, backup_dir, db_manager):
        """Test that invalid backups are cleaned up."""
        manager = BackupManager(backup_dir, db_manager)

        # Create a file that will fail validation
        # We'll mock the validation by creating a corrupt file after copy
        # For this test, we just verify the behavior with corrupt source
        corrupt_path = tmp_path / "corrupt.docx"
        with open(corrupt_path, 'w', encoding='utf-8') as f:
            f.write("Not a valid docx")

        backup_path = manager.create_backup(corrupt_path)

        # Backup should fail and return None
        assert backup_path is None

        # No backup files should exist
        daily_dir = backup_dir / datetime.now().strftime('%Y-%m-%d')
        if daily_dir.exists():
            backup_files = list(daily_dir.glob("*.backup.docx"))
            # Should be empty since validation failed
            assert len(backup_files) == 0


class TestBackupRestoration:
    """Test suite for backup restoration functionality."""

    def test_restore_to_original_location(self, sample_docx, backup_dir, db_manager, tmp_path):
        """Test restoring backup to original location."""
        manager = BackupManager(backup_dir, db_manager)

        # Create backup
        backup_path = manager.create_backup(sample_docx)

        # Delete original
        sample_docx.unlink()
        assert not sample_docx.exists()

        # Restore backup (should restore to original location)
        success = manager.restore_backup(backup_path)

        assert success is True
        assert sample_docx.exists()

        # Verify content matches
        restored_doc = Document(sample_docx)
        backup_doc = Document(backup_path)
        assert [p.text for p in restored_doc.paragraphs] == [p.text for p in backup_doc.paragraphs]

    def test_restore_to_custom_location(self, sample_docx, backup_dir, db_manager, tmp_path):
        """Test restoring backup to custom location."""
        manager = BackupManager(backup_dir, db_manager)

        # Create backup
        backup_path = manager.create_backup(sample_docx)

        # Restore to different location
        target_path = tmp_path / "restored" / "document.docx"
        success = manager.restore_backup(backup_path, target_path)

        assert success is True
        assert target_path.exists()

        # Verify content matches
        restored_doc = Document(target_path)
        backup_doc = Document(backup_path)
        assert [p.text for p in restored_doc.paragraphs] == [p.text for p in backup_doc.paragraphs]

    def test_restore_updates_database(self, sample_docx, backup_dir, db_manager):
        """Test restore updates database with restored flag."""
        manager = BackupManager(backup_dir, db_manager)

        # Create backup
        backup_path = manager.create_backup(sample_docx)

        # Verify initially not restored
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT restored, restored_at FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()
            assert row["restored"] == 0
            assert row["restored_at"] is None

        # Restore backup
        manager.restore_backup(backup_path)

        # Verify marked as restored
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT restored, restored_at FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()
            assert row["restored"] == 1
            assert row["restored_at"] is not None

    def test_restore_nonexistent_backup(self, backup_dir, db_manager, tmp_path):
        """Test restore fails for nonexistent backup."""
        manager = BackupManager(backup_dir, db_manager)

        nonexistent = tmp_path / "nonexistent.backup.docx"
        success = manager.restore_backup(nonexistent)

        assert success is False

    def test_restore_backup_not_in_database(self, sample_docx, backup_dir, db_manager, tmp_path):
        """Test restore fails when backup not found in database."""
        manager = BackupManager(backup_dir, db_manager)

        # Create a backup file manually (not through create_backup)
        manual_backup = tmp_path / "manual_backup.docx"
        shutil.copy2(sample_docx, manual_backup)

        # Try to restore (should fail - not in database)
        success = manager.restore_backup(manual_backup)

        assert success is False

    def test_restore_creates_target_directory(self, sample_docx, backup_dir, db_manager, tmp_path):
        """Test restore creates target directory if it doesn't exist."""
        manager = BackupManager(backup_dir, db_manager)

        # Create backup
        backup_path = manager.create_backup(sample_docx)

        # Restore to nested path that doesn't exist
        target_path = tmp_path / "level1" / "level2" / "level3" / "restored.docx"
        assert not target_path.parent.exists()

        success = manager.restore_backup(backup_path, target_path)

        assert success is True
        assert target_path.exists()
        assert target_path.parent.exists()

    def test_restore_with_unicode_paths(self, unicode_docx, backup_dir, db_manager, tmp_path):
        """Test restore with Unicode paths."""
        manager = BackupManager(backup_dir, db_manager)

        # Create backup of Unicode file
        backup_path = manager.create_backup(unicode_docx)

        # Delete original
        unicode_docx.unlink()

        # Restore
        success = manager.restore_backup(backup_path)

        assert success is True
        assert unicode_docx.exists()

    def test_restore_integration(self, tmp_path, backup_dir, db_manager):
        """Test complete backup and restore workflow."""
        manager = BackupManager(backup_dir, db_manager)

        # Create original document
        original_path = tmp_path / "original.docx"
        original_doc = Document()
        original_doc.add_paragraph("Original content")
        original_doc.add_paragraph("Second paragraph")
        original_doc.save(original_path)

        # Create backup
        backup_path = manager.create_backup(original_path)

        # Modify original
        modified_doc = Document(original_path)
        modified_doc.add_paragraph("Modified content")
        modified_doc.save(original_path)

        # Verify modified
        modified_doc = Document(original_path)
        assert len(modified_doc.paragraphs) == 3

        # Restore from backup
        success = manager.restore_backup(backup_path)

        # Verify restored to original state
        assert success is True
        restored_doc = Document(original_path)
        assert len(restored_doc.paragraphs) == 2
        assert restored_doc.paragraphs[0].text == "Original content"


class TestBackupCleanup:
    """Test suite for backup cleanup functionality."""

    def test_cleanup_old_backups(self, sample_docx, backup_dir, db_manager):
        """Test cleanup removes backups older than retention period."""
        from datetime import timedelta

        manager = BackupManager(backup_dir, db_manager, retention_days=7)

        # Create a backup
        backup_path = manager.create_backup(sample_docx)

        # Manually set created_at to 10 days ago in database
        old_date = (datetime.now() - timedelta(days=10)).isoformat()
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE backups SET created_at = ? WHERE backup_path = ?",
                (old_date, str(backup_path.resolve()))
            )
            conn.commit()

        # Run cleanup
        deleted_count = manager.cleanup_old_backups()

        # Verify backup was deleted
        assert deleted_count == 1
        assert not backup_path.exists()

        # Verify database record removed
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT COUNT(*) as count FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            assert cursor.fetchone()["count"] == 0

    def test_cleanup_keeps_recent_backups(self, sample_docx, backup_dir, db_manager):
        """Test cleanup keeps backups within retention period."""
        manager = BackupManager(backup_dir, db_manager, retention_days=7)

        # Create a recent backup
        backup_path = manager.create_backup(sample_docx)

        # Run cleanup
        deleted_count = manager.cleanup_old_backups()

        # Verify no backups deleted
        assert deleted_count == 0
        assert backup_path.exists()

    def test_cleanup_keeps_restored_backups(self, sample_docx, backup_dir, db_manager):
        """Test cleanup preserves backups that have been restored."""
        from datetime import timedelta

        manager = BackupManager(backup_dir, db_manager, retention_days=7)

        # Create and restore a backup
        backup_path = manager.create_backup(sample_docx)
        manager.restore_backup(backup_path)

        # Set backup to be old
        old_date = (datetime.now() - timedelta(days=10)).isoformat()
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE backups SET created_at = ? WHERE backup_path = ?",
                (old_date, str(backup_path.resolve()))
            )
            conn.commit()

        # Run cleanup
        deleted_count = manager.cleanup_old_backups()

        # Verify restored backup NOT deleted (restored=1)
        assert deleted_count == 0
        assert backup_path.exists()

    def test_cleanup_multiple_backups(self, tmp_path, backup_dir, db_manager):
        """Test cleanup with multiple old backups."""
        from datetime import timedelta

        manager = BackupManager(backup_dir, db_manager, retention_days=7)

        # Create multiple backups
        backup_paths = []
        for i in range(5):
            doc_path = tmp_path / f"doc_{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.save(doc_path)
            backup_paths.append(manager.create_backup(doc_path))
            time.sleep(0.01)

        # Make all backups old
        old_date = (datetime.now() - timedelta(days=10)).isoformat()
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            for backup_path in backup_paths:
                cursor.execute(
                    "UPDATE backups SET created_at = ? WHERE backup_path = ?",
                    (old_date, str(backup_path.resolve()))
                )
            conn.commit()

        # Run cleanup
        deleted_count = manager.cleanup_old_backups()

        # Verify all deleted
        assert deleted_count == 5
        assert all(not bp.exists() for bp in backup_paths)

    def test_cleanup_handles_missing_files(self, sample_docx, backup_dir, db_manager):
        """Test cleanup handles case where backup file already deleted."""
        from datetime import timedelta

        manager = BackupManager(backup_dir, db_manager, retention_days=7)

        # Create backup
        backup_path = manager.create_backup(sample_docx)

        # Manually delete file (but keep database record)
        backup_path.unlink()

        # Make backup old in database
        old_date = (datetime.now() - timedelta(days=10)).isoformat()
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE backups SET created_at = ? WHERE backup_path = ?",
                (old_date, str(backup_path.resolve()))
            )
            conn.commit()

        # Run cleanup (should not error)
        deleted_count = manager.cleanup_old_backups()

        # File already gone, but database record should be removed
        assert deleted_count == 0  # No files actually deleted


class TestBackupStorageMetrics:
    """Test suite for backup storage metrics."""

    def test_storage_usage_empty(self, backup_dir, db_manager):
        """Test storage usage with no backups."""
        manager = BackupManager(backup_dir, db_manager)

        usage = manager.get_backup_storage_usage()

        assert usage == 0

    def test_storage_usage_single_backup(self, sample_docx, backup_dir, db_manager):
        """Test storage usage with single backup."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)
        expected_size = backup_path.stat().st_size

        usage = manager.get_backup_storage_usage()

        assert usage == expected_size

    def test_storage_usage_multiple_backups(self, tmp_path, backup_dir, db_manager):
        """Test storage usage with multiple backups."""
        manager = BackupManager(backup_dir, db_manager)

        # Create multiple backups of different sizes
        total_expected = 0
        for i in range(3):
            doc_path = tmp_path / f"doc_{i}.docx"
            doc = Document()
            for j in range(i + 1):  # Varying content
                doc.add_paragraph(f"Document {i} paragraph {j}")
            doc.save(doc_path)

            backup_path = manager.create_backup(doc_path)
            total_expected += backup_path.stat().st_size
            time.sleep(0.01)

        usage = manager.get_backup_storage_usage()

        assert usage == total_expected

    def test_storage_usage_ignores_other_files(self, backup_dir, db_manager):
        """Test storage usage ignores non-backup files."""
        manager = BackupManager(backup_dir, db_manager)

        # Create a non-backup file in backup directory
        other_file = backup_dir / "other.txt"
        with open(other_file, 'w', encoding='utf-8') as f:
            f.write("This is not a backup file")

        usage = manager.get_backup_storage_usage()

        # Should not count other.txt
        assert usage == 0


class TestBackupEdgeCases:
    """Test suite for edge cases and error handling."""

    def test_backup_with_spaces_in_path(self, tmp_path, db_manager):
        """Test backup with spaces in directory path."""
        backup_dir = tmp_path / "backup directory with spaces"
        manager = BackupManager(backup_dir, db_manager)

        # Create test file
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(doc_path)

        backup_path = manager.create_backup(doc_path)

        assert backup_path is not None
        assert backup_path.exists()

    def test_backup_preserves_metadata(self, sample_docx, backup_dir, db_manager):
        """Test shutil.copy2 preserves file metadata."""
        manager = BackupManager(backup_dir, db_manager)

        original_stat = sample_docx.stat()
        backup_path = manager.create_backup(sample_docx)
        backup_stat = backup_path.stat()

        # Modification time should be preserved (within 1 second tolerance)
        assert abs(original_stat.st_mtime - backup_stat.st_mtime) < 1.0

    def test_backup_file_size_recorded(self, sample_docx, backup_dir, db_manager):
        """Test file size is recorded correctly in database."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)
        actual_size = backup_path.stat().st_size

        # Check database
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT size_bytes FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()

        assert row["size_bytes"] == actual_size

    def test_backup_absolute_paths_stored(self, sample_docx, backup_dir, db_manager):
        """Test that absolute paths are stored in database."""
        manager = BackupManager(backup_dir, db_manager)

        backup_path = manager.create_backup(sample_docx)

        # Check database stores absolute paths
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT original_path, backup_path FROM backups WHERE backup_path = ?",
                (str(backup_path.resolve()),)
            )
            row = cursor.fetchone()

        original_path = Path(row["original_path"])
        stored_backup_path = Path(row["backup_path"])

        assert original_path.is_absolute()
        assert stored_backup_path.is_absolute()
