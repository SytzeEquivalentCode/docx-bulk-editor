"""Unit tests for search/replace pattern compilation."""

import re
import pytest
from src.processors.search_replace import _compile_pattern


def test_literal_case_sensitive():
    """Test literal search with case sensitivity."""
    pattern = _compile_pattern('Test', use_regex=False, case_sensitive=True, whole_words=False)

    assert pattern.search('Test') is not None
    assert pattern.search('test') is None  # Should not match different case
    assert pattern.search('Testing') is not None  # Should match as substring


def test_literal_case_insensitive():
    """Test literal search without case sensitivity."""
    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)

    assert pattern.search('test') is not None
    assert pattern.search('Test') is not None  # Should match different case
    assert pattern.search('TEST') is not None  # Should match all caps
    assert pattern.search('Testing') is not None  # Should match as substring


def test_literal_whole_words_only():
    """Test literal search with whole words only."""
    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=True)

    assert pattern.search('test') is not None  # Exact match
    assert pattern.search('a test here') is not None  # Word in sentence
    assert pattern.search('testing') is None  # Should not match as substring
    assert pattern.search('pretest') is None  # Should not match as substring


def test_literal_special_chars_escaped():
    """Test literal search escapes special regex characters."""
    # Test with various special regex characters
    special_chars = r'.*+?[]{}()^$|\\'

    pattern = _compile_pattern(special_chars, use_regex=False, case_sensitive=True, whole_words=False)

    # Should match the literal string, not interpret as regex
    assert pattern.search(special_chars) is not None
    assert pattern.search('anything') is None  # .* should not match anything


def test_literal_parentheses_escaped():
    """Test literal search with parentheses (common in legal text)."""
    pattern = _compile_pattern('(a)', use_regex=False, case_sensitive=True, whole_words=False)

    assert pattern.search('(a)') is not None
    assert pattern.search('a') is None  # Should not match without parentheses


def test_literal_brackets_escaped():
    """Test literal search with brackets."""
    pattern = _compile_pattern('[test]', use_regex=False, case_sensitive=True, whole_words=False)

    assert pattern.search('[test]') is not None
    assert pattern.search('test') is None  # Should not match without brackets


def test_regex_basic_pattern():
    """Test regex search with basic pattern."""
    pattern = _compile_pattern(r'\d+', use_regex=True, case_sensitive=True, whole_words=False)

    assert pattern.search('123') is not None
    assert pattern.search('abc') is None
    assert pattern.search('test 456 here') is not None


def test_regex_case_insensitive():
    """Test regex search with case insensitivity."""
    pattern = _compile_pattern(r'[a-z]+', use_regex=True, case_sensitive=False, whole_words=False)

    assert pattern.search('abc') is not None
    assert pattern.search('ABC') is not None  # Should match with IGNORECASE flag
    assert pattern.search('123') is None


def test_regex_case_sensitive():
    """Test regex search with case sensitivity."""
    pattern = _compile_pattern(r'[a-z]+', use_regex=True, case_sensitive=True, whole_words=False)

    assert pattern.search('abc') is not None
    assert pattern.search('ABC') is None  # Should not match uppercase


def test_regex_capture_groups():
    """Test regex search with capture groups."""
    pattern = _compile_pattern(r'(\w+)@(\w+)\.com', use_regex=True, case_sensitive=False, whole_words=False)

    match = pattern.search('user@example.com')
    assert match is not None
    assert match.group(1) == 'user'
    assert match.group(2) == 'example'


def test_regex_word_boundaries():
    """Test regex search with word boundaries (user-defined)."""
    # In regex mode, user controls word boundaries explicitly
    pattern = _compile_pattern(r'\btest\b', use_regex=True, case_sensitive=False, whole_words=False)

    assert pattern.search('test') is not None
    assert pattern.search('a test here') is not None
    assert pattern.search('testing') is None  # Should not match
    assert pattern.search('pretest') is None  # Should not match


def test_regex_alternation():
    """Test regex search with alternation (OR)."""
    pattern = _compile_pattern(r'cat|dog', use_regex=True, case_sensitive=False, whole_words=False)

    assert pattern.search('I have a cat') is not None
    assert pattern.search('I have a dog') is not None
    assert pattern.search('I have a bird') is None


def test_whole_words_ignored_in_regex_mode():
    """Test that whole_words parameter is ignored in regex mode."""
    # whole_words only applies to literal mode
    pattern = _compile_pattern(r'test', use_regex=True, case_sensitive=False, whole_words=True)

    # Should match as substring since regex mode ignores whole_words
    assert pattern.search('testing') is not None


def test_pattern_subn_replacement():
    """Test that compiled pattern can perform substitution."""
    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)

    text = 'This is a test. Testing 123.'
    new_text, count = pattern.subn('REPLACED', text)

    assert count == 2  # Should match 'test' and 'Test' in 'Testing'
    assert 'REPLACED' in new_text


def test_pattern_subn_whole_words():
    """Test substitution with whole words only."""
    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=True)

    text = 'This is a test. Testing 123.'
    new_text, count = pattern.subn('REPLACED', text)

    assert count == 1  # Should only match 'test', not 'Test' in 'Testing'
    assert 'This is a REPLACED. Testing 123.' == new_text


def test_unicode_search():
    """Test pattern compilation with Unicode characters."""
    pattern = _compile_pattern('测试', use_regex=False, case_sensitive=True, whole_words=False)

    assert pattern.search('这是测试文本') is not None
    assert pattern.search('This is a test') is None


def test_emoji_search():
    """Test pattern compilation with emoji characters."""
    pattern = _compile_pattern('😀', use_regex=False, case_sensitive=True, whole_words=False)

    assert pattern.search('Hello 😀 World') is not None
    assert pattern.search('Hello World') is None


def test_empty_search_term():
    """Test pattern compilation with empty search term."""
    # Empty string should compile but match everywhere
    pattern = _compile_pattern('', use_regex=False, case_sensitive=True, whole_words=False)

    # Empty pattern matches at every position
    text = 'test'
    matches = list(pattern.finditer(text))
    assert len(matches) > 0


def test_multiline_regex():
    """Test regex pattern with multiline text."""
    pattern = _compile_pattern(r'^test', use_regex=True, case_sensitive=False, whole_words=False)

    # Without MULTILINE flag, ^ only matches start of string
    assert pattern.search('test\nmore') is not None
    assert pattern.search('start\ntest') is None


def test_process_document_success(docx_factory):
    """Test process_document end-to-end success case."""
    from src.processors.search_replace import process_document
    from docx import Document

    # Create test document
    doc_path = docx_factory(content="Hello world\nHello again")

    config = {
        'search_term': 'Hello',
        'replace_term': 'Goodbye',
        'use_regex': False,
        'case_sensitive': True,
        'whole_words': False,
        'search_body': True,
        'search_tables': False,
        'search_headers': False,
        'search_footers': False
    }

    result = process_document(str(doc_path), config)

    assert result.success is True
    assert result.changes_made == 2
    assert result.duration_seconds > 0

    # Verify document was actually modified
    doc = Document(str(doc_path))
    text = '\n'.join([para.text for para in doc.paragraphs])
    assert 'Goodbye' in text
    assert 'Hello' not in text


def test_process_document_file_not_found():
    """Test process_document handles missing file gracefully."""
    from src.processors.search_replace import process_document

    config = {
        'find': 'test',
        'replace': 'sample',
        'use_regex': False,
        'case_sensitive': True,
        'whole_words': False
    }

    result = process_document('/nonexistent/file.docx', config)

    assert result.success is False
    assert result.changes_made == 0
    assert result.error_message is not None


def test_empty_run_text(docx_factory):
    """Test replacing in paragraphs with empty runs (edge case for line 168 branch)."""
    from src.processors.search_replace import _replace_in_paragraph, _compile_pattern
    from docx import Document

    # Create document with content
    doc_path = docx_factory(content="Test paragraph")
    doc = Document(str(doc_path))
    para = doc.paragraphs[0]

    # Add a run with empty text to trigger the if run.text condition
    para.add_run("")  # Empty run

    pattern = _compile_pattern('Test', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(para, pattern, 'Sample')

    assert count == 1
    # Verify replacement happened despite empty run
    assert 'Sample' in para.text
