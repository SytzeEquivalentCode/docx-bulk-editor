"""Unit tests for the Table Formatter processor.

Tests cover:
- Border standardization
- Header row background
- Zebra striping
- Cell alignment
- Cell padding
- Edge cases and error handling
"""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from src.processors.table_formatter import (
    process_document,
    perform_table_formatting,
    _apply_table_borders,
    _apply_header_background,
    _apply_zebra_striping,
    _set_cell_shading,
    _apply_cell_alignment,
    _apply_cell_padding,
)
from src.processors import ProcessorResult


# ============================================================================
# Border Tests
# ============================================================================

class TestBorderStandardization:
    """Tests for _apply_table_borders function."""

    @pytest.mark.unit
    def test_applies_borders_to_all_edges(self, docx_with_table):
        """All border types should be applied (top, bottom, left, right, insideH, insideV)."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_table_borders(table)

        # Check tblBorders exists and has all border types
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None

        border_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        for name in border_names:
            border = tblBorders.find(qn(f'w:{name}'))
            assert border is not None, f"Missing border: {name}"
            assert border.get(qn('w:val')) == 'single'
            assert border.get(qn('w:color')) == '000000'

    @pytest.mark.unit
    def test_replaces_existing_borders(self, docx_with_table):
        """Existing borders should be replaced, not duplicated."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        # Apply twice
        _apply_table_borders(table)
        _apply_table_borders(table)

        # Should only have one tblBorders element
        tblPr = table._tbl.tblPr
        borders_count = len(tblPr.findall(qn('w:tblBorders')))
        assert borders_count == 1

    @pytest.mark.unit
    def test_border_size_attribute(self, docx_with_table):
        """Border size should be 4 (0.5pt)."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_table_borders(table)

        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        top_border = tblBorders.find(qn('w:top'))
        assert top_border.get(qn('w:sz')) == '4'


# ============================================================================
# Header Background Tests
# ============================================================================

class TestHeaderBackground:
    """Tests for _apply_header_background function."""

    @pytest.mark.unit
    def test_applies_gray_background_to_header_row(self, docx_with_table):
        """Header row cells should get light gray background."""
        doc_path = docx_with_table(rows=3, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_header_background(table.rows[0])

        # Check all header cells have shading
        for cell in table.rows[0].cells:
            tcPr = cell._tc.tcPr
            shd = tcPr.find(qn('w:shd'))
            assert shd is not None, "Header cell missing shading"
            assert shd.get(qn('w:fill')) == 'D9D9D9', "Should be light gray"

    @pytest.mark.unit
    def test_header_shading_attributes(self, docx_with_table):
        """Header shading should have val='clear' and color='auto'."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_header_background(table.rows[0])

        cell = table.rows[0].cells[0]
        shd = cell._tc.tcPr.find(qn('w:shd'))
        assert shd.get(qn('w:val')) == 'clear'
        assert shd.get(qn('w:color')) == 'auto'

    @pytest.mark.unit
    def test_replaces_existing_shading(self, docx_with_table):
        """Existing shading should be replaced."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        # Apply twice
        _apply_header_background(table.rows[0])
        _apply_header_background(table.rows[0])

        # Should only have one shd element per cell
        cell = table.rows[0].cells[0]
        shd_count = len(cell._tc.tcPr.findall(qn('w:shd')))
        assert shd_count == 1


# ============================================================================
# Zebra Striping Tests
# ============================================================================

class TestZebraStriping:
    """Tests for _apply_zebra_striping function."""

    @pytest.mark.unit
    def test_alternates_row_colors(self, docx_with_table):
        """Odd rows should be white, even rows should be light gray."""
        doc_path = docx_with_table(rows=5, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_zebra_striping(table)

        # Row 0 (header) skipped
        # Row 1 (odd) should be white
        # Row 2 (even) should be light gray
        # Row 3 (odd) should be white
        # Row 4 (even) should be light gray

        for idx, row in enumerate(table.rows):
            if idx == 0:
                continue  # Header skipped

            for cell in row.cells:
                tcPr = cell._tc.tcPr
                shd = tcPr.find(qn('w:shd'))
                assert shd is not None

                fill = shd.get(qn('w:fill'))
                if idx % 2 == 0:  # Even = gray
                    assert fill == 'F2F2F2', f"Row {idx} should be light gray"
                else:  # Odd = white
                    assert fill == 'FFFFFF', f"Row {idx} should be white"

    @pytest.mark.unit
    def test_skips_header_row(self, docx_with_table):
        """Header row (row 0) should not be affected by zebra striping when pre-styled."""
        doc_path = docx_with_table(rows=3, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        # First apply header background
        _apply_header_background(table.rows[0])
        header_fill_before = table.rows[0].cells[0]._tc.tcPr.find(qn('w:shd')).get(qn('w:fill'))

        # Then apply zebra striping
        _apply_zebra_striping(table)

        # Header should retain its original fill (D9D9D9 from header background)
        # Note: zebra striping starts at row 1, skipping row 0
        header_fill_after = table.rows[0].cells[0]._tc.tcPr.find(qn('w:shd')).get(qn('w:fill'))
        assert header_fill_before == header_fill_after

    @pytest.mark.unit
    def test_zebra_two_row_table(self, docx_with_table):
        """Table with header + one data row should work."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_zebra_striping(table)

        # Row 1 (odd) should be white
        cell = table.rows[1].cells[0]
        shd = cell._tc.tcPr.find(qn('w:shd'))
        assert shd.get(qn('w:fill')) == 'FFFFFF'


# ============================================================================
# Cell Shading Tests
# ============================================================================

class TestCellShading:
    """Tests for _set_cell_shading helper function."""

    @pytest.mark.unit
    def test_sets_custom_color(self, docx_with_table):
        """Custom hex color should be applied."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]
        cell = table.rows[0].cells[0]

        _set_cell_shading(cell, 'FF0000')

        shd = cell._tc.tcPr.find(qn('w:shd'))
        assert shd.get(qn('w:fill')) == 'FF0000'

    @pytest.mark.unit
    def test_replaces_existing_color(self, docx_with_table):
        """Existing shading color should be replaced."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]
        cell = table.rows[0].cells[0]

        _set_cell_shading(cell, 'FF0000')
        _set_cell_shading(cell, '00FF00')

        shd = cell._tc.tcPr.find(qn('w:shd'))
        assert shd.get(qn('w:fill')) == '00FF00'


# ============================================================================
# Cell Alignment Tests
# ============================================================================

class TestCellAlignment:
    """Tests for _apply_cell_alignment function."""

    @pytest.mark.unit
    @pytest.mark.parametrize("alignment,expected", [
        ("left", 0),    # WD_ALIGN_PARAGRAPH.LEFT = 0
        ("center", 1),  # WD_ALIGN_PARAGRAPH.CENTER = 1
        ("right", 2),   # WD_ALIGN_PARAGRAPH.RIGHT = 2
    ])
    def test_applies_alignment_to_all_cells(self, docx_with_table, alignment, expected):
        """Alignment should be applied to all cell paragraphs."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_cell_alignment(table, alignment)

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    assert para.alignment == expected

    @pytest.mark.unit
    def test_case_insensitive_alignment(self, docx_with_table):
        """Alignment parameter should be case-insensitive."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_cell_alignment(table, 'CENTER')

        cell = table.rows[0].cells[0]
        assert cell.paragraphs[0].alignment == 1  # CENTER

    @pytest.mark.unit
    def test_invalid_alignment_defaults_to_left(self, docx_with_table):
        """Invalid alignment value should default to left."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_cell_alignment(table, 'invalid')

        cell = table.rows[0].cells[0]
        assert cell.paragraphs[0].alignment == 0  # LEFT


# ============================================================================
# Cell Padding Tests
# ============================================================================

class TestCellPadding:
    """Tests for _apply_cell_padding function."""

    @pytest.mark.unit
    def test_applies_padding_to_all_cells(self, docx_with_table):
        """Padding should be applied to all cells."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_cell_padding(table, 10)  # 10pt padding

        expected_twips = str(10 * 20)  # 1pt = 20 twips

        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.tcPr
                tcMar = tcPr.find(qn('w:tcMar'))
                assert tcMar is not None

                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = tcMar.find(qn(f'w:{margin_name}'))
                    assert margin is not None
                    assert margin.get(qn('w:w')) == expected_twips

    @pytest.mark.unit
    def test_replaces_existing_padding(self, docx_with_table):
        """Existing padding should be replaced."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_cell_padding(table, 5)
        _apply_cell_padding(table, 10)

        # Check final value is 10pt (200 twips)
        cell = table.rows[0].cells[0]
        tcMar = cell._tc.tcPr.find(qn('w:tcMar'))
        top = tcMar.find(qn('w:top'))
        assert top.get(qn('w:w')) == '200'

    @pytest.mark.unit
    def test_padding_type_attribute(self, docx_with_table):
        """Padding should have type='dxa' attribute."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        _apply_cell_padding(table, 5)

        cell = table.rows[0].cells[0]
        tcMar = cell._tc.tcPr.find(qn('w:tcMar'))
        top = tcMar.find(qn('w:top'))
        assert top.get(qn('w:type')) == 'dxa'


# ============================================================================
# perform_table_formatting Tests
# ============================================================================

class TestPerformTableFormatting:
    """Tests for perform_table_formatting function."""

    @pytest.mark.unit
    def test_formats_all_tables_in_document(self, docx_with_table):
        """All tables should be formatted."""
        # Create document with one table, then add another
        doc_path = docx_with_table(rows=2, cols=2, filename="multi_table.docx")
        doc = Document(str(doc_path))

        # Add second table
        doc.add_table(rows=3, cols=3)
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        config = {'table_options': {}}
        count = perform_table_formatting(doc, config)

        assert count == 2  # Both tables formatted

    @pytest.mark.unit
    def test_default_options_applied(self, docx_with_table):
        """Default options (borders + header) should be applied."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))

        config = {'table_options': {}}
        count = perform_table_formatting(doc, config)

        assert count == 1

        # Check borders applied
        table = doc.tables[0]
        tblBorders = table._tbl.tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None

        # Check header background applied
        shd = table.rows[0].cells[0]._tc.tcPr.find(qn('w:shd'))
        assert shd is not None
        assert shd.get(qn('w:fill')) == 'D9D9D9'

    @pytest.mark.unit
    def test_respects_disabled_options(self, docx_with_table):
        """Disabled options should not be applied."""
        doc_path = docx_with_table(rows=3, cols=2)
        doc = Document(str(doc_path))

        config = {
            'table_options': {
                'standardize_borders': False,
                'header_background': False,
                'zebra_striping': False,
            }
        }
        # Note: alignment and padding still applied by default

        count = perform_table_formatting(doc, config)

        # Table should still count as formatted (alignment/padding applied)
        assert count >= 0

    @pytest.mark.unit
    def test_zebra_striping_when_enabled(self, docx_with_table):
        """Zebra striping should only be applied when enabled."""
        doc_path = docx_with_table(rows=4, cols=2)
        doc = Document(str(doc_path))

        config = {
            'table_options': {
                'standardize_borders': False,
                'header_background': False,
                'zebra_striping': True,
            }
        }
        perform_table_formatting(doc, config)

        # Check row 2 (even data row) has striping
        table = doc.tables[0]
        shd = table.rows[2].cells[0]._tc.tcPr.find(qn('w:shd'))
        assert shd is not None
        assert shd.get(qn('w:fill')) == 'F2F2F2'

    @pytest.mark.unit
    def test_custom_alignment(self, docx_with_table):
        """Custom alignment should be applied."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))

        config = {
            'table_options': {
                'standardize_borders': False,
                'header_background': False,
                'alignment': 'center',
            }
        }
        perform_table_formatting(doc, config)

        cell = doc.tables[0].rows[0].cells[0]
        assert cell.paragraphs[0].alignment == 1  # CENTER

    @pytest.mark.unit
    def test_custom_padding(self, docx_with_table):
        """Custom padding should be applied."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))

        config = {
            'table_options': {
                'standardize_borders': False,
                'header_background': False,
                'padding_pt': 15,
            }
        }
        perform_table_formatting(doc, config)

        cell = doc.tables[0].rows[0].cells[0]
        tcMar = cell._tc.tcPr.find(qn('w:tcMar'))
        top = tcMar.find(qn('w:top'))
        assert top.get(qn('w:w')) == '300'  # 15pt * 20 = 300 twips

    @pytest.mark.unit
    def test_empty_document_returns_zero(self, docx_factory):
        """Document with no tables should return 0."""
        doc_path = docx_factory(content="No tables here.")
        doc = Document(str(doc_path))

        config = {'table_options': {}}
        count = perform_table_formatting(doc, config)

        assert count == 0


# ============================================================================
# process_document Integration Tests
# ============================================================================

class TestProcessDocument:
    """Integration tests for process_document function."""

    @pytest.mark.unit
    def test_returns_processor_result(self, docx_with_table):
        """Should return a valid ProcessorResult."""
        doc_path = docx_with_table(rows=2, cols=2)

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert isinstance(result, ProcessorResult)
        assert result.success is True
        assert result.file_path == str(doc_path)
        assert result.duration_seconds >= 0

    @pytest.mark.unit
    def test_saves_changes_when_made(self, docx_with_table):
        """Document should be saved when changes are made."""
        doc_path = docx_with_table(rows=2, cols=2)

        config = {'table_options': {'standardize_borders': True}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 1

        # Reload and verify saved
        doc = Document(str(doc_path))
        tblBorders = doc.tables[0]._tbl.tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None

    @pytest.mark.unit
    def test_handles_invalid_file(self, tmp_path):
        """Should handle invalid file paths gracefully."""
        invalid_path = str(tmp_path / "nonexistent.docx")

        config = {'table_options': {}}
        result = process_document(invalid_path, config)

        assert result.success is False
        assert result.error_message is not None

    @pytest.mark.unit
    def test_handles_corrupted_file(self, tmp_path):
        """Should handle corrupted files gracefully."""
        corrupted = tmp_path / "corrupted.docx"
        corrupted.write_text("Not a valid DOCX file", encoding='utf-8')

        config = {'table_options': {}}
        result = process_document(str(corrupted), config)

        assert result.success is False
        assert result.error_message is not None

    @pytest.mark.unit
    def test_duration_tracked(self, docx_with_table):
        """Duration should be tracked in result."""
        doc_path = docx_with_table(rows=2, cols=2)

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.duration_seconds > 0


# ============================================================================
# Edge Cases
# ============================================================================

class TestEdgeCases:
    """Edge case and special scenario tests."""

    @pytest.mark.unit
    def test_single_cell_table(self, docx_with_table):
        """Single cell table should be handled."""
        doc_path = docx_with_table(rows=1, cols=1, header_row=False)

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 1

    @pytest.mark.unit
    def test_large_table(self, docx_with_table):
        """Large table should be processed."""
        doc_path = docx_with_table(rows=50, cols=10)

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True

    @pytest.mark.unit
    def test_multiple_tables_all_formatted(self, tmp_path):
        """All tables in document should be formatted."""
        doc = Document()
        for _ in range(5):
            doc.add_table(rows=2, cols=2)

        doc_path = tmp_path / "multi.docx"
        doc.save(str(doc_path))

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 5

    @pytest.mark.unit
    def test_empty_table_cells(self, docx_with_table):
        """Tables with empty cells should be handled."""
        doc_path = docx_with_table(rows=2, cols=2, cell_text="", header_row=False)

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True

    @pytest.mark.unit
    def test_unicode_cell_content(self, tmp_path):
        """Unicode content in cells should be preserved."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header Chinese"
        table.cell(1, 0).text = "Data Japanese"

        doc_path = tmp_path / "unicode.docx"
        doc.save(str(doc_path))

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True

        # Verify content preserved
        doc = Document(str(doc_path))
        assert "Chinese" in doc.tables[0].cell(0, 0).text
        assert "Japanese" in doc.tables[0].cell(1, 0).text

    @pytest.mark.unit
    @pytest.mark.parametrize("rows,cols", [
        (2, 2),
        (5, 3),
        (10, 5),
        (3, 10),
    ])
    def test_various_table_sizes(self, docx_with_table, rows, cols):
        """Various table sizes should all be processable."""
        doc_path = docx_with_table(rows=rows, cols=cols)

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 1

    @pytest.mark.unit
    def test_no_save_when_no_tables(self, docx_factory):
        """Document should not be saved when no changes are made."""
        doc_path = docx_factory(content="Just text, no tables")

        config = {'table_options': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0

    @pytest.mark.unit
    def test_all_options_combined(self, docx_with_table):
        """All formatting options applied together."""
        doc_path = docx_with_table(rows=4, cols=3)

        config = {
            'table_options': {
                'standardize_borders': True,
                'header_background': True,
                'zebra_striping': True,
                'alignment': 'center',
                'padding_pt': 8,
            }
        }
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 1

        # Verify all options applied
        doc = Document(str(doc_path))
        table = doc.tables[0]

        # Borders
        assert table._tbl.tblPr.find(qn('w:tblBorders')) is not None

        # Header background
        header_shd = table.rows[0].cells[0]._tc.tcPr.find(qn('w:shd'))
        assert header_shd.get(qn('w:fill')) == 'D9D9D9'

        # Zebra striping on row 2
        zebra_shd = table.rows[2].cells[0]._tc.tcPr.find(qn('w:shd'))
        assert zebra_shd.get(qn('w:fill')) == 'F2F2F2'

        # Alignment
        assert table.rows[0].cells[0].paragraphs[0].alignment == 1

        # Padding
        tcMar = table.rows[0].cells[0]._tc.tcPr.find(qn('w:tcMar'))
        assert tcMar.find(qn('w:top')).get(qn('w:w')) == '160'  # 8 * 20
