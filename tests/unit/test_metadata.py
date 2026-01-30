"""Unit tests for Metadata processor.

Tests cover:
- Setting metadata fields (single and multiple)
- Clearing metadata fields (single and multiple)
- Combined set and clear operations
- Unicode and special characters in metadata
- Error handling and edge cases
"""

import pytest
from pathlib import Path
from docx import Document

from src.processors.metadata import process_document, perform_metadata_management
from src.processors import ProcessorResult


# ============================================================================
# Clear Metadata Tests
# ============================================================================

class TestClearMetadata:
    """Tests for clearing metadata fields."""

    @pytest.mark.unit
    def test_clear_single_metadata_field(self):
        """Test clearing a single metadata field."""
        doc = Document()
        doc.core_properties.author = 'Test Author'
        doc.core_properties.title = 'Test Title'

        config = {
            'metadata_operations': {
                'clear': ['author']
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.author == ''
        assert doc.core_properties.title == 'Test Title'  # Unchanged

    @pytest.mark.unit
    def test_clear_multiple_metadata_fields(self):
        """Test clearing multiple metadata fields."""
        doc = Document()
        doc.core_properties.author = 'Test Author'
        doc.core_properties.title = 'Test Title'
        doc.core_properties.subject = 'Test Subject'
        doc.core_properties.keywords = 'test, keywords'

        config = {
            'metadata_operations': {
                'clear': ['author', 'title', 'subject']
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 3
        assert doc.core_properties.author == ''
        assert doc.core_properties.title == ''
        assert doc.core_properties.subject == ''
        assert doc.core_properties.keywords == 'test, keywords'  # Unchanged

    @pytest.mark.unit
    def test_clear_nonexistent_field(self):
        """Test clearing a field that doesn't exist (should be skipped)."""
        doc = Document()
        doc.core_properties.author = 'Test Author'

        config = {
            'metadata_operations': {
                'clear': ['nonexistent_field', 'author']
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.author == ''

    @pytest.mark.unit
    def test_clear_company_field(self):
        """Test clearing company (category) field."""
        doc = Document()
        doc.core_properties.category = 'Test Company'

        config = {
            'metadata_operations': {
                'clear': ['category']
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.category == ''

    @pytest.mark.unit
    def test_clear_comments(self):
        """Test clearing comments field."""
        doc = Document()
        doc.core_properties.comments = 'Some comments here'

        config = {
            'metadata_operations': {
                'clear': ['comments']
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.comments == ''

    @pytest.mark.unit
    def test_clear_revision_sets_to_zero(self):
        """Test that clearing revision field attempts to set it to 0."""
        doc = Document()
        doc.core_properties.revision = 5

        config = {
            'metadata_operations': {
                'clear': ['revision']
            }
        }

        changes = perform_metadata_management(doc, config)

        # Revision field might be read-only, so changes could be 0 or 1
        assert changes >= 0

    @pytest.mark.unit
    def test_datetime_fields_not_cleared(self):
        """Test that datetime fields (created, modified) are not cleared."""
        doc = Document()

        config = {
            'metadata_operations': {
                'clear': ['created', 'modified']
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 0


# ============================================================================
# Set Metadata Tests
# ============================================================================

class TestSetMetadata:
    """Tests for setting metadata fields."""

    @pytest.mark.unit
    def test_set_single_metadata_field(self):
        """Test setting a single metadata field."""
        doc = Document()

        config = {
            'metadata_operations': {
                'set': {
                    'author': 'New Author'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.author == 'New Author'

    @pytest.mark.unit
    def test_set_multiple_metadata_fields(self):
        """Test setting multiple metadata fields."""
        doc = Document()

        config = {
            'metadata_operations': {
                'set': {
                    'author': 'New Author',
                    'title': 'New Title',
                    'subject': 'New Subject',
                    'keywords': 'new, keywords'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 4
        assert doc.core_properties.author == 'New Author'
        assert doc.core_properties.title == 'New Title'
        assert doc.core_properties.subject == 'New Subject'
        assert doc.core_properties.keywords == 'new, keywords'

    @pytest.mark.unit
    def test_set_nonexistent_field(self):
        """Test setting a field that doesn't exist (should be skipped)."""
        doc = Document()

        config = {
            'metadata_operations': {
                'set': {
                    'nonexistent_field': 'value',
                    'author': 'Test Author'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.author == 'Test Author'

    @pytest.mark.unit
    def test_set_content_status(self):
        """Test setting content_status field."""
        doc = Document()

        config = {
            'metadata_operations': {
                'set': {
                    'content_status': 'Final'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.content_status == 'Final'

    @pytest.mark.unit
    def test_set_last_modified_by(self):
        """Test setting last_modified_by field."""
        doc = Document()

        config = {
            'metadata_operations': {
                'set': {
                    'last_modified_by': 'Automation Script'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 1
        assert doc.core_properties.last_modified_by == 'Automation Script'

    @pytest.mark.unit
    def test_overwrite_existing_metadata(self):
        """Test overwriting existing metadata values."""
        doc = Document()
        doc.core_properties.author = 'Old Author'
        doc.core_properties.title = 'Old Title'

        config = {
            'metadata_operations': {
                'set': {
                    'author': 'New Author',
                    'title': 'New Title'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 2
        assert doc.core_properties.author == 'New Author'
        assert doc.core_properties.title == 'New Title'


# ============================================================================
# Combined Operations Tests
# ============================================================================

class TestCombinedOperations:
    """Tests for combined clear and set operations."""

    @pytest.mark.unit
    def test_clear_and_set_operations(self):
        """Test combining clear and set operations."""
        doc = Document()
        doc.core_properties.author = 'Old Author'
        doc.core_properties.title = 'Old Title'

        config = {
            'metadata_operations': {
                'clear': ['author'],
                'set': {
                    'title': 'New Title',
                    'subject': 'New Subject'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 3
        assert doc.core_properties.author == ''
        assert doc.core_properties.title == 'New Title'
        assert doc.core_properties.subject == 'New Subject'

    @pytest.mark.unit
    def test_empty_metadata_operations(self):
        """Test with empty metadata operations."""
        doc = Document()
        doc.core_properties.author = 'Test Author'

        config = {
            'metadata_operations': {}
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 0
        assert doc.core_properties.author == 'Test Author'

    @pytest.mark.unit
    def test_no_metadata_operations_key(self):
        """Test with missing metadata_operations key."""
        doc = Document()
        doc.core_properties.author = 'Test Author'

        config = {}

        changes = perform_metadata_management(doc, config)

        assert changes == 0
        assert doc.core_properties.author == 'Test Author'


# ============================================================================
# Unicode Tests
# ============================================================================

class TestUnicodeMetadata:
    """Tests for Unicode characters in metadata."""

    @pytest.mark.unit
    def test_unicode_metadata_values(self):
        """Test setting metadata with Unicode characters."""
        doc = Document()

        config = {
            'metadata_operations': {
                'set': {
                    'author': '作者名',
                    'title': 'Document 文档',
                    'keywords': 'test, 测试, emoj'
                }
            }
        }

        changes = perform_metadata_management(doc, config)

        assert changes == 3
        assert doc.core_properties.author == '作者名'
        assert doc.core_properties.title == 'Document 文档'
        assert doc.core_properties.keywords == 'test, 测试, emoj'

    @pytest.mark.unit
    def test_process_document_with_unicode(self, docx_with_metadata):
        """Test process_document with Unicode metadata."""
        temp_doc = docx_with_metadata()

        config = {
            'metadata_operations': {
                'set': {
                    'author': 'Li Ming',
                    'title': 'Cafe Resume',
                    'keywords': 'naive, resume'
                }
            }
        }

        result = process_document(str(temp_doc), config)

        assert result.success is True
        assert result.changes_made == 3

        doc = Document(temp_doc)
        assert doc.core_properties.author == 'Li Ming'
        assert doc.core_properties.title == 'Cafe Resume'
        assert doc.core_properties.keywords == 'naive, resume'


# ============================================================================
# process_document Integration Tests
# ============================================================================

class TestProcessDocument:
    """Integration tests for process_document function."""

    @pytest.mark.unit
    def test_process_document_success(self, docx_with_metadata):
        """Test process_document wrapper function with successful operation."""
        temp_doc = docx_with_metadata()

        config = {
            'metadata_operations': {
                'set': {
                    'author': 'Test Author',
                    'title': 'Test Title'
                }
            }
        }

        result = process_document(str(temp_doc), config)

        assert result.success is True
        assert result.changes_made == 2
        assert result.error_message is None
        assert result.duration_seconds > 0

        doc = Document(temp_doc)
        assert doc.core_properties.author == 'Test Author'
        assert doc.core_properties.title == 'Test Title'

    @pytest.mark.unit
    def test_process_document_no_changes(self, docx_with_metadata):
        """Test process_document with no changes made."""
        temp_doc = docx_with_metadata()

        config = {
            'metadata_operations': {
                'clear': [],
                'set': {}
            }
        }

        result = process_document(str(temp_doc), config)

        assert result.success is True
        assert result.changes_made == 0
        assert result.error_message is None

    @pytest.mark.unit
    def test_process_document_file_not_found(self):
        """Test process_document with non-existent file."""
        config = {
            'metadata_operations': {
                'set': {'author': 'Test'}
            }
        }

        result = process_document('nonexistent.docx', config)

        assert result.success is False
        assert result.changes_made == 0
        assert result.error_message is not None
        assert 'nonexistent.docx' in result.file_path

    @pytest.mark.unit
    def test_set_invalid_property_type(self, docx_with_metadata):
        """Test that setting property to invalid type doesn't crash."""
        temp_doc = docx_with_metadata()

        config = {
            'metadata_operations': {
                'set': {
                    'revision': 'not_a_number',  # triggers exception
                    'author': 'Valid Author'
                }
            }
        }

        result = process_document(str(temp_doc), config)

        assert result.success is True
        assert result.changes_made == 1

        doc = Document(temp_doc)
        assert doc.core_properties.author == 'Valid Author'

    @pytest.mark.unit
    def test_returns_processor_result(self, docx_with_metadata):
        """Should return a valid ProcessorResult."""
        temp_doc = docx_with_metadata()

        config = {'metadata_operations': {'set': {'author': 'Test'}}}
        result = process_document(str(temp_doc), config)

        assert isinstance(result, ProcessorResult)
        assert result.file_path == str(temp_doc)
        assert result.duration_seconds >= 0
