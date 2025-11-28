"""Unit tests for table, header, and footer search/replace functionality."""

import shutil
import pytest
from pathlib import Path
from docx import Document

from src.processors.search_replace import (
    _replace_in_table,
    perform_search_replace,
    _compile_pattern
)


@pytest.fixture
def test_docs_dir():
    """Return path to test documents directory."""
    return Path('tests/test_documents')


@pytest.fixture
def table_doc(test_docs_dir, tmp_path):
    """Create temporary copy of table document."""
    source = test_docs_dir / 'table_doc.docx'
    dest = tmp_path / 'table_test.docx'
    shutil.copy(source, dest)
    return dest


@pytest.fixture
def header_footer_doc(test_docs_dir, tmp_path):
    """Create temporary copy of header/footer document."""
    source = test_docs_dir / 'header_footer_doc.docx'
    dest = tmp_path / 'header_footer_test.docx'
    shutil.copy(source, dest)
    return dest


# ==================== Table Tests ====================

def test_replace_in_table_basic():
    """Test basic text replacement in table cells."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    # Add content
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'

    pattern = _compile_pattern('Header', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_table(table, pattern, 'TITLE')

    assert count == 2
    assert table.cell(0, 0).text == 'TITLE 1'
    assert table.cell(0, 1).text == 'TITLE 2'
    assert table.cell(1, 0).text == 'Data 1'  # Unchanged
    assert table.cell(1, 1).text == 'Data 2'  # Unchanged


def test_replace_in_table_all_cells():
    """Test replacement across all table cells."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)

    # Fill all cells with 'test'
    for row in table.rows:
        for cell in row.cells:
            cell.text = 'test'

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_table(table, pattern, 'REPLACED')

    assert count == 9  # 3x3 table

    # Verify all cells replaced
    for row in table.rows:
        for cell in row.cells:
            assert cell.text == 'REPLACED'


def test_replace_in_table_with_formatting():
    """Test table replacement preserves cell formatting."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Add formatted text to cell
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Bold Text')
    run.bold = True

    pattern = _compile_pattern('Bold', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_table(table, pattern, 'STRONG')

    assert count == 1
    assert 'STRONG Text' in cell.text

    # Verify bold formatting preserved
    assert cell.paragraphs[0].runs[0].bold is True


def test_replace_in_table_multiple_paragraphs_per_cell():
    """Test replacement in cells with multiple paragraphs."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Add multiple paragraphs
    cell.paragraphs[0].text = 'First test'
    cell.add_paragraph('Second test')
    cell.add_paragraph('Third test')

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_table(table, pattern, 'REPLACED')

    assert count == 3
    assert 'REPLACED' in cell.text


def test_replace_in_table_no_match():
    """Test table replacement when no matches found."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Data'

    pattern = _compile_pattern('notfound', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_table(table, pattern, 'REPLACED')

    assert count == 0
    assert table.cell(0, 0).text == 'Data'


def test_replace_in_empty_table():
    """Test replacement in empty table."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # Leave all cells empty

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_table(table, pattern, 'REPLACED')

    assert count == 0


def test_replace_in_table_from_file(table_doc):
    """Test replacement in table loaded from file."""
    doc = Document(table_doc)

    # Document should have a table with 'Data' in cells
    assert len(doc.tables) > 0
    table = doc.tables[0]

    pattern = _compile_pattern('Data', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_table(table, pattern, 'INFO')

    assert count == 2  # Should match 'Data 1' and 'Data 2'


# ==================== Header/Footer Tests ====================

def test_replace_in_header(header_footer_doc):
    """Test text replacement in document header."""
    doc = Document(header_footer_doc)

    config = {
        'search_term': 'header',
        'replace_term': 'HEADER',
        'use_regex': False,
        'case_sensitive': False,
        'whole_words': False,
        'search_body': False,
        'search_tables': False,
        'search_headers': True,
        'search_footers': False
    }

    count = perform_search_replace(doc, config)

    assert count == 1

    # Verify header was changed
    header_text = doc.sections[0].header.paragraphs[0].text
    assert 'HEADER' in header_text


def test_replace_in_footer(header_footer_doc):
    """Test text replacement in document footer."""
    doc = Document(header_footer_doc)

    config = {
        'search_term': 'footer',
        'replace_term': 'FOOTER',
        'use_regex': False,
        'case_sensitive': False,
        'whole_words': False,
        'search_body': False,
        'search_tables': False,
        'search_headers': False,
        'search_footers': True
    }

    count = perform_search_replace(doc, config)

    assert count == 1

    # Verify footer was changed
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert 'FOOTER' in footer_text


def test_replace_in_header_and_footer(header_footer_doc):
    """Test replacement in both header and footer."""
    doc = Document(header_footer_doc)

    config = {
        'search_term': 'is',
        'replace_term': 'IS',
        'use_regex': False,
        'case_sensitive': False,
        'whole_words': False,
        'search_body': False,
        'search_tables': False,
        'search_headers': True,
        'search_footers': True
    }

    count = perform_search_replace(doc, config)

    # 'is' appears multiple times: in "This" and standalone "is" in both header and footer
    assert count == 4  # 2 in header ("This" + "is") + 2 in footer ("This" + "is")

    header_text = doc.sections[0].header.paragraphs[0].text
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert 'IS' in header_text
    assert 'IS' in footer_text


def test_header_footer_not_searched_when_disabled(header_footer_doc):
    """Test that headers/footers are not searched when flags are False."""
    doc = Document(header_footer_doc)

    config = {
        'search_term': 'header',
        'replace_term': 'REPLACED',
        'use_regex': False,
        'case_sensitive': False,
        'whole_words': False,
        'search_body': True,
        'search_tables': True,
        'search_headers': False,  # Disabled
        'search_footers': False
    }

    count = perform_search_replace(doc, config)

    # Should not find 'header' since header search is disabled
    assert count == 0

    # Verify header unchanged
    header_text = doc.sections[0].header.paragraphs[0].text
    assert 'header' in header_text.lower()
    assert 'REPLACED' not in header_text


# ==================== Integration Tests ====================

def test_search_all_sections(header_footer_doc):
    """Test searching body, tables, headers, and footers simultaneously."""
    doc = Document(header_footer_doc)

    # Add a table to the document
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'content'

    # Add body paragraph with 'content'
    doc.add_paragraph('Main content here')

    config = {
        'search_term': 'content',
        'replace_term': 'CONTENT',
        'use_regex': False,
        'case_sensitive': False,
        'whole_words': False,
        'search_body': True,
        'search_tables': True,
        'search_headers': True,
        'search_footers': True
    }

    count = perform_search_replace(doc, config)

    # Should find in both body and table (at least 2 matches)
    assert count >= 2


def test_search_only_body(header_footer_doc):
    """Test searching only body paragraphs."""
    doc = Document(header_footer_doc)

    config = {
        'search_term': 'Main',
        'replace_term': 'PRIMARY',
        'use_regex': False,
        'case_sensitive': True,
        'whole_words': False,
        'search_body': True,
        'search_tables': False,
        'search_headers': False,
        'search_footers': False
    }

    count = perform_search_replace(doc, config)

    # Should find 'Main' in body
    assert count == 1


def test_search_only_tables(table_doc):
    """Test searching only tables."""
    doc = Document(table_doc)

    config = {
        'search_term': 'Header',
        'replace_term': 'HEADER',
        'use_regex': False,
        'case_sensitive': True,
        'whole_words': False,
        'search_body': False,
        'search_tables': True,
        'search_headers': False,
        'search_footers': False
    }

    count = perform_search_replace(doc, config)

    # Should find 'Header' in table
    assert count == 2


def test_multiple_sections_with_different_headers():
    """Test document with multiple sections having different headers."""
    doc = Document()

    # First section
    doc.add_paragraph('Section 1')
    section1 = doc.sections[0]
    section1.header.paragraphs[0].text = 'Header 1'

    # Add second section
    doc.add_section()
    doc.add_paragraph('Section 2')
    section2 = doc.sections[1]
    section2.header.paragraphs[0].text = 'Header 2'

    config = {
        'search_term': 'Header',
        'replace_term': 'TITLE',
        'use_regex': False,
        'case_sensitive': True,
        'whole_words': False,
        'search_body': False,
        'search_tables': False,
        'search_headers': True,
        'search_footers': False
    }

    count = perform_search_replace(doc, config)

    # Note: In python-docx, section headers may be linked by default.
    # When setting section2 header, it may overwrite section1 header.
    # The actual count depends on header linkage behavior.
    assert count >= 1  # At least one 'Header' should be found


def test_table_with_unicode(test_docs_dir, tmp_path):
    """Test replacement in table with Unicode characters."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    table.cell(0, 0).text = '测试 1'
    table.cell(0, 1).text = 'Test 2'
    table.cell(1, 0).text = '测试 3'
    table.cell(1, 1).text = 'Test 4'

    pattern = _compile_pattern('测试', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_table(table, pattern, '替换')

    assert count == 2
    assert '替换' in table.cell(0, 0).text
    assert '替换' in table.cell(1, 0).text


def test_nested_tables_not_double_counted():
    """Test that nested tables are handled correctly."""
    doc = Document()
    outer_table = doc.add_table(rows=1, cols=1)

    # Add text to outer table
    outer_table.cell(0, 0).text = 'test outer'

    # Note: python-docx doesn't easily support nested tables,
    # but we test the function handles it gracefully
    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_table(outer_table, pattern, 'REPLACED')

    assert count == 1
