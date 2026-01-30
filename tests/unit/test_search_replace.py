"""Unit tests for Search & Replace processor.

Tests cover:
- Pattern compilation (literal, regex, case sensitivity, whole words)
- Paragraph replacement with formatting preservation
- Table cell replacement
- Header and footer replacement
- Edge cases and error handling

Consolidated from: test_search_replace_patterns.py, test_paragraph_replacement.py,
                  test_table_header_footer.py
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

from src.processors.search_replace import (
    process_document,
    perform_search_replace,
    _compile_pattern,
    _replace_in_paragraph,
    _replace_in_table,
)
from src.processors import ProcessorResult


# ============================================================================
# Pattern Compilation Tests
# ============================================================================

class TestPatternCompilation:
    """Tests for _compile_pattern function."""

    @pytest.mark.unit
    def test_literal_case_sensitive(self):
        """Literal search with case sensitivity."""
        pattern = _compile_pattern('Test', use_regex=False, case_sensitive=True, whole_words=False)

        assert pattern.search('Test') is not None
        assert pattern.search('test') is None
        assert pattern.search('Testing') is not None

    @pytest.mark.unit
    def test_literal_case_insensitive(self):
        """Literal search without case sensitivity."""
        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)

        assert pattern.search('test') is not None
        assert pattern.search('Test') is not None
        assert pattern.search('TEST') is not None

    @pytest.mark.unit
    def test_literal_whole_words_only(self):
        """Literal search with whole words only."""
        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=True)

        assert pattern.search('test') is not None
        assert pattern.search('a test here') is not None
        assert pattern.search('testing') is None
        assert pattern.search('pretest') is None

    @pytest.mark.unit
    def test_literal_escapes_special_chars(self):
        """Literal search escapes regex special characters."""
        special_chars = r'.*+?[]{}()^$|\\'
        pattern = _compile_pattern(special_chars, use_regex=False, case_sensitive=True, whole_words=False)

        assert pattern.search(special_chars) is not None
        assert pattern.search('anything') is None

    @pytest.mark.unit
    def test_literal_parentheses(self):
        """Literal search with parentheses (common in legal text)."""
        pattern = _compile_pattern('(a)', use_regex=False, case_sensitive=True, whole_words=False)

        assert pattern.search('(a)') is not None
        assert pattern.search('a') is None

    @pytest.mark.unit
    def test_literal_brackets(self):
        """Literal search with brackets."""
        pattern = _compile_pattern('[test]', use_regex=False, case_sensitive=True, whole_words=False)

        assert pattern.search('[test]') is not None
        assert pattern.search('test') is None

    @pytest.mark.unit
    def test_regex_basic_pattern(self):
        """Regex search with basic pattern."""
        pattern = _compile_pattern(r'\d+', use_regex=True, case_sensitive=True, whole_words=False)

        assert pattern.search('123') is not None
        assert pattern.search('abc') is None

    @pytest.mark.unit
    def test_regex_case_insensitive(self):
        """Regex search with case insensitivity."""
        pattern = _compile_pattern(r'[a-z]+', use_regex=True, case_sensitive=False, whole_words=False)

        assert pattern.search('abc') is not None
        assert pattern.search('ABC') is not None

    @pytest.mark.unit
    def test_regex_case_sensitive(self):
        """Regex search with case sensitivity."""
        pattern = _compile_pattern(r'[a-z]+', use_regex=True, case_sensitive=True, whole_words=False)

        assert pattern.search('abc') is not None
        assert pattern.search('ABC') is None

    @pytest.mark.unit
    def test_regex_capture_groups(self):
        """Regex search with capture groups."""
        pattern = _compile_pattern(r'(\w+)@(\w+)\.com', use_regex=True, case_sensitive=False, whole_words=False)

        match = pattern.search('user@example.com')
        assert match is not None
        assert match.group(1) == 'user'
        assert match.group(2) == 'example'

    @pytest.mark.unit
    def test_regex_word_boundaries(self):
        """Regex search with user-defined word boundaries."""
        pattern = _compile_pattern(r'\btest\b', use_regex=True, case_sensitive=False, whole_words=False)

        assert pattern.search('test') is not None
        assert pattern.search('a test here') is not None
        assert pattern.search('testing') is None
        assert pattern.search('pretest') is None

    @pytest.mark.unit
    def test_regex_alternation(self):
        """Regex search with alternation (OR)."""
        pattern = _compile_pattern(r'cat|dog', use_regex=True, case_sensitive=False, whole_words=False)

        assert pattern.search('I have a cat') is not None
        assert pattern.search('I have a dog') is not None
        assert pattern.search('I have a bird') is None

    @pytest.mark.unit
    def test_whole_words_ignored_in_regex_mode(self):
        """whole_words parameter is ignored in regex mode."""
        pattern = _compile_pattern(r'test', use_regex=True, case_sensitive=False, whole_words=True)

        # Should match as substring since regex mode ignores whole_words
        assert pattern.search('testing') is not None

    @pytest.mark.unit
    def test_unicode_search(self):
        """Pattern compilation with Unicode characters."""
        pattern = _compile_pattern('测试', use_regex=False, case_sensitive=True, whole_words=False)

        assert pattern.search('这是测试文本') is not None
        assert pattern.search('This is a test') is None

    @pytest.mark.unit
    def test_emoji_search(self):
        """Pattern compilation with emoji characters."""
        pattern = _compile_pattern('😀', use_regex=False, case_sensitive=True, whole_words=False)

        assert pattern.search('Hello 😀 World') is not None
        assert pattern.search('Hello World') is None

    @pytest.mark.unit
    def test_subn_replacement(self):
        """Compiled pattern can perform substitution."""
        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)

        text = 'This is a test. Testing 123.'
        new_text, count = pattern.subn('REPLACED', text)

        assert count == 2
        assert 'REPLACED' in new_text

    @pytest.mark.unit
    def test_subn_whole_words(self):
        """Substitution with whole words only."""
        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=True)

        text = 'This is a test. Testing 123.'
        new_text, count = pattern.subn('REPLACED', text)

        assert count == 1
        assert new_text == 'This is a REPLACED. Testing 123.'

    @pytest.mark.unit
    def test_multiline_regex(self):
        """Regex pattern with multiline text."""
        pattern = _compile_pattern(r'^test', use_regex=True, case_sensitive=False, whole_words=False)

        # Without MULTILINE flag, ^ only matches start of string
        assert pattern.search('test\nmore') is not None
        assert pattern.search('start\ntest') is None


# ============================================================================
# Paragraph Replacement Tests (Formatting Preservation)
# ============================================================================

class TestParagraphReplacement:
    """Tests for _replace_in_paragraph function with formatting preservation."""

    @pytest.mark.unit
    def test_replace_in_empty_paragraph(self, docx_factory):
        """Replacement in empty paragraph doesn't crash."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))
        # Add empty paragraph
        paragraph = doc.add_paragraph('')

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_paragraph(paragraph, pattern, 'REPLACED')

        assert count == 0

    @pytest.mark.unit
    def test_preserves_bold_formatting(self, tmp_path):
        """Bold formatting preserved after replacement."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Hello world')
        run.bold = True

        doc_path = tmp_path / "bold.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('Hello', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'Goodbye')

        assert count == 1
        assert 'Goodbye' in doc.paragraphs[0].text
        assert doc.paragraphs[0].runs[0].bold is True

    @pytest.mark.unit
    def test_preserves_italic_formatting(self, tmp_path):
        """Italic formatting preserved after replacement."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Important text')
        run.italic = True

        doc_path = tmp_path / "italic.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('Important', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'Critical')

        assert count == 1
        assert doc.paragraphs[0].runs[0].italic is True

    @pytest.mark.unit
    def test_preserves_underline_formatting(self, tmp_path):
        """Underline formatting preserved after replacement."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Underlined text')
        run.underline = True

        doc_path = tmp_path / "underline.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('text', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'content')

        assert count == 1
        assert any(run.underline for run in doc.paragraphs[0].runs)

    @pytest.mark.unit
    def test_multiple_occurrences_same_paragraph(self, tmp_path):
        """Multiple occurrences in same paragraph all replaced."""
        doc = Document()
        doc.add_paragraph('test test test')

        doc_path = tmp_path / "multi.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'REPLACED')

        assert count == 3
        assert doc.paragraphs[0].text == 'REPLACED REPLACED REPLACED'

    @pytest.mark.unit
    def test_case_insensitive_preserves_formatting(self, tmp_path):
        """Case-insensitive replacement preserves formatting on both matches."""
        doc = Document()
        p = doc.add_paragraph()
        run1 = p.add_run('Test')
        run1.bold = True
        p.add_run(' and ')
        run2 = p.add_run('test')
        run2.italic = True

        doc_path = tmp_path / "case.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'REPLACED')

        assert count == 2
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].runs[2].italic is True

    @pytest.mark.unit
    def test_regex_replacement_preserves_formatting(self, tmp_path):
        """Regex replacement preserves formatting."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Call 555-1234 or 555-5678')
        run.bold = True

        doc_path = tmp_path / "regex.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern(r'\d{3}-\d{4}', use_regex=True, case_sensitive=False, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'XXX-XXXX')

        assert count == 2
        assert 'XXX-XXXX' in doc.paragraphs[0].text
        assert doc.paragraphs[0].runs[0].bold is True

    @pytest.mark.unit
    def test_unicode_replacement(self, tmp_path):
        """Unicode text replacement preserves characters."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('测试 test 测试')
        run.bold = True

        doc_path = tmp_path / "unicode.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('测试', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, '替换')

        assert count == 2
        assert '替换 test 替换' == doc.paragraphs[0].text

    @pytest.mark.unit
    def test_whole_word_replacement_with_formatting(self, tmp_path):
        """Whole word replacement preserves formatting."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('test testing tested')
        run.underline = True

        doc_path = tmp_path / "whole.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=True)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'REPLACED')

        assert count == 1  # Only 'test', not 'testing' or 'tested'
        assert 'REPLACED' in doc.paragraphs[0].text
        assert doc.paragraphs[0].runs[0].underline is True

    @pytest.mark.unit
    def test_no_replacement_returns_zero(self, tmp_path):
        """No match returns zero count."""
        doc = Document()
        doc.add_paragraph('This is a test')

        doc_path = tmp_path / "nomatch.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('notfound', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'REPLACED')

        assert count == 0
        assert doc.paragraphs[0].text == 'This is a test'

    @pytest.mark.unit
    def test_replace_emoji(self, tmp_path):
        """Replacing text with emoji preserves formatting."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Hello World')
        run.italic = True

        doc_path = tmp_path / "emoji.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('World', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, '😀')

        assert count == 1
        assert 'Hello 😀' == doc.paragraphs[0].text
        assert doc.paragraphs[0].runs[0].italic is True

    @pytest.mark.unit
    def test_replace_partial_run_text(self, tmp_path):
        """Replacing part of a run's text."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('abcdefgh')
        run.bold = True

        doc_path = tmp_path / "partial.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('cde', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'XYZ')

        assert count == 1
        assert doc.paragraphs[0].text == 'abXYZfgh'
        assert doc.paragraphs[0].runs[0].bold is True

    @pytest.mark.unit
    def test_replace_across_multiple_runs(self, tmp_path):
        """Replacement when text spans multiple runs."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run('Hello ')
        p.add_run('World')

        doc_path = tmp_path / "multirun.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        # The pattern looks at paragraph.text which joins all runs
        pattern = _compile_pattern('Hello World', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, 'Goodbye Earth')

        # Should find and replace the text
        assert count == 1

    @pytest.mark.unit
    def test_empty_run_handling(self, docx_factory):
        """Replacing in paragraphs with empty runs."""
        doc_path = docx_factory(content="Test paragraph")
        doc = Document(str(doc_path))
        para = doc.paragraphs[0]

        # Add a run with empty text
        para.add_run("")

        pattern = _compile_pattern('Test', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(para, pattern, 'Sample')

        assert count == 1
        assert 'Sample' in para.text


# ============================================================================
# Table Replacement Tests
# ============================================================================

class TestTableReplacement:
    """Tests for _replace_in_table function."""

    @pytest.mark.unit
    def test_replace_in_table_basic(self, docx_with_table):
        """Basic text replacement in table cells."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))
        table = doc.tables[0]

        pattern = _compile_pattern('Header', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_table(table, pattern, 'TITLE')

        assert count == 2
        assert table.cell(0, 0).text == 'TITLE 1'
        assert table.cell(0, 1).text == 'TITLE 2'

    @pytest.mark.unit
    def test_replace_in_all_cells(self, tmp_path):
        """Replacement across all table cells."""
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        for row in table.rows:
            for cell in row.cells:
                cell.text = 'test'

        doc_path = tmp_path / "allcells.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, 'REPLACED')

        assert count == 9

    @pytest.mark.unit
    def test_replace_preserves_cell_formatting(self, tmp_path):
        """Table replacement preserves cell formatting."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        run = cell.paragraphs[0].add_run('Bold Text')
        run.bold = True

        doc_path = tmp_path / "formatted.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('Bold', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, 'STRONG')

        assert count == 1
        assert 'STRONG' in doc.tables[0].cell(0, 0).text
        assert doc.tables[0].cell(0, 0).paragraphs[0].runs[0].bold is True

    @pytest.mark.unit
    def test_multiple_paragraphs_per_cell(self, tmp_path):
        """Replacement in cells with multiple paragraphs."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = 'First test'
        cell.add_paragraph('Second test')
        cell.add_paragraph('Third test')

        doc_path = tmp_path / "multipara.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, 'REPLACED')

        assert count == 3

    @pytest.mark.unit
    def test_no_match_in_table(self, docx_with_table):
        """Table replacement when no matches found."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))

        pattern = _compile_pattern('notfound', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, 'REPLACED')

        assert count == 0

    @pytest.mark.unit
    def test_empty_table(self, tmp_path):
        """Replacement in empty table."""
        doc = Document()
        doc.add_table(rows=2, cols=2)

        doc_path = tmp_path / "empty.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, 'REPLACED')

        assert count == 0

    @pytest.mark.unit
    def test_unicode_in_table(self, tmp_path):
        """Unicode replacement in table cells."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '测试 1'
        table.cell(1, 0).text = '测试 2'

        doc_path = tmp_path / "unicode_table.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('测试', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, '替换')

        assert count == 2
        assert '替换' in doc.tables[0].cell(0, 0).text

    @pytest.mark.unit
    def test_nested_table_handling(self, tmp_path):
        """Test that table replacement handles gracefully."""
        doc = Document()
        outer_table = doc.add_table(rows=1, cols=1)
        outer_table.cell(0, 0).text = 'test outer'

        doc_path = tmp_path / "nested.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
        count = _replace_in_table(doc.tables[0], pattern, 'REPLACED')

        assert count == 1


# ============================================================================
# Header/Footer Replacement Tests
# ============================================================================

class TestHeaderFooterReplacement:
    """Tests for header and footer search/replace."""

    @pytest.mark.unit
    def test_replace_in_header(self, tmp_path):
        """Text replacement in document header."""
        doc = Document()
        doc.add_paragraph('Body content')
        section = doc.sections[0]
        section.header.paragraphs[0].text = 'This is the header text'

        doc_path = tmp_path / "header.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        config = {
            'search_term': 'header',
            'replace_term': 'HEADER',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': False,
            'search_tables': False,
            'search_headers': True,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)

        assert count == 1
        assert 'HEADER' in doc.sections[0].header.paragraphs[0].text

    @pytest.mark.unit
    def test_replace_in_footer(self, tmp_path):
        """Text replacement in document footer."""
        doc = Document()
        doc.add_paragraph('Body content')
        section = doc.sections[0]
        section.footer.paragraphs[0].text = 'This is the footer text'

        doc_path = tmp_path / "footer.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        config = {
            'search_term': 'footer',
            'replace_term': 'FOOTER',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': False,
            'search_tables': False,
            'search_headers': False,
            'search_footers': True
        }

        count = perform_search_replace(doc, config)

        assert count == 1
        assert 'FOOTER' in doc.sections[0].footer.paragraphs[0].text

    @pytest.mark.unit
    def test_replace_in_both_header_and_footer(self, tmp_path):
        """Replacement in both header and footer."""
        doc = Document()
        doc.add_paragraph('Body content')
        section = doc.sections[0]
        section.header.paragraphs[0].text = 'Header with test'
        section.footer.paragraphs[0].text = 'Footer with test'

        doc_path = tmp_path / "both.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        config = {
            'search_term': 'test',
            'replace_term': 'TEST',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': False,
            'search_tables': False,
            'search_headers': True,
            'search_footers': True
        }

        count = perform_search_replace(doc, config)

        assert count == 2
        assert 'TEST' in doc.sections[0].header.paragraphs[0].text
        assert 'TEST' in doc.sections[0].footer.paragraphs[0].text

    @pytest.mark.unit
    def test_header_footer_disabled(self, tmp_path):
        """Headers/footers not searched when flags are False."""
        doc = Document()
        doc.add_paragraph('Body content')
        section = doc.sections[0]
        section.header.paragraphs[0].text = 'Header with test'
        section.footer.paragraphs[0].text = 'Footer with test'

        doc_path = tmp_path / "disabled.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        config = {
            'search_term': 'test',
            'replace_term': 'REPLACED',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': True,
            'search_tables': True,
            'search_headers': False,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)

        assert count == 0  # No matches in body/tables
        assert 'test' in doc.sections[0].header.paragraphs[0].text  # Unchanged

    @pytest.mark.unit
    def test_multiple_sections_with_headers(self, tmp_path):
        """Document with multiple sections having headers."""
        doc = Document()
        doc.add_paragraph('Section 1')
        section1 = doc.sections[0]
        section1.header.paragraphs[0].text = 'Header 1'

        # Add second section
        doc.add_section()
        doc.add_paragraph('Section 2')
        section2 = doc.sections[1]
        section2.header.paragraphs[0].text = 'Header 2'

        doc_path = tmp_path / "multi_section.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        config = {
            'search_term': 'Header',
            'replace_term': 'TITLE',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': False,
            'search_tables': False,
            'search_headers': True,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)

        # At least one 'Header' should be found (section headers may be linked)
        assert count >= 1


# ============================================================================
# perform_search_replace Tests
# ============================================================================

class TestPerformSearchReplace:
    """Tests for perform_search_replace function."""

    @pytest.mark.unit
    def test_search_all_sections(self, tmp_path):
        """Search body, tables, headers, footers simultaneously."""
        doc = Document()
        doc.add_paragraph('Body test')
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = 'Table test'
        doc.sections[0].header.paragraphs[0].text = 'Header test'
        doc.sections[0].footer.paragraphs[0].text = 'Footer test'

        doc_path = tmp_path / "all.docx"
        doc.save(str(doc_path))
        doc = Document(str(doc_path))

        config = {
            'search_term': 'test',
            'replace_term': 'REPLACED',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': True,
            'search_tables': True,
            'search_headers': True,
            'search_footers': True
        }

        count = perform_search_replace(doc, config)

        assert count == 4

    @pytest.mark.unit
    def test_search_only_body(self, docx_factory):
        """Search only body paragraphs."""
        doc_path = docx_factory(content="Main test content")
        doc = Document(str(doc_path))

        config = {
            'search_term': 'test',
            'replace_term': 'REPLACED',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)

        assert count == 1

    @pytest.mark.unit
    def test_search_only_tables(self, docx_with_table):
        """Search only tables."""
        doc_path = docx_with_table(rows=2, cols=2)
        doc = Document(str(doc_path))

        config = {
            'search_term': 'Header',
            'replace_term': 'HEADER',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': False,
            'search_tables': True,
            'search_headers': False,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)

        assert count == 2


# ============================================================================
# process_document Integration Tests
# ============================================================================

class TestProcessDocument:
    """Integration tests for process_document function."""

    @pytest.mark.unit
    def test_returns_processor_result(self, docx_factory):
        """Should return a valid ProcessorResult."""
        doc_path = docx_factory(content="Hello world\nHello again")

        config = {
            'search_term': 'Hello',
            'replace_term': 'Goodbye',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        result = process_document(str(doc_path), config)

        assert isinstance(result, ProcessorResult)
        assert result.success is True
        assert result.changes_made == 2
        assert result.duration_seconds > 0

    @pytest.mark.unit
    def test_saves_document_with_changes(self, docx_factory):
        """Document saved when changes made."""
        doc_path = docx_factory(content="Hello world")

        config = {
            'search_term': 'Hello',
            'replace_term': 'Goodbye',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        result = process_document(str(doc_path), config)

        assert result.success is True

        # Reload and verify
        doc = Document(str(doc_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert 'Goodbye' in text
        assert 'Hello' not in text

    @pytest.mark.unit
    def test_handles_file_not_found(self):
        """Handles missing file gracefully."""
        config = {
            'search_term': 'test',
            'replace_term': 'sample',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False
        }

        result = process_document('/nonexistent/file.docx', config)

        assert result.success is False
        assert result.changes_made == 0
        assert result.error_message is not None

    @pytest.mark.unit
    def test_handles_corrupted_file(self, tmp_path):
        """Handles corrupted files gracefully."""
        corrupted = tmp_path / "corrupted.docx"
        corrupted.write_text("Not a valid DOCX", encoding='utf-8')

        config = {
            'search_term': 'test',
            'replace_term': 'sample',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False
        }

        result = process_document(str(corrupted), config)

        assert result.success is False
        assert result.error_message is not None

    @pytest.mark.unit
    def test_no_match_returns_zero_changes(self, docx_factory):
        """No matches returns zero changes."""
        doc_path = docx_factory(content="Hello world")

        config = {
            'search_term': 'notfound',
            'replace_term': 'replacement',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0


# ============================================================================
# Edge Cases
# ============================================================================

class TestEdgeCases:
    """Edge case and special scenario tests."""

    @pytest.mark.unit
    def test_empty_document(self, docx_factory):
        """Empty document handled gracefully."""
        doc_path = docx_factory(content="")

        config = {
            'search_term': 'test',
            'replace_term': 'replacement',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0

    @pytest.mark.unit
    def test_empty_search_term(self, docx_factory):
        """Empty search term behavior."""
        doc_path = docx_factory(content="test")

        pattern = _compile_pattern('', use_regex=False, case_sensitive=True, whole_words=False)
        # Empty pattern matches everywhere
        assert pattern.search('test') is not None

    @pytest.mark.unit
    def test_replacement_with_empty_string(self, docx_factory):
        """Replacing with empty string removes text."""
        doc_path = docx_factory(content="Hello world")
        doc = Document(str(doc_path))

        pattern = _compile_pattern('Hello ', use_regex=False, case_sensitive=True, whole_words=False)
        count = _replace_in_paragraph(doc.paragraphs[0], pattern, '')

        assert count == 1
        assert doc.paragraphs[0].text == 'world'

    @pytest.mark.unit
    @pytest.mark.parametrize("search,replace,expected_count", [
        ("test", "TEST", 3),
        ("TEST", "test", 0),  # Case sensitive
        ("test", "replacement", 3),
    ])
    def test_various_replacement_scenarios(self, docx_factory, search, replace, expected_count):
        """Various replacement scenarios."""
        doc_path = docx_factory(content="test test test")
        doc = Document(str(doc_path))

        config = {
            'search_term': search,
            'replace_term': replace,
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)
        assert count == expected_count

    @pytest.mark.unit
    def test_large_document_handling(self, large_docx):
        """Large document processing works correctly."""
        config = {
            'search_term': 'content',
            'replace_term': 'CONTENT',
            'use_regex': False,
            'case_sensitive': False,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        result = process_document(str(large_docx), config)

        assert result.success is True
        assert result.changes_made >= 1  # At least some matches

    @pytest.mark.unit
    def test_special_chars_in_replacement(self, docx_factory):
        """Special characters in replacement text."""
        doc_path = docx_factory(content="Hello world")
        doc = Document(str(doc_path))

        config = {
            'search_term': 'world',
            'replace_term': '$100 & <more>',
            'use_regex': False,
            'case_sensitive': True,
            'whole_words': False,
            'search_body': True,
            'search_tables': False,
            'search_headers': False,
            'search_footers': False
        }

        count = perform_search_replace(doc, config)

        assert count == 1
        assert '$100 & <more>' in doc.paragraphs[0].text
