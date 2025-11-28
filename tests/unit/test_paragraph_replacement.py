"""Unit tests for paragraph-level text replacement with formatting preservation."""

import re
import shutil
import pytest
from pathlib import Path
from docx import Document

from src.processors.search_replace import _replace_in_paragraph, _compile_pattern


@pytest.fixture
def test_docs_dir():
    """Return path to test documents directory."""
    return Path('tests/test_documents')


@pytest.fixture
def temp_doc(test_docs_dir, tmp_path):
    """Create a temporary copy of formatted_text.docx for testing."""
    source = test_docs_dir / 'formatted_text.docx'
    dest = tmp_path / 'test.docx'
    shutil.copy(source, dest)
    return dest


def test_replace_in_empty_paragraph():
    """Test replacement in empty paragraph doesn't crash."""
    doc = Document()
    paragraph = doc.add_paragraph('')

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'REPLACED')

    assert count == 0
    assert paragraph.text == ''


def test_replace_preserves_bold_formatting(temp_doc):
    """Test that bold formatting is preserved after replacement."""
    doc = Document(temp_doc)

    # First paragraph has "Hello" in bold
    paragraph = doc.paragraphs[0]
    assert paragraph.text == 'Hello world'

    # Get original formatting
    first_run = paragraph.runs[0]
    original_bold = first_run.bold

    # Replace 'Hello' with 'Goodbye'
    pattern = _compile_pattern('Hello', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'Goodbye')

    assert count == 1
    assert paragraph.text == 'Goodbye world'

    # Verify bold formatting is preserved on first run
    assert paragraph.runs[0].bold == original_bold
    assert paragraph.runs[0].bold is True


def test_replace_preserves_mixed_formatting(temp_doc):
    """Test that mixed formatting (bold + italic) is preserved."""
    doc = Document(temp_doc)

    # Second paragraph has "important" in bold+italic
    paragraph = doc.paragraphs[1]
    assert 'important' in paragraph.text

    # Find the run with 'important'
    important_run = None
    for run in paragraph.runs:
        if 'important' in run.text:
            important_run = run
            break

    assert important_run is not None
    assert important_run.bold is True
    assert important_run.italic is True

    # Replace 'important'
    pattern = _compile_pattern('important', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'CRITICAL')

    assert count == 1

    # Find the replaced run
    critical_run = None
    for run in paragraph.runs:
        if 'CRITICAL' in run.text:
            critical_run = run
            break

    assert critical_run is not None
    assert critical_run.bold is True
    assert critical_run.italic is True


def test_replace_preserves_font_size(temp_doc):
    """Test that font sizes are preserved after replacement."""
    doc = Document(temp_doc)

    # Third paragraph has different font sizes
    paragraph = doc.paragraphs[2]

    # Get original font sizes
    original_sizes = [run.font.size for run in paragraph.runs]

    # Replace 'medium' which should be in the middle run
    pattern = _compile_pattern('medium', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'MEDIUM')

    assert count == 1

    # Verify font sizes are still different (preserved)
    new_sizes = [run.font.size for run in paragraph.runs]

    # At least one size should be different from others (mixed formatting preserved)
    assert len(set(s for s in new_sizes if s is not None)) > 1


def test_replace_preserves_color(temp_doc):
    """Test that text colors are preserved after replacement."""
    doc = Document(temp_doc)

    # Fourth paragraph has colored text
    paragraph = doc.paragraphs[3]

    # Get original colors
    original_colors = []
    for run in paragraph.runs:
        if run.font.color.rgb:
            original_colors.append(run.font.color.rgb)

    assert len(original_colors) > 0  # Should have colored text

    # Replace 'green'
    pattern = _compile_pattern('green', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'GREEN')

    assert count == 1

    # Verify colors are still present
    new_colors = []
    for run in paragraph.runs:
        if run.font.color.rgb:
            new_colors.append(run.font.color.rgb)

    assert len(new_colors) > 0


def test_replace_preserves_underline(temp_doc):
    """Test that underline formatting is preserved."""
    doc = Document(temp_doc)

    # Fifth paragraph has underlined text
    paragraph = doc.paragraphs[4]
    assert paragraph.text == 'Underlined text'

    # Verify original underline
    assert any(run.underline for run in paragraph.runs)

    # Replace 'text'
    pattern = _compile_pattern('text', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'content')

    assert count == 1
    assert 'content' in paragraph.text

    # Verify underline is still present
    assert any(run.underline for run in paragraph.runs)


def test_replace_multiple_occurrences_in_paragraph():
    """Test replacing multiple occurrences in same paragraph."""
    doc = Document()
    paragraph = doc.add_paragraph('test test test')

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'REPLACED')

    assert count == 3
    assert paragraph.text == 'REPLACED REPLACED REPLACED'


def test_replace_with_case_insensitive():
    """Test case-insensitive replacement preserves formatting."""
    doc = Document()
    p = doc.add_paragraph()
    run1 = p.add_run('Test')
    run1.bold = True
    p.add_run(' and ')
    run2 = p.add_run('test')
    run2.italic = True

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_paragraph(p, pattern, 'REPLACED')

    assert count == 2
    assert 'REPLACED' in p.text

    # Check formatting preserved on both runs
    assert p.runs[0].bold is True
    assert p.runs[2].italic is True


def test_replace_with_regex_pattern():
    """Test regex replacement preserves formatting."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run('Call 555-1234 or 555-5678')
    run.bold = True

    pattern = _compile_pattern(r'\d{3}-\d{4}', use_regex=True, case_sensitive=False, whole_words=False)
    count = _replace_in_paragraph(p, pattern, 'XXX-XXXX')

    assert count == 2
    assert 'XXX-XXXX' in p.text
    assert p.runs[0].bold is True


def test_no_replacement_returns_zero():
    """Test that no match returns zero count."""
    doc = Document()
    paragraph = doc.add_paragraph('This is a test')

    pattern = _compile_pattern('notfound', use_regex=False, case_sensitive=False, whole_words=False)
    count = _replace_in_paragraph(paragraph, pattern, 'REPLACED')

    assert count == 0
    assert paragraph.text == 'This is a test'


def test_replace_unicode_text():
    """Test replacing Unicode text preserves characters."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run('测试 test 测试')
    run.bold = True

    pattern = _compile_pattern('测试', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(p, pattern, '替换')

    assert count == 2
    assert '替换 test 替换' == p.text
    assert p.runs[0].bold is True


def test_replace_emoji():
    """Test replacing text with emoji preserves formatting."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run('Hello World')
    run.italic = True

    pattern = _compile_pattern('World', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(p, pattern, '😀')

    assert count == 1
    assert 'Hello 😀' == p.text
    assert p.runs[0].italic is True


def test_replace_partial_run_text():
    """Test replacing part of a run's text."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run('abcdefgh')
    run.bold = True

    pattern = _compile_pattern('cde', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(p, pattern, 'XYZ')

    assert count == 1
    assert p.text == 'abXYZfgh'
    assert p.runs[0].bold is True


def test_replace_across_multiple_runs():
    """Test replacement when text spans multiple runs."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('Hello ')
    p.add_run('World')

    # The pattern looks at paragraph.text which joins all runs
    pattern = _compile_pattern('Hello World', use_regex=False, case_sensitive=True, whole_words=False)
    count = _replace_in_paragraph(p, pattern, 'Goodbye Earth')

    # Should find and replace the text
    assert count == 1


def test_whole_word_replacement_with_formatting():
    """Test whole word replacement preserves formatting."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run('test testing tested')
    run.underline = True

    pattern = _compile_pattern('test', use_regex=False, case_sensitive=False, whole_words=True)
    count = _replace_in_paragraph(p, pattern, 'REPLACED')

    # Should only replace 'test', not 'testing' or 'tested'
    assert count == 1
    assert 'REPLACED' in p.text
    assert p.runs[0].underline is True
