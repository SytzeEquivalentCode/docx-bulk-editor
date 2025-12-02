"""Unit tests for the Validator processor.

Tests cover:
- Heading hierarchy validation
- Empty paragraph detection
- Placeholder pattern detection
- Whitespace issue detection
- Report formatting
- Edge cases and error handling
"""

import pytest
from pathlib import Path
from docx import Document

from src.processors.validator import (
    process_document,
    perform_validation,
    _check_heading_hierarchy,
    _check_empty_paragraphs,
    _check_placeholders,
    _check_whitespace,
    _format_validation_report,
)
from src.processors import ProcessorResult


# ============================================================================
# Heading Hierarchy Tests
# ============================================================================

class TestHeadingHierarchy:
    """Tests for _check_heading_hierarchy function."""

    @pytest.mark.unit
    def test_valid_hierarchy_no_issues(self, docx_factory):
        """Valid heading hierarchy should produce no issues."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter 1', level=1)
        doc.add_paragraph('Content.')
        doc.add_heading('Section 1.1', level=2)
        doc.add_paragraph('Content.')
        doc.add_heading('Subsection 1.1.1', level=3)
        doc.add_paragraph('Content.')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_heading_hierarchy(doc)

        assert len(issues) == 0

    @pytest.mark.unit
    def test_detects_skipped_h2(self, docx_factory):
        """Should detect H1 -> H3 (skipped H2)."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter 1', level=1)
        doc.add_paragraph('Content.')
        doc.add_heading('Skipped to H3', level=3)  # Skipped H2
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_heading_hierarchy(doc)

        assert len(issues) == 1
        assert issues[0]['type'] == 'heading_hierarchy'
        assert 'Heading 3' in issues[0]['message']
        assert 'Heading 1' in issues[0]['message']

    @pytest.mark.unit
    def test_detects_skipped_h3(self, docx_factory):
        """Should detect H2 -> H4 (skipped H3)."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter 1', level=1)
        doc.add_heading('Section 1.1', level=2)
        doc.add_heading('Skipped to H4', level=4)  # Skipped H3
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_heading_hierarchy(doc)

        assert len(issues) == 1
        assert issues[0]['type'] == 'heading_hierarchy'
        assert 'Heading 4' in issues[0]['message']

    @pytest.mark.unit
    def test_multiple_hierarchy_violations(self, docx_factory):
        """Should detect multiple hierarchy violations."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter 1', level=1)
        doc.add_heading('Skipped to H3', level=3)  # First violation
        doc.add_heading('Chapter 2', level=1)
        doc.add_heading('Skipped to H4', level=4)  # Second violation
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_heading_hierarchy(doc)

        assert len(issues) == 2

    @pytest.mark.unit
    def test_going_back_levels_is_ok(self, docx_factory):
        """Going from deeper to shallower levels should not be flagged."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter 1', level=1)
        doc.add_heading('Section 1.1', level=2)
        doc.add_heading('Subsection 1.1.1', level=3)
        doc.add_heading('Chapter 2', level=1)  # Going back to H1
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_heading_hierarchy(doc)

        assert len(issues) == 0

    @pytest.mark.unit
    def test_preview_text_truncated(self, docx_factory):
        """Long heading text should be truncated in preview."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        long_title = 'A' * 100
        doc.add_heading('Chapter 1', level=1)
        doc.add_heading(long_title, level=3)  # Skipped H2
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_heading_hierarchy(doc)

        assert len(issues) == 1
        assert len(issues[0]['preview']) <= 53  # 50 chars + '...'


# ============================================================================
# Empty Paragraphs Tests
# ============================================================================

class TestEmptyParagraphs:
    """Tests for _check_empty_paragraphs function."""

    @pytest.mark.unit
    def test_single_empty_paragraph_ok(self, docx_factory):
        """Single empty paragraph should not be flagged."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_paragraph('Content')
        doc.add_paragraph('')  # Single empty
        doc.add_paragraph('More content')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_empty_paragraphs(doc)

        assert len(issues) == 0

    @pytest.mark.unit
    def test_two_empty_paragraphs_ok(self, docx_factory):
        """Two consecutive empty paragraphs should not be flagged."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_paragraph('Content')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('More content')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_empty_paragraphs(doc)

        assert len(issues) == 0

    @pytest.mark.unit
    def test_three_empty_paragraphs_flagged(self, docx_factory):
        """Three consecutive empty paragraphs should be flagged."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_paragraph('Content')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('More content')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_empty_paragraphs(doc)

        assert len(issues) == 1
        assert issues[0]['type'] == 'empty_paragraph'
        assert issues[0]['severity'] == 'info'

    @pytest.mark.unit
    def test_multiple_empty_clusters(self, docx_factory):
        """Multiple clusters of empty paragraphs should each be flagged."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_paragraph('Content')
        # First cluster
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('Middle content')
        # Second cluster
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('End content')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_empty_paragraphs(doc)

        assert len(issues) == 2

    @pytest.mark.unit
    def test_whitespace_only_counts_as_empty(self, docx_factory):
        """Paragraphs with only whitespace should count as empty."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_paragraph('Content')
        doc.add_paragraph('   ')  # whitespace only
        doc.add_paragraph('  \t  ')  # tabs and spaces
        doc.add_paragraph('')  # truly empty
        doc.add_paragraph('More content')
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_empty_paragraphs(doc)

        assert len(issues) == 1


# ============================================================================
# Placeholder Detection Tests
# ============================================================================

class TestPlaceholderDetection:
    """Tests for _check_placeholders function."""

    @pytest.mark.unit
    def test_detects_template_variables(self, docx_factory):
        """Should detect {{variable}} patterns."""
        doc_path = docx_factory(content="Hello {{name}}, your order is {{order_id}}.")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 2
        assert all(i['type'] == 'placeholder' for i in issues)
        assert any('{{name}}' in i['message'] for i in issues)
        assert any('{{order_id}}' in i['message'] for i in issues)

    @pytest.mark.unit
    def test_detects_todo_marker(self, docx_factory):
        """Should detect [TODO] markers."""
        doc_path = docx_factory(content="This section needs work [TODO].")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 1
        assert 'TODO' in issues[0]['message']

    @pytest.mark.unit
    def test_detects_tbd_marker(self, docx_factory):
        """Should detect [TBD] markers."""
        doc_path = docx_factory(content="Details to be determined [TBD].")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 1
        assert 'TBD' in issues[0]['message']

    @pytest.mark.unit
    def test_detects_insert_placeholder(self, docx_factory):
        """Should detect [INSERT...] placeholders."""
        doc_path = docx_factory(content="Please [INSERT DATE] here.")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 1
        assert 'INSERT' in issues[0]['message']

    @pytest.mark.unit
    def test_detects_placeholder_marker(self, docx_factory):
        """Should detect [PLACEHOLDER] markers."""
        doc_path = docx_factory(content="Content [PLACEHOLDER] for future.")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 1
        assert 'PLACEHOLDER' in issues[0]['message']

    @pytest.mark.unit
    def test_detects_fill_in_blanks(self, docx_factory):
        """Should detect ____ (4+ underscores) fill-in blanks."""
        doc_path = docx_factory(content="Signature: ____________")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 1
        assert 'Fill-in blanks' in issues[0]['message']

    @pytest.mark.unit
    def test_ignores_short_underscores(self, docx_factory):
        """Should not flag fewer than 4 underscores."""
        doc_path = docx_factory(content="Test_with_underscores and ___three.")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 0

    @pytest.mark.unit
    def test_detects_empty_brackets(self, docx_factory):
        """Should detect empty brackets [ ]."""
        doc_path = docx_factory(content="Check box: [ ]")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 1
        assert 'Empty brackets' in issues[0]['message']

    @pytest.mark.unit
    def test_detects_placeholders_in_tables(self, docx_factory):
        """Should detect placeholders in table cells."""
        doc_path = docx_factory(content="Intro", add_table=True)
        doc = Document(str(doc_path))

        doc.tables[0].cell(1, 0).text = '{{cell_value}}'
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        table_issues = [i for i in issues if 'Table' in i['location']]
        assert len(table_issues) >= 1

    @pytest.mark.unit
    def test_multiple_placeholders_same_paragraph(self, docx_factory):
        """Should detect multiple placeholders in same paragraph."""
        doc_path = docx_factory(content="{{field1}} and {{field2}} with [TODO]")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) == 3

    @pytest.mark.unit
    def test_case_insensitive_detection(self, docx_factory):
        """Placeholder detection should be case-insensitive."""
        doc_path = docx_factory(content="[todo] and [TBD] and [ToDo]")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        # Should detect all three
        assert len(issues) >= 2


# ============================================================================
# Whitespace Tests
# ============================================================================

class TestWhitespaceDetection:
    """Tests for _check_whitespace function."""

    @pytest.mark.unit
    def test_detects_trailing_spaces(self, docx_factory):
        """Should detect trailing spaces."""
        doc_path = docx_factory(content="This has trailing spaces.   ")

        doc = Document(str(doc_path))
        issues = _check_whitespace(doc)

        assert len(issues) >= 1
        assert any('Trailing whitespace' in i['message'] for i in issues)

    @pytest.mark.unit
    def test_detects_multiple_spaces(self, docx_factory):
        """Should detect multiple consecutive spaces."""
        doc_path = docx_factory(content="This  has  multiple  spaces.")

        doc = Document(str(doc_path))
        issues = _check_whitespace(doc)

        assert len(issues) >= 1
        assert any('Multiple consecutive spaces' in i['message'] for i in issues)

    @pytest.mark.unit
    def test_clean_text_no_issues(self, docx_factory):
        """Clean text should have no whitespace issues."""
        doc_path = docx_factory(content="This is clean text with single spaces.")

        doc = Document(str(doc_path))
        issues = _check_whitespace(doc)

        assert len(issues) == 0

    @pytest.mark.unit
    def test_skips_empty_paragraphs(self, docx_factory):
        """Empty paragraphs should be skipped for whitespace checks."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))
        doc.add_paragraph('')
        doc.add_paragraph('   ')  # whitespace only
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        issues = _check_whitespace(doc)

        # No trailing whitespace issues for empty paragraphs
        assert all('Trailing' not in i['message'] for i in issues)

    @pytest.mark.unit
    def test_combined_whitespace_issues(self, docx_factory):
        """Should detect both trailing and multiple spaces."""
        doc_path = docx_factory(content="Multiple  spaces  and trailing.  ")

        doc = Document(str(doc_path))
        issues = _check_whitespace(doc)

        assert len(issues) == 2


# ============================================================================
# Report Formatting Tests
# ============================================================================

class TestReportFormatting:
    """Tests for _format_validation_report function."""

    @pytest.mark.unit
    def test_empty_issues_empty_report(self):
        """Empty issues list should return empty string."""
        report = _format_validation_report([])
        assert report == ""

    @pytest.mark.unit
    def test_single_issue_formatted(self):
        """Single issue should be formatted correctly."""
        issues = [{
            'type': 'placeholder',
            'severity': 'warning',
            'message': 'TODO marker found',
            'location': 'Paragraph 1',
            'preview': 'Test text'
        }]

        report = _format_validation_report(issues)

        assert 'Found 1 issue' in report
        assert 'Placeholder' in report
        assert 'TODO marker found' in report

    @pytest.mark.unit
    def test_groups_by_type(self):
        """Issues should be grouped by type."""
        issues = [
            {'type': 'placeholder', 'severity': 'warning', 'message': 'msg1', 'location': 'loc1', 'preview': ''},
            {'type': 'placeholder', 'severity': 'warning', 'message': 'msg2', 'location': 'loc2', 'preview': ''},
            {'type': 'whitespace', 'severity': 'info', 'message': 'msg3', 'location': 'loc3', 'preview': ''},
        ]

        report = _format_validation_report(issues)

        assert 'Placeholder (2)' in report
        assert 'Whitespace (1)' in report

    @pytest.mark.unit
    def test_truncates_after_five(self):
        """Should show first 5 and indicate more."""
        issues = [
            {'type': 'placeholder', 'severity': 'warning', 'message': f'msg{i}', 'location': f'loc{i}', 'preview': ''}
            for i in range(10)
        ]

        report = _format_validation_report(issues)

        assert '... and 5 more' in report


# ============================================================================
# perform_validation Tests
# ============================================================================

class TestPerformValidation:
    """Tests for the main perform_validation function."""

    @pytest.mark.unit
    def test_all_rules_enabled_by_default(self, docx_factory):
        """All rules except whitespace should be enabled by default."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter', level=1)
        doc.add_heading('Skipped', level=3)  # Hierarchy issue
        doc.add_paragraph('{{placeholder}}')  # Placeholder
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')  # Empty cluster
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        config = {'validation_rules': {}}
        issues = perform_validation(doc, config)

        # Should have hierarchy and placeholder issues
        assert len(issues) >= 2

    @pytest.mark.unit
    def test_whitespace_disabled_by_default(self, docx_factory):
        """Whitespace check should be disabled by default."""
        doc_path = docx_factory(content="Multiple  spaces  here.")

        doc = Document(str(doc_path))
        config = {'validation_rules': {}}
        issues = perform_validation(doc, config)

        whitespace_issues = [i for i in issues if i['type'] == 'whitespace']
        assert len(whitespace_issues) == 0

    @pytest.mark.unit
    def test_whitespace_enabled_when_configured(self, docx_factory):
        """Whitespace check should work when enabled."""
        doc_path = docx_factory(content="Multiple  spaces  here.")

        doc = Document(str(doc_path))
        config = {'validation_rules': {'check_whitespace': True}}
        issues = perform_validation(doc, config)

        whitespace_issues = [i for i in issues if i['type'] == 'whitespace']
        assert len(whitespace_issues) >= 1

    @pytest.mark.unit
    def test_disable_specific_rules(self, docx_factory):
        """Individual rules should be disableable."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter', level=1)
        doc.add_heading('Skipped', level=3)  # Would be flagged
        doc.add_paragraph('{{placeholder}}')  # Would be flagged
        doc.save(str(doc_path))

        doc = Document(str(doc_path))
        config = {
            'validation_rules': {
                'check_heading_hierarchy': False,
                'check_placeholders': False,
            }
        }
        issues = perform_validation(doc, config)

        # Both should be disabled
        hierarchy_issues = [i for i in issues if i['type'] == 'heading_hierarchy']
        placeholder_issues = [i for i in issues if i['type'] == 'placeholder']
        assert len(hierarchy_issues) == 0
        assert len(placeholder_issues) == 0


# ============================================================================
# process_document Integration Tests
# ============================================================================

class TestProcessDocument:
    """Integration tests for process_document function."""

    @pytest.mark.unit
    def test_returns_processor_result(self, docx_factory):
        """Should return a valid ProcessorResult."""
        doc_path = docx_factory(content="Clean document content.")

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert isinstance(result, ProcessorResult)
        assert result.success is True
        assert result.file_path == str(doc_path)

    @pytest.mark.unit
    def test_changes_made_reflects_issue_count(self, docx_factory):
        """changes_made should reflect number of issues found."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_paragraph('{{var1}}')
        doc.add_paragraph('{{var2}}')
        doc.add_paragraph('{{var3}}')
        doc.save(str(doc_path))

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.changes_made == 3

    @pytest.mark.unit
    def test_error_message_contains_report(self, docx_factory):
        """error_message should contain validation report when issues found."""
        doc_path = docx_factory(content="Has a [TODO] here.")

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.error_message is not None
        assert 'Found' in result.error_message
        assert 'issue' in result.error_message

    @pytest.mark.unit
    def test_no_report_when_clean(self, docx_factory):
        """error_message should be None when no issues found."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter 1', level=1)
        doc.add_heading('Section 1.1', level=2)
        doc.add_paragraph('Clean content.')
        doc.save(str(doc_path))

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0
        assert result.error_message is None

    @pytest.mark.unit
    def test_handles_invalid_file(self, tmp_path):
        """Should handle invalid file paths gracefully."""
        invalid_path = str(tmp_path / "nonexistent.docx")

        config = {'validation_rules': {}}
        result = process_document(invalid_path, config)

        assert result.success is False
        assert 'Validation error' in result.error_message

    @pytest.mark.unit
    def test_handles_corrupted_file(self, tmp_path):
        """Should handle corrupted files gracefully."""
        corrupted_file = tmp_path / "corrupted.docx"
        corrupted_file.write_text("Not a valid DOCX")

        config = {'validation_rules': {}}
        result = process_document(str(corrupted_file), config)

        assert result.success is False
        assert result.error_message is not None

    @pytest.mark.unit
    def test_duration_tracked(self, docx_factory):
        """Duration should be tracked in result."""
        doc_path = docx_factory(content="Test content.")

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.duration_seconds >= 0


# ============================================================================
# Edge Cases
# ============================================================================

class TestEdgeCases:
    """Edge case and special scenario tests."""

    @pytest.mark.unit
    def test_empty_document(self, docx_factory):
        """Empty document should be handled gracefully."""
        doc_path = docx_factory(content="")

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0

    @pytest.mark.unit
    def test_document_only_headings(self, docx_factory):
        """Document with only valid headings should pass."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Title', level=0)
        doc.add_heading('Chapter', level=1)
        doc.add_heading('Section', level=2)
        doc.save(str(doc_path))

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 0

    @pytest.mark.unit
    def test_unicode_content(self, docx_factory):
        """Unicode content should be handled correctly."""
        doc_path = docx_factory(content="\u6d4b\u8bd5 {{name}} caf\xe9 [TODO]")

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 2  # placeholder and TODO

    @pytest.mark.unit
    def test_very_long_document(self, docx_factory):
        """Long document should be processed."""
        content = "\n".join([f"Paragraph {i} content." for i in range(100)])
        doc_path = docx_factory(content=content)

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True

    @pytest.mark.unit
    def test_many_tables(self, docx_factory):
        """Document with many tables should be processed."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        for i in range(10):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = f'Table {i} Header'
            table.cell(1, 0).text = '{{value}}'  # Placeholder
        doc.save(str(doc_path))

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        assert result.changes_made == 10  # One placeholder per table

    @pytest.mark.unit
    @pytest.mark.parametrize("placeholder,expected_type", [
        ("{{variable}}", "Template variable"),
        ("[TODO]", "TODO marker"),
        ("[TBD]", "TBD marker"),
        ("[INSERT SOMETHING]", "INSERT placeholder"),
        ("[PLACEHOLDER]", "PLACEHOLDER marker"),
        ("________", "Fill-in blanks"),
        ("[ ]", "Empty brackets"),
    ])
    def test_all_placeholder_types(self, docx_factory, placeholder, expected_type):
        """All placeholder types should be detected."""
        doc_path = docx_factory(content=f"Contains {placeholder} here.")

        doc = Document(str(doc_path))
        issues = _check_placeholders(doc)

        assert len(issues) >= 1
        assert any(expected_type in i['message'] for i in issues)


# ============================================================================
# Complex Scenarios
# ============================================================================

class TestComplexScenarios:
    """Tests for complex real-world scenarios."""

    @pytest.mark.unit
    def test_mixed_issues_all_detected(self, docx_factory):
        """Document with multiple issue types should have all detected."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        # Heading hierarchy issue
        doc.add_heading('Chapter', level=1)
        doc.add_heading('Skipped', level=3)

        # Placeholder
        doc.add_paragraph('{{variable}} here')

        # Empty paragraphs
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')

        # Whitespace (if enabled)
        doc.add_paragraph('Multiple  spaces  and trailing.  ')

        doc.save(str(doc_path))

        config = {'validation_rules': {'check_whitespace': True}}
        result = process_document(str(doc_path), config)

        assert result.success is True
        # Should have: hierarchy(1) + placeholder(1) + empty(1) + whitespace(2)
        assert result.changes_made >= 4

    @pytest.mark.unit
    def test_report_format_comprehensive(self, docx_factory):
        """Report should be well-formatted for complex documents."""
        doc_path = docx_factory(content="")
        doc = Document(str(doc_path))

        doc.add_heading('Chapter', level=1)
        doc.add_heading('Skipped', level=3)
        doc.add_paragraph('{{var1}}')
        doc.add_paragraph('{{var2}}')
        doc.add_paragraph('[TODO]')
        doc.save(str(doc_path))

        config = {'validation_rules': {}}
        result = process_document(str(doc_path), config)

        assert result.error_message is not None
        # Report should have counts
        assert 'Found' in result.error_message
        # Report should group by type
        assert 'Placeholder' in result.error_message or 'placeholder' in result.error_message.lower()
