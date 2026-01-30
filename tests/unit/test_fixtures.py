"""Validation tests for test infrastructure fixtures.

These tests verify the fixtures in conftest.py work correctly,
serving as both validation and documentation of fixture behavior.

NOTE: This file only tests fixtures that DON'T depend on src/ modules.
The test_db and test_config fixtures are blocked until src/database/db_manager.py
and src/core/config.py are implemented.

Working fixtures validated here:
- docx_factory: Creates .docx files programmatically
- sample_docx: Pre-made simple document
- complex_docx: Pre-made document with tables
- large_docx: 50+ paragraph document for performance tests
- multiple_docx: Batch of 10 documents
- temp_workspace: Isolated directory structure
- mock_config: Dictionary-based config (no imports)
- unicode_filename: Path with non-ASCII characters
- Helper functions: assert_docx_contains_text, assert_docx_property

Blocked fixtures (require src/ implementation):
- test_db: Imports src.database.db_manager.DatabaseManager
- test_config: Imports src.core.config.ConfigManager
"""

import pytest
from pathlib import Path
from docx import Document

# Import helper functions from conftest
from conftest import assert_docx_contains_text, assert_docx_property


@pytest.mark.unit
class TestDocxFactory:
    """Tests for docx_factory fixture."""

    def test_creates_file_that_exists(self, docx_factory):
        """docx_factory returns a Path that exists on disk."""
        doc_path = docx_factory(content="Test content")
        assert doc_path.exists(), "Created document should exist"
        assert doc_path.is_file(), "Should be a file, not directory"

    def test_creates_readable_docx(self, docx_factory):
        """Created file is a valid .docx readable by python-docx."""
        doc_path = docx_factory(content="Readable test")
        # This will raise if file is not valid docx
        doc = Document(str(doc_path))
        assert doc is not None

    def test_content_creates_paragraphs(self, docx_factory):
        """Content string creates paragraphs, split on newlines."""
        doc_path = docx_factory(content="Line one\nLine two\nLine three")
        doc = Document(str(doc_path))
        # Filter empty paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 3
        assert paragraphs[0] == "Line one"
        assert paragraphs[1] == "Line two"
        assert paragraphs[2] == "Line three"

    def test_empty_lines_skipped(self, docx_factory):
        """Empty lines in content are not created as paragraphs."""
        doc_path = docx_factory(content="Line one\n\n\nLine two")
        doc = Document(str(doc_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 2

    def test_author_metadata_set(self, docx_factory):
        """Author parameter sets document metadata."""
        doc_path = docx_factory(content="Test", author="Jane Doe")
        doc = Document(str(doc_path))
        assert doc.core_properties.author == "Jane Doe"

    def test_title_metadata_set(self, docx_factory):
        """Title parameter sets document metadata."""
        doc_path = docx_factory(content="Test", title="My Title")
        doc = Document(str(doc_path))
        assert doc.core_properties.title == "My Title"

    def test_add_heading_creates_heading_paragraph(self, docx_factory):
        """add_heading=True creates a heading as first paragraph."""
        doc_path = docx_factory(content="Body text", add_heading=True)
        doc = Document(str(doc_path))
        # First paragraph should be the heading
        assert doc.paragraphs[0].text == "Test Document"
        # Heading style check (level 0 = Title style)
        assert "Title" in doc.paragraphs[0].style.name or "Heading" in doc.paragraphs[0].style.name

    def test_add_table_creates_2x2_table(self, docx_factory):
        """add_table=True creates a 2x2 table with expected content."""
        doc_path = docx_factory(content="Before table", add_table=True)
        doc = Document(str(doc_path))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "Header 1"
        assert table.cell(0, 1).text == "Header 2"
        assert table.cell(1, 0).text == "Data 1"
        assert table.cell(1, 1).text == "Data 2"

    def test_custom_filename(self, docx_factory):
        """filename parameter controls output filename."""
        doc_path = docx_factory(content="Test", filename="custom_name.docx")
        assert doc_path.name == "custom_name.docx"

    def test_returns_path_object(self, docx_factory):
        """Return value is a pathlib.Path object."""
        doc_path = docx_factory(content="Test")
        assert isinstance(doc_path, Path)


@pytest.mark.unit
class TestPreMadeFixtures:
    """Tests for sample_docx and complex_docx fixtures."""

    def test_sample_docx_exists(self, sample_docx):
        """sample_docx fixture creates a file that exists."""
        assert sample_docx.exists()
        assert sample_docx.is_file()

    def test_sample_docx_has_heading_and_paragraphs(self, sample_docx):
        """sample_docx has heading + 3 content paragraphs = 4 total."""
        doc = Document(str(sample_docx))
        # Heading + 3 paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 4
        assert paragraphs[0] == "Test Document"  # Heading
        assert "paragraph one" in paragraphs[1]
        assert "paragraph two" in paragraphs[2]
        assert "paragraph three" in paragraphs[3]

    def test_sample_docx_has_correct_metadata(self, sample_docx):
        """sample_docx has expected author and title metadata."""
        doc = Document(str(sample_docx))
        assert doc.core_properties.author == "Test Author"
        assert doc.core_properties.title == "Sample Document"

    def test_complex_docx_exists(self, complex_docx):
        """complex_docx fixture creates a file that exists."""
        assert complex_docx.exists()
        assert complex_docx.is_file()

    def test_complex_docx_has_table(self, complex_docx):
        """complex_docx contains a table."""
        doc = Document(str(complex_docx))
        assert len(doc.tables) == 1

    def test_complex_docx_has_correct_metadata(self, complex_docx):
        """complex_docx has expected author and title metadata."""
        doc = Document(str(complex_docx))
        assert doc.core_properties.author == "Complex Author"
        assert doc.core_properties.title == "Complex Document"

    def test_large_docx_has_50_plus_paragraphs(self, large_docx):
        """large_docx has 50+ paragraphs for performance testing."""
        doc = Document(str(large_docx))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) >= 50

    def test_multiple_docx_returns_list_of_10(self, multiple_docx):
        """multiple_docx returns list of 10 document paths."""
        assert isinstance(multiple_docx, list)
        assert len(multiple_docx) == 10

    def test_multiple_docx_all_exist(self, multiple_docx):
        """All documents in multiple_docx exist and are readable."""
        for doc_path in multiple_docx:
            assert doc_path.exists()
            # Verify readable
            doc = Document(str(doc_path))
            assert doc is not None


@pytest.mark.unit
class TestTempWorkspace:
    """Tests for temp_workspace fixture."""

    def test_workspace_exists(self, temp_workspace):
        """temp_workspace creates a directory that exists."""
        assert temp_workspace.exists()
        assert temp_workspace.is_dir()

    def test_workspace_has_data_dir(self, temp_workspace):
        """temp_workspace has data/ subdirectory."""
        data_dir = temp_workspace / "data"
        assert data_dir.exists()
        assert data_dir.is_dir()

    def test_workspace_has_backups_dir(self, temp_workspace):
        """temp_workspace has backups/ subdirectory."""
        backups_dir = temp_workspace / "backups"
        assert backups_dir.exists()
        assert backups_dir.is_dir()

    def test_workspace_has_logs_dir(self, temp_workspace):
        """temp_workspace has logs/ subdirectory."""
        logs_dir = temp_workspace / "logs"
        assert logs_dir.exists()
        assert logs_dir.is_dir()


@pytest.mark.unit
class TestMockConfig:
    """Tests for mock_config fixture."""

    def test_mock_config_is_dict(self, mock_config):
        """mock_config returns a dictionary."""
        assert isinstance(mock_config, dict)

    def test_mock_config_has_expected_keys(self, mock_config):
        """mock_config contains expected configuration keys."""
        expected_keys = [
            "backup_enabled",
            "backup_retention_days",
            "max_workers",
            "log_level",
            "database_path",
            "backup_directory",
        ]
        for key in expected_keys:
            assert key in mock_config, f"Expected key '{key}' in mock_config"

    def test_mock_config_uses_memory_db(self, mock_config):
        """mock_config uses in-memory database for test isolation."""
        assert mock_config["database_path"] == ":memory:"


@pytest.mark.unit
class TestUnicodeFilename:
    """Tests for unicode_filename fixture."""

    def test_unicode_filename_is_path(self, unicode_filename):
        """unicode_filename returns a Path object."""
        assert isinstance(unicode_filename, Path)

    def test_unicode_filename_has_non_ascii(self, unicode_filename):
        """unicode_filename contains non-ASCII characters."""
        name = unicode_filename.name
        # Should contain Latin extended, Chinese, and Cyrillic
        assert any(ord(c) > 127 for c in name), "Filename should have non-ASCII chars"

    def test_can_create_docx_with_unicode_name(self, docx_factory, unicode_filename):
        """Can create and read a .docx with Unicode filename."""
        # Create document with unicode filename
        doc_path = docx_factory(
            content="Unicode test content",
            filename=unicode_filename.name
        )
        assert doc_path.exists()
        # Verify readable
        doc = Document(str(doc_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Unicode test content" in paragraphs


@pytest.mark.unit
class TestHelperFunctions:
    """Tests for assertion helper functions."""

    def test_assert_docx_contains_text_finds_text(self, docx_factory):
        """assert_docx_contains_text returns True when text exists."""
        doc_path = docx_factory(content="The quick brown fox")
        result = assert_docx_contains_text(doc_path, "quick brown")
        assert result is True

    def test_assert_docx_contains_text_raises_on_missing(self, docx_factory):
        """assert_docx_contains_text raises AssertionError for missing text."""
        doc_path = docx_factory(content="The quick brown fox")
        with pytest.raises(AssertionError) as exc_info:
            assert_docx_contains_text(doc_path, "lazy dog")
        assert "not found in document" in str(exc_info.value)

    def test_assert_docx_property_validates_author(self, docx_factory):
        """assert_docx_property validates author metadata correctly."""
        doc_path = docx_factory(content="Test", author="John Smith")
        result = assert_docx_property(doc_path, "author", "John Smith")
        assert result is True

    def test_assert_docx_property_validates_title(self, docx_factory):
        """assert_docx_property validates title metadata correctly."""
        doc_path = docx_factory(content="Test", title="Document Title")
        result = assert_docx_property(doc_path, "title", "Document Title")
        assert result is True

    def test_assert_docx_property_raises_on_mismatch(self, docx_factory):
        """assert_docx_property raises AssertionError on property mismatch."""
        doc_path = docx_factory(content="Test", author="John Smith")
        with pytest.raises(AssertionError) as exc_info:
            assert_docx_property(doc_path, "author", "Jane Doe")
        assert "John Smith" in str(exc_info.value)
        assert "Jane Doe" in str(exc_info.value)


# ============================================================================
# Documentation: Blocked fixtures requiring src/ implementation
# ============================================================================

# ============================================================================
# Extended Fixture Tests (Phase 02)
# ============================================================================

@pytest.mark.unit
class TestDocxWithTable:
    """Tests for docx_with_table fixture."""

    def test_creates_default_table(self, docx_with_table):
        """Default call creates 3x2 table with headers."""
        doc_path = docx_with_table()
        doc = Document(str(doc_path))

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 2

        # Verify header row
        assert table.cell(0, 0).text == "Header 1"
        assert table.cell(0, 1).text == "Header 2"

        # Verify data rows
        assert table.cell(1, 0).text == "R1C0"

    def test_custom_dimensions(self, docx_with_table):
        """Can create tables with custom row/column counts."""
        doc_path = docx_with_table(rows=5, cols=4)
        doc = Document(str(doc_path))

        table = doc.tables[0]
        assert len(table.rows) == 5
        assert len(table.columns) == 4

    def test_custom_cell_text(self, docx_with_table):
        """Can use custom cell text template."""
        doc_path = docx_with_table(rows=2, cols=2, cell_text="Cell-{row}-{col}", header_row=False)
        doc = Document(str(doc_path))

        table = doc.tables[0]
        assert table.cell(0, 0).text == "Cell-0-0"
        assert table.cell(1, 1).text == "Cell-1-1"


@pytest.mark.unit
class TestDocxWithMetadata:
    """Tests for docx_with_metadata fixture."""

    def test_creates_document_with_default_metadata(self, docx_with_metadata):
        """Default call creates document with all metadata fields populated."""
        doc_path = docx_with_metadata()
        doc = Document(str(doc_path))

        assert doc.core_properties.author == "Test Author"
        assert doc.core_properties.title == "Test Document"
        assert doc.core_properties.subject == "Test Subject"
        assert doc.core_properties.keywords == "test, keywords"
        assert doc.core_properties.category == "Test Category"
        assert doc.core_properties.comments == "Test comments"

    def test_custom_metadata_values(self, docx_with_metadata):
        """Can create document with custom metadata values."""
        doc_path = docx_with_metadata(
            author="Custom Author",
            title="Custom Title"
        )
        doc = Document(str(doc_path))

        assert doc.core_properties.author == "Custom Author"
        assert doc.core_properties.title == "Custom Title"

    def test_unicode_metadata(self, docx_with_metadata):
        """Handles Unicode in metadata fields."""
        doc_path = docx_with_metadata(
            author="作者名",
            title="Document 文档"
        )
        doc = Document(str(doc_path))

        assert doc.core_properties.author == "作者名"
        assert doc.core_properties.title == "Document 文档"


# ============================================================================
# Documentation: Blocked fixtures requiring src/ implementation
# ============================================================================

@pytest.mark.unit
@pytest.mark.skip(reason="test_db fixture requires src.database.db_manager.DatabaseManager")
class TestTestDb:
    """
    Tests for test_db fixture - BLOCKED.

    These tests are skipped because the test_db fixture imports from
    src.database.db_manager which does not exist yet.

    Implementation needed:
    - src/database/__init__.py
    - src/database/db_manager.py with DatabaseManager class

    Once implemented, remove the skip marker and these tests should pass.
    """

    def test_db_is_initialized(self, test_db):
        """test_db should return initialized DatabaseManager."""
        assert test_db is not None

    def test_db_has_jobs_table(self, test_db):
        """test_db should have jobs table in schema."""
        # Implementation: query sqlite_master for table
        pass

    def test_db_isolation_write(self, test_db):
        """First test writes a value."""
        # Implementation: test_db.execute("INSERT INTO settings ...")
        pass

    def test_db_isolation_read(self, test_db):
        """Second test should NOT see first test's value (isolation)."""
        # Implementation: test_db.execute("SELECT FROM settings ...")
        pass


@pytest.mark.unit
@pytest.mark.skip(reason="test_config fixture requires src.core.config.ConfigManager")
class TestTestConfig:
    """
    Tests for test_config fixture - BLOCKED.

    These tests are skipped because the test_config fixture imports from
    src.core.config which does not exist yet.

    Implementation needed:
    - src/core/__init__.py
    - src/core/config.py with ConfigManager class

    Once implemented, remove the skip marker and these tests should pass.
    """

    def test_config_is_initialized(self, test_config):
        """test_config should return initialized ConfigManager."""
        assert test_config is not None
