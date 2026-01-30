"""
Shared pytest fixtures for DOCX Bulk Editor test suite.

This module provides reusable test fixtures for:
- PySide6 QApplication instance
- Test document generation (programmatic .docx creation)
- Temporary workspace with proper directory structure
- In-memory SQLite database
- Mock configurations
- Windows-specific test utilities
"""

import sys
import pytest
from pathlib import Path
from typing import Callable
from docx import Document
from docx.shared import Pt, Inches
from PySide6.QtWidgets import QApplication


# ============================================================================
# Session-scoped fixtures (created once per test session)
# ============================================================================

@pytest.fixture(scope="session")
def qapp():
    """
    Provide a single QApplication instance for all GUI tests.

    pytest-qt handles this automatically, but we define it explicitly
    for clarity and to ensure proper initialization.

    Scope: session (one instance for entire test run)
    """
    app = QApplication.instance()
    if app is None:
        app = QApplication([])
    yield app
    # pytest-qt handles cleanup automatically


# ============================================================================
# Function-scoped fixtures (created fresh for each test)
# ============================================================================

@pytest.fixture
def temp_workspace(tmp_path):
    """
    Create isolated temporary workspace with application directory structure.

    Creates:
        workspace/
        ├── data/       # For SQLite database
        ├── backups/    # For document backups
        └── logs/       # For application logs

    Args:
        tmp_path: pytest built-in fixture providing temp directory

    Returns:
        Path: Root workspace directory

    Example:
        def test_backup_creation(temp_workspace):
            backup_dir = temp_workspace / "backups"
            assert backup_dir.exists()
    """
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "data").mkdir()
    (workspace / "backups").mkdir()
    (workspace / "logs").mkdir()
    return workspace


@pytest.fixture
def docx_factory(tmp_path):
    """
    Factory fixture for creating test .docx files programmatically.

    Returns a callable that creates custom Word documents with specified
    content, metadata, and formatting.

    Args:
        tmp_path: pytest built-in fixture providing temp directory

    Returns:
        Callable: Function to create .docx files

    Example:
        def test_search_replace(docx_factory):
            doc_path = docx_factory(
                content="Hello world\\nGoodbye world",
                filename="test.docx",
                author="Test Author"
            )
            assert doc_path.exists()
    """
    def _create_docx(
        content: str = "",
        filename: str = "test.docx",
        author: str | None = None,
        title: str | None = None,
        add_table: bool = False,
        add_heading: bool = False,
        **kwargs
    ) -> Path:
        """
        Create a test .docx file with specified content and properties.

        Args:
            content: Text content (paragraphs separated by \\n)
            filename: Output filename
            author: Document author metadata
            title: Document title metadata
            add_table: If True, adds a 2x2 sample table
            add_heading: If True, adds a heading at the top
            **kwargs: Additional python-docx operations

        Returns:
            Path: Absolute path to created .docx file
        """
        doc = Document()

        # Add heading if requested
        if add_heading:
            doc.add_heading("Test Document", level=0)

        # Add content paragraphs
        if content:
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

        # Add sample table if requested
        if add_table:
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Header 1"
            table.cell(0, 1).text = "Header 2"
            table.cell(1, 0).text = "Data 1"
            table.cell(1, 1).text = "Data 2"

        # Set metadata
        if author:
            doc.core_properties.author = author
        if title:
            doc.core_properties.title = title

        # Save to temp directory
        file_path = tmp_path / filename
        doc.save(str(file_path))
        return file_path

    return _create_docx


@pytest.fixture
def sample_docx(docx_factory):
    """
    Pre-created simple test document for quick tests.

    Contains:
    - Heading: "Test Document"
    - 3 paragraphs of plain text
    - Author: "Test Author"
    - Title: "Sample Document"

    Example:
        def test_read_document(sample_docx):
            doc = Document(sample_docx)
            assert len(doc.paragraphs) == 4  # heading + 3 paragraphs
    """
    return docx_factory(
        content="This is paragraph one.\nThis is paragraph two.\nThis is paragraph three.",
        filename="sample.docx",
        author="Test Author",
        title="Sample Document",
        add_heading=True
    )


@pytest.fixture
def complex_docx(docx_factory):
    """
    Pre-created complex test document with tables and mixed formatting.

    Contains:
    - Heading
    - Multiple paragraphs
    - 2x2 table
    - Metadata

    Use for integration tests that need realistic documents.
    """
    return docx_factory(
        content="Introduction paragraph.\nBody paragraph with details.\nConclusion paragraph.",
        filename="complex.docx",
        author="Complex Author",
        title="Complex Document",
        add_heading=True,
        add_table=True
    )


@pytest.fixture
def docx_with_table(tmp_path):
    """Create document with configurable table structure.

    Args:
        tmp_path: pytest built-in fixture

    Returns:
        Callable that creates documents with tables

    Example:
        def test_table_formatting(docx_with_table):
            doc_path = docx_with_table(rows=3, cols=2, cell_text="Cell {row},{col}")
            doc = Document(str(doc_path))
            assert len(doc.tables) == 1
    """
    def _create(
        rows: int = 3,
        cols: int = 2,
        cell_text: str | None = None,
        header_row: bool = True,
        filename: str = "table_test.docx"
    ) -> Path:
        """Create document with table.

        Args:
            rows: Number of rows including header
            cols: Number of columns
            cell_text: Template for cell text, uses {row} and {col} placeholders
                      If None, uses "R{row}C{col}" format
            header_row: If True, first row gets "Header {col}" text
            filename: Output filename
        """
        doc = Document()
        table = doc.add_table(rows=rows, cols=cols)

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if i == 0 and header_row:
                    cell.text = f"Header {j + 1}"
                elif cell_text:
                    cell.text = cell_text.format(row=i, col=j)
                else:
                    cell.text = f"R{i}C{j}"

        file_path = tmp_path / filename
        doc.save(str(file_path))
        return file_path

    return _create


@pytest.fixture
def docx_with_metadata(tmp_path):
    """Create document with pre-populated metadata fields.

    Returns:
        Callable that creates documents with known metadata

    Example:
        def test_metadata_clear(docx_with_metadata):
            doc_path = docx_with_metadata()
            doc = Document(str(doc_path))
            assert doc.core_properties.author == "Test Author"
    """
    def _create(
        author: str = "Test Author",
        title: str = "Test Document",
        subject: str = "Test Subject",
        keywords: str = "test, keywords",
        category: str = "Test Category",
        comments: str = "Test comments",
        content: str = "Document content.",
        filename: str = "metadata_test.docx"
    ) -> Path:
        """Create document with specified metadata.

        All metadata fields have defaults to enable quick testing
        of clear/modify operations.
        """
        doc = Document()

        # Add content
        if content:
            doc.add_paragraph(content)

        # Set all metadata fields
        doc.core_properties.author = author
        doc.core_properties.title = title
        doc.core_properties.subject = subject
        doc.core_properties.keywords = keywords
        doc.core_properties.category = category
        doc.core_properties.comments = comments

        file_path = tmp_path / filename
        doc.save(str(file_path))
        return file_path

    return _create


@pytest.fixture
def mock_config():
    """
    Standard mock configuration for testing.

    Returns dict with typical application settings but optimized for tests
    (e.g., lower worker count, debug logging).

    Example:
        def test_worker_initialization(mock_config):
            worker = JobWorker(config=mock_config)
            assert worker.max_workers == 2
    """
    return {
        "backup_enabled": True,
        "backup_retention_days": 7,
        "max_workers": 2,  # Lower for tests to avoid resource contention
        "log_level": "DEBUG",
        "database_path": ":memory:",  # Use in-memory SQLite for tests
        "backup_directory": None,  # Will be set by individual tests
    }


@pytest.fixture
def test_config(temp_workspace):
    """
    ConfigManager instance with test configuration.

    Creates a ConfigManager with minimal settings for testing,
    using temp_workspace for paths.

    Args:
        temp_workspace: Workspace fixture

    Returns:
        ConfigManager: Initialized configuration manager

    Example:
        def test_main_window(test_config, test_db):
            window = MainWindow(test_config, test_db)
            assert window.config == test_config
    """
    from src.core.config import ConfigManager
    import json

    # Create minimal test config file
    config_data = {
        "backup": {
            "directory": str(temp_workspace / "backups"),
            "retention_days": 7
        },
        "multiprocessing": {
            "pool_size": 2
        },
        "logging": {
            "level": "DEBUG"
        }
    }

    config_file = temp_workspace / "test_config.json"
    with open(config_file, 'w', encoding='utf-8') as f:
        json.dump(config_data, f)

    return ConfigManager(config_path=config_file)


# ============================================================================
# Windows-specific fixtures
# ============================================================================

# Skip decorator for Windows-only tests
windows_only = pytest.mark.skipif(
    sys.platform != "win32",
    reason="Windows-specific test (requires Win32 platform)"
)


@pytest.fixture
def unicode_filename(tmp_path):
    """
    Create a file path with Unicode characters (non-ASCII).

    Tests Windows Unicode handling (cp1252 vs UTF-8 encoding issues).

    Returns:
        Path: Path with Unicode characters in filename

    Example:
        @pytest.mark.windows
        def test_unicode_paths(unicode_filename, docx_factory):
            doc_path = docx_factory(
                content="Test",
                filename=unicode_filename.name
            )
            assert doc_path.exists()
    """
    # Mix of Latin extended, Chinese, and Cyrillic characters
    filename = "tëst_文档_файл.docx"
    return tmp_path / filename


@pytest.fixture
def unc_path(tmp_path):
    """
    Create UNC (Universal Naming Convention) network path for testing.

    Converts local path to \\\\localhost\\C$\\... format to simulate
    network share access.

    Only runs on Windows (skipped on other platforms).

    Returns:
        Path: UNC-formatted path

    Example:
        @pytest.mark.windows
        def test_unc_path_handling(unc_path):
            assert str(unc_path).startswith("\\\\\\\\")
    """
    if sys.platform != "win32":
        pytest.skip("Windows-specific fixture (requires Win32 platform)")
    local_path = str(tmp_path.resolve())

    # Convert C:\... to \\localhost\C$\...
    drive = local_path[0]
    rest = local_path[3:].replace("\\", "/")
    unc = f"\\\\localhost\\{drive}$\\{rest}"

    return Path(unc)


# ============================================================================
# Database fixtures
# ============================================================================

@pytest.fixture
def test_db(temp_workspace):
    """
    Create fresh in-memory SQLite database for each test.

    Database is initialized with schema and destroyed after test completes.
    Each test gets isolated database to prevent cross-test contamination.

    Args:
        temp_workspace: Workspace fixture (for alternative file-based DB)

    Yields:
        DatabaseManager: Initialized database connection

    Example:
        def test_job_creation(test_db):
            job_id = test_db.create_job(operation="search_replace")
            assert job_id > 0

    Note:
        Currently uses :memory: for speed. For integration tests that need
        persistence, can switch to: temp_workspace / "data" / "test.db"
    """
    # Import here to avoid circular dependencies
    from src.database.db_manager import DatabaseManager

    # Create in-memory database for each test
    db = DatabaseManager(":memory:")
    yield db
    db.close()


@pytest.fixture(autouse=True)
def cleanup_workers():
    """
    Automatically cleanup any running workers after each test.

    This prevents multiprocessing pool issues when tests end with workers
    still running in the background.
    """
    yield
    # Cleanup happens after test completes
    import gc
    gc.collect()  # Force garbage collection to cleanup any lingering workers


# ============================================================================
# Performance testing fixtures
# ============================================================================

@pytest.fixture
def large_docx(docx_factory):
    """
    Create large document for performance testing.

    Contains 50+ paragraphs to simulate real-world document sizes.
    Use with @pytest.mark.performance decorator.

    Returns:
        Path: Path to large test document
    """
    content = "\n".join([f"Paragraph {i} with some test content." for i in range(50)])
    return docx_factory(
        content=content,
        filename="large.docx",
        title="Performance Test Document"
    )


@pytest.fixture
def multiple_docx(docx_factory, tmp_path):
    """
    Create multiple test documents for batch processing tests.

    Args:
        docx_factory: Document factory fixture
        tmp_path: Temporary directory

    Returns:
        list[Path]: List of 10 test document paths

    Example:
        def test_batch_processing(multiple_docx):
            assert len(multiple_docx) == 10
            for doc_path in multiple_docx:
                assert doc_path.exists()
    """
    docs = []
    for i in range(10):
        doc_path = docx_factory(
            content=f"Document {i} content.\nAnother paragraph.",
            filename=f"doc_{i:02d}.docx",
            title=f"Document {i}"
        )
        docs.append(doc_path)
    return docs


# ============================================================================
# Utility functions available to all tests
# ============================================================================

def assert_docx_contains_text(docx_path: Path, expected_text: str) -> bool:
    """
    Helper to verify text exists in a .docx file.

    Searches all paragraphs for the expected text.

    Args:
        docx_path: Path to .docx file
        expected_text: Text to search for

    Returns:
        bool: True if text found

    Raises:
        AssertionError: If text not found
    """
    doc = Document(str(docx_path))
    full_text = "\n".join([para.text for para in doc.paragraphs])
    assert expected_text in full_text, f"Expected text '{expected_text}' not found in document"
    return True


def assert_docx_property(docx_path: Path, property_name: str, expected_value: str) -> bool:
    """
    Helper to verify document metadata properties.

    Args:
        docx_path: Path to .docx file
        property_name: Property name (e.g., 'author', 'title')
        expected_value: Expected property value

    Returns:
        bool: True if property matches

    Raises:
        AssertionError: If property doesn't match
    """
    doc = Document(str(docx_path))
    actual_value = getattr(doc.core_properties, property_name)
    assert actual_value == expected_value, f"Property '{property_name}' = '{actual_value}', expected '{expected_value}'"
    return True


# Make helper functions available to all tests
pytest.assert_docx_contains_text = assert_docx_contains_text
pytest.assert_docx_property = assert_docx_property
