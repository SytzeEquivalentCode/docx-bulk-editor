"""Script to create test DOCX files with various formatting."""

from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path


def create_formatted_text_doc():
    """Create document with formatted text for testing."""
    doc = Document()

    # Paragraph with bold text
    p1 = doc.add_paragraph()
    run1 = p1.add_run('Hello')
    run1.bold = True
    p1.add_run(' world')

    # Paragraph with mixed formatting
    p2 = doc.add_paragraph()
    run2 = p2.add_run('This is ')
    run3 = p2.add_run('important')
    run3.bold = True
    run3.italic = True
    p2.add_run(' text')

    # Paragraph with font sizes
    p3 = doc.add_paragraph()
    run4 = p3.add_run('Small ')
    run4.font.size = Pt(8)
    run5 = p3.add_run('medium ')
    run5.font.size = Pt(12)
    run6 = p3.add_run('large')
    run6.font.size = Pt(16)

    # Paragraph with colors
    p4 = doc.add_paragraph()
    run7 = p4.add_run('Red ')
    run7.font.color.rgb = RGBColor(255, 0, 0)
    run8 = p4.add_run('green ')
    run8.font.color.rgb = RGBColor(0, 255, 0)
    run9 = p4.add_run('blue')
    run9.font.color.rgb = RGBColor(0, 0, 255)

    # Paragraph with underline
    p5 = doc.add_paragraph()
    run10 = p5.add_run('Underlined text')
    run10.underline = True

    doc.save('tests/test_documents/formatted_text.docx')
    print("Created: formatted_text.docx")


def create_table_doc():
    """Create document with table for testing."""
    doc = Document()

    doc.add_paragraph('Document with table')

    # Create 2x2 table
    table = doc.add_table(rows=2, cols=2)

    # Add content to cells
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'

    doc.save('tests/test_documents/table_doc.docx')
    print("Created: table_doc.docx")


def create_header_footer_doc():
    """Create document with headers and footers."""
    doc = Document()

    # Add content
    doc.add_paragraph('Main content')

    # Add header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = 'This is the header'

    # Add footer
    footer = section.footer
    footer.paragraphs[0].text = 'This is the footer'

    doc.save('tests/test_documents/header_footer_doc.docx')
    print("Created: header_footer_doc.docx")


def create_unicode_doc():
    """Create document with Unicode characters."""
    doc = Document()

    # Chinese characters
    doc.add_paragraph('测试文本 test 测试')

    # Emoji
    doc.add_paragraph('Hello 😀 World 🌟 Test')

    # Accented characters
    doc.add_paragraph('Café résumé naïve')

    doc.save('tests/test_documents/unicode_doc.docx')
    print("Created: unicode_doc.docx")


def create_empty_doc():
    """Create empty document."""
    doc = Document()
    doc.add_paragraph('')  # Empty paragraph
    doc.save('tests/test_documents/empty_doc.docx')
    print("Created: empty_doc.docx")


def create_simple_doc():
    """Create simple document for basic tests."""
    doc = Document()
    doc.add_paragraph('This is a test document.')
    doc.add_paragraph('It contains multiple paragraphs.')
    doc.add_paragraph('Testing search and replace functionality.')
    doc.save('tests/test_documents/simple_doc.docx')
    print("Created: simple_doc.docx")


if __name__ == '__main__':
    Path('tests/test_documents').mkdir(parents=True, exist_ok=True)

    create_formatted_text_doc()
    create_table_doc()
    create_header_footer_doc()
    create_unicode_doc()
    create_empty_doc()
    create_simple_doc()

    print("\nAll test documents created successfully!")
