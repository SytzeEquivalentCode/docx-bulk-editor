"""Script to create specialized test DOCX files for processor testing.

This script creates test documents specifically designed to test:
- Style Enforcer: varied fonts, heading detection, spacing issues
- Validator: heading hierarchy issues, placeholders, whitespace problems
- Table Formatter: various table configurations

Run from project root:
    python tests/test_documents/create_specialized_test_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


# ============================================================================
# Style Enforcer Test Documents
# ============================================================================

def create_mixed_fonts_doc(output_dir: Path):
    """Create document with varied fonts for standardization testing."""
    doc = Document()

    # Paragraph with Arial 10pt
    p1 = doc.add_paragraph()
    run1 = p1.add_run('This text is Arial 10pt.')
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)

    # Paragraph with Times New Roman 14pt
    p2 = doc.add_paragraph()
    run2 = p2.add_run('This text is Times New Roman 14pt.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(14)

    # Paragraph with Calibri 11pt (standard, should not change)
    p3 = doc.add_paragraph()
    run3 = p3.add_run('This text is Calibri 11pt (standard).')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(11)

    # Paragraph with mixed fonts in same paragraph
    p4 = doc.add_paragraph()
    run4a = p4.add_run('This is Arial ')
    run4a.font.name = 'Arial'
    run4a.font.size = Pt(12)
    run4b = p4.add_run('and this is Courier ')
    run4b.font.name = 'Courier New'
    run4b.font.size = Pt(10)
    run4c = p4.add_run('mixed in one paragraph.')
    run4c.font.name = 'Verdana'
    run4c.font.size = Pt(9)

    # Table with varied fonts
    table = doc.add_table(rows=2, cols=2)
    cell00 = table.cell(0, 0)
    run_t1 = cell00.paragraphs[0].add_run('Arial in table')
    run_t1.font.name = 'Arial'
    run_t1.font.size = Pt(8)

    cell01 = table.cell(0, 1)
    run_t2 = cell01.paragraphs[0].add_run('Times in table')
    run_t2.font.name = 'Times New Roman'
    run_t2.font.size = Pt(16)

    doc.save(output_dir / 'mixed_fonts_doc.docx')
    print("Created: mixed_fonts_doc.docx")


def create_potential_headings_doc(output_dir: Path):
    """Create document with potential headings for detection testing."""
    doc = Document()

    # Bold short text (should be detected as heading)
    p1 = doc.add_paragraph()
    run1 = p1.add_run('Introduction')
    run1.bold = True

    # Normal body text
    doc.add_paragraph('This is normal body text that follows the heading.')

    # Large font short text (should be detected as H1)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Chapter One Title')
    run2.font.size = Pt(16)

    doc.add_paragraph('More body text after the chapter title.')

    # Medium font short text (should be detected as H2)
    p3 = doc.add_paragraph()
    run3 = p3.add_run('Section 1.1')
    run3.font.size = Pt(14)
    run3.bold = True

    doc.add_paragraph('Body text in section 1.1.')

    # Smaller bold text (should be H3)
    p4 = doc.add_paragraph()
    run4 = p4.add_run('Subsection Details')
    run4.font.size = Pt(12)
    run4.bold = True

    doc.add_paragraph('Details in the subsection.')

    # Long bold text (should NOT be detected as heading - too long)
    p5 = doc.add_paragraph()
    run5 = p5.add_run('This is a very long bold text that should not be detected as a heading because it exceeds the character limit for heading detection in the style enforcer.')
    run5.bold = True

    # Already a heading (should be skipped)
    doc.add_heading('Existing Heading', level=2)

    doc.add_paragraph('Final paragraph.')

    doc.save(output_dir / 'potential_headings_doc.docx')
    print("Created: potential_headings_doc.docx")


def create_inconsistent_spacing_doc(output_dir: Path):
    """Create document with inconsistent paragraph spacing."""
    doc = Document()

    # Paragraph with no spacing
    p1 = doc.add_paragraph('Paragraph with no spacing after.')
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)

    # Paragraph with excessive spacing
    p2 = doc.add_paragraph('Paragraph with 20pt spacing after.')
    p2.paragraph_format.space_after = Pt(20)
    p2.paragraph_format.space_before = Pt(10)

    # Paragraph with standard spacing (8pt after)
    p3 = doc.add_paragraph('Paragraph with standard 8pt spacing.')
    p3.paragraph_format.space_after = Pt(8)
    p3.paragraph_format.space_before = Pt(0)

    # Paragraph with odd line spacing
    p4 = doc.add_paragraph('Paragraph with 2.0 line spacing.')
    p4.paragraph_format.line_spacing = 2.0

    # Paragraph with single line spacing
    p5 = doc.add_paragraph('Paragraph with single line spacing.')
    p5.paragraph_format.line_spacing = 1.0

    # Add heading (should be skipped by spacing normalization)
    doc.add_heading('Section Heading', level=1)

    doc.add_paragraph('Normal paragraph after heading.')

    doc.save(output_dir / 'inconsistent_spacing_doc.docx')
    print("Created: inconsistent_spacing_doc.docx")


def create_style_enforcer_complex_doc(output_dir: Path):
    """Create complex document for comprehensive style enforcement testing."""
    doc = Document()

    # Title with heading style
    doc.add_heading('Document Title', level=0)

    # Introduction - bold potential heading
    p_intro = doc.add_paragraph()
    run_intro = p_intro.add_run('Introduction')
    run_intro.bold = True
    run_intro.font.size = Pt(14)

    # Body with mixed fonts
    p1 = doc.add_paragraph()
    run1a = p1.add_run('This document contains ')
    run1a.font.name = 'Arial'
    run1b = p1.add_run('mixed formatting ')
    run1b.font.name = 'Times New Roman'
    run1b.italic = True
    run1c = p1.add_run('for testing.')
    run1c.font.name = 'Verdana'

    # Table with various fonts
    doc.add_paragraph('Table with varied formatting:')
    table = doc.add_table(rows=3, cols=3)

    fonts = ['Arial', 'Times New Roman', 'Courier New']
    sizes = [8, 11, 14]

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            run = cell.paragraphs[0].add_run(f'R{row_idx}C{col_idx}')
            run.font.name = fonts[col_idx]
            run.font.size = Pt(sizes[row_idx])

    # Another potential heading
    p_section = doc.add_paragraph()
    run_section = p_section.add_run('Methods Section')
    run_section.bold = True
    run_section.font.size = Pt(16)

    # Paragraphs with inconsistent spacing
    p2 = doc.add_paragraph('Paragraph with excessive spacing.')
    p2.paragraph_format.space_after = Pt(24)

    p3 = doc.add_paragraph('Paragraph with no spacing.')
    p3.paragraph_format.space_after = Pt(0)

    doc.save(output_dir / 'style_enforcer_complex_doc.docx')
    print("Created: style_enforcer_complex_doc.docx")


# ============================================================================
# Validator Test Documents
# ============================================================================

def create_heading_hierarchy_issues_doc(output_dir: Path):
    """Create document with heading hierarchy violations."""
    doc = Document()

    # H1 - good start
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_paragraph('Introduction content.')

    # H2 - good
    doc.add_heading('Section 1.1: Background', level=2)
    doc.add_paragraph('Background content.')

    # H4 - BAD: skipped H3
    doc.add_heading('Detail 1.1.1.1: Specific', level=4)
    doc.add_paragraph('Specific detail content.')

    # H2 - back to good
    doc.add_heading('Section 1.2: Overview', level=2)
    doc.add_paragraph('Overview content.')

    # H3 - good
    doc.add_heading('Subsection 1.2.1: Details', level=3)
    doc.add_paragraph('Details content.')

    # H1 - new chapter
    doc.add_heading('Chapter 2: Methods', level=1)
    doc.add_paragraph('Methods content.')

    # H3 - BAD: skipped H2
    doc.add_heading('Sub-Method A', level=3)
    doc.add_paragraph('Sub-method content.')

    doc.save(output_dir / 'heading_hierarchy_issues_doc.docx')
    print("Created: heading_hierarchy_issues_doc.docx")


def create_placeholder_doc(output_dir: Path):
    """Create document with various placeholder patterns."""
    doc = Document()

    doc.add_heading('Document with Placeholders', level=1)

    # Template variables
    doc.add_paragraph('Dear {{customer_name}},')
    doc.add_paragraph('Your order {{order_id}} has been processed.')

    # TODO markers
    doc.add_paragraph('This section needs review [TODO].')
    doc.add_paragraph('Details to be determined [TBD].')

    # INSERT placeholders
    doc.add_paragraph('Please [INSERT DATE] here.')
    doc.add_paragraph('Add [INSERT NAME] to the document.')

    # PLACEHOLDER marker
    doc.add_paragraph('Content [PLACEHOLDER] for future use.')

    # XML-style placeholders
    doc.add_paragraph('The <company_placeholder> logo goes here.')
    doc.add_paragraph('Insert <value_placeholder> below.')

    # Fill-in blanks (underscores)
    doc.add_paragraph('Signature: ____________')
    doc.add_paragraph('Date: ____________________')

    # Empty brackets
    doc.add_paragraph('Check if complete: [ ]')
    doc.add_paragraph('Status: [  ]')

    # Multiple placeholders in one paragraph
    doc.add_paragraph('{{field1}} and {{field2}} with [TODO] and ______')

    # Table with placeholders
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Name'
    table.cell(0, 1).text = '{{name}}'
    table.cell(1, 0).text = 'Date'
    table.cell(1, 1).text = '[INSERT DATE]'

    doc.save(output_dir / 'placeholder_doc.docx')
    print("Created: placeholder_doc.docx")


def create_empty_paragraphs_doc(output_dir: Path):
    """Create document with multiple consecutive empty paragraphs."""
    doc = Document()

    doc.add_heading('Document with Empty Paragraphs', level=1)

    doc.add_paragraph('First paragraph with content.')

    # Single empty paragraph (OK)
    doc.add_paragraph('')

    doc.add_paragraph('After one empty paragraph.')

    # Two empty paragraphs (OK)
    doc.add_paragraph('')
    doc.add_paragraph('')

    doc.add_paragraph('After two empty paragraphs.')

    # Three empty paragraphs (should trigger warning)
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    doc.add_paragraph('After three empty paragraphs.')

    # Five empty paragraphs (should trigger warning)
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    doc.add_paragraph('After five empty paragraphs.')

    doc.save(output_dir / 'empty_paragraphs_doc.docx')
    print("Created: empty_paragraphs_doc.docx")


def create_whitespace_issues_doc(output_dir: Path):
    """Create document with various whitespace issues."""
    doc = Document()

    doc.add_heading('Document with Whitespace Issues', level=1)

    # Trailing whitespace
    doc.add_paragraph('This line has trailing spaces.   ')
    doc.add_paragraph('This line has trailing tabs.\t\t')

    # Multiple consecutive spaces
    doc.add_paragraph('This  line  has  multiple  spaces.')
    doc.add_paragraph('Extra   spaces    throughout.')

    # Combined issues
    doc.add_paragraph('Multiple  spaces  and trailing.  ')

    # Clean paragraph for comparison
    doc.add_paragraph('This is a clean paragraph.')

    # Whitespace only
    doc.add_paragraph('   ')

    doc.save(output_dir / 'whitespace_issues_doc.docx')
    print("Created: whitespace_issues_doc.docx")


def create_validator_complex_doc(output_dir: Path):
    """Create complex document with multiple validation issues."""
    doc = Document()

    # H1
    doc.add_heading('Annual Report 2024', level=1)

    # Content with placeholder
    doc.add_paragraph('Prepared by {{author_name}} on {{date}}.')

    # H3 - skipped H2 (hierarchy issue)
    doc.add_heading('Financial Summary', level=3)

    doc.add_paragraph('Total revenue: ____________')
    doc.add_paragraph('Net profit: [TBD]')

    # Empty paragraphs cluster
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    # H2
    doc.add_heading('Operations', level=2)

    # Multiple spaces and trailing
    doc.add_paragraph('Operations  overview  with extra spaces.  ')

    # Table with placeholders
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Metric'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Growth Rate'
    table.cell(1, 1).text = '{{growth_rate}}%'

    # H4 - skipped H3 (hierarchy issue)
    doc.add_heading('Specific Metric', level=4)

    doc.add_paragraph('[INSERT detailed analysis here]')

    doc.save(output_dir / 'validator_complex_doc.docx')
    print("Created: validator_complex_doc.docx")


def create_clean_validation_doc(output_dir: Path):
    """Create document that passes all validation rules."""
    doc = Document()

    doc.add_heading('Clean Document', level=1)
    doc.add_paragraph('This document should pass all validation rules.')

    doc.add_heading('Section 1', level=2)
    doc.add_paragraph('Content for section 1.')

    doc.add_heading('Subsection 1.1', level=3)
    doc.add_paragraph('Content for subsection 1.1.')

    doc.add_heading('Section 2', level=2)
    doc.add_paragraph('Content for section 2.')

    doc.save(output_dir / 'clean_validation_doc.docx')
    print("Created: clean_validation_doc.docx")


# ============================================================================
# Table Formatter Test Documents
# ============================================================================

def create_unformatted_table_doc(output_dir: Path):
    """Create document with unformatted tables."""
    doc = Document()

    doc.add_heading('Tables for Formatting', level=1)

    # Simple 3x3 table without formatting
    doc.add_paragraph('Simple 3x3 table:')
    table1 = doc.add_table(rows=3, cols=3)
    for row_idx, row in enumerate(table1.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f'R{row_idx+1}C{col_idx+1}'

    doc.add_paragraph('')  # Spacer

    # 5x4 table with header row
    doc.add_paragraph('5x4 table with header:')
    table2 = doc.add_table(rows=5, cols=4)
    headers = ['ID', 'Name', 'Value', 'Status']
    for col_idx, header in enumerate(headers):
        table2.cell(0, col_idx).text = header

    for row_idx in range(1, 5):
        table2.cell(row_idx, 0).text = str(row_idx)
        table2.cell(row_idx, 1).text = f'Item {row_idx}'
        table2.cell(row_idx, 2).text = f'${row_idx * 100}'
        table2.cell(row_idx, 3).text = 'Active'

    doc.add_paragraph('')  # Spacer

    # Single row table
    doc.add_paragraph('Single row table:')
    table3 = doc.add_table(rows=1, cols=4)
    for col_idx in range(4):
        table3.cell(0, col_idx).text = f'Header {col_idx+1}'

    # Single column table
    doc.add_paragraph('Single column table:')
    table4 = doc.add_table(rows=4, cols=1)
    for row_idx in range(4):
        table4.cell(row_idx, 0).text = f'Row {row_idx+1}'

    doc.save(output_dir / 'unformatted_table_doc.docx')
    print("Created: unformatted_table_doc.docx")


def create_multiple_tables_doc(output_dir: Path):
    """Create document with multiple tables for batch formatting."""
    doc = Document()

    doc.add_heading('Multiple Tables Document', level=1)

    for i in range(5):
        doc.add_paragraph(f'Table {i+1}:')
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = 'Header A'
        table.cell(0, 1).text = 'Header B'
        table.cell(0, 2).text = 'Header C'
        table.cell(1, 0).text = f'Data {i}A'
        table.cell(1, 1).text = f'Data {i}B'
        table.cell(1, 2).text = f'Data {i}C'
        table.cell(2, 0).text = f'More {i}A'
        table.cell(2, 1).text = f'More {i}B'
        table.cell(2, 2).text = f'More {i}C'
        doc.add_paragraph('')  # Spacer

    doc.save(output_dir / 'multiple_tables_doc.docx')
    print("Created: multiple_tables_doc.docx")


# ============================================================================
# Edge Case Test Documents
# ============================================================================

def create_empty_document(output_dir: Path):
    """Create completely empty document."""
    doc = Document()
    doc.save(output_dir / 'truly_empty_doc.docx')
    print("Created: truly_empty_doc.docx")


def create_only_headings_doc(output_dir: Path):
    """Create document with only headings, no body text."""
    doc = Document()

    doc.add_heading('Title', level=0)
    doc.add_heading('Chapter 1', level=1)
    doc.add_heading('Section 1.1', level=2)
    doc.add_heading('Section 1.2', level=2)
    doc.add_heading('Chapter 2', level=1)
    doc.add_heading('Section 2.1', level=2)

    doc.save(output_dir / 'only_headings_doc.docx')
    print("Created: only_headings_doc.docx")


def create_unicode_intensive_doc(output_dir: Path):
    """Create document with intensive Unicode content."""
    doc = Document()

    doc.add_heading('Unicode Test Document', level=1)

    # Chinese
    doc.add_paragraph('Chinese: \u6d4b\u8bd5\u6587\u672c\u3001\u4e2d\u6587\u5185\u5bb9\u3001\u4e2d\u534e\u4eba\u6c11\u5171\u548c\u56fd')

    # Japanese
    doc.add_paragraph('Japanese: \u30c6\u30b9\u30c8\u3001\u65e5\u672c\u8a9e\u3001\u3053\u3093\u306b\u3061\u306f')

    # Korean
    doc.add_paragraph('Korean: \ud14c\uc2a4\ud2b8\u3001\ud55c\uae00\u3001\uc548\ub155\ud558\uc138\uc694')

    # Russian
    doc.add_paragraph('Russian: \u0422\u0435\u0441\u0442\u043e\u0432\u044b\u0439 \u0442\u0435\u043a\u0441\u0442\u3001\u041f\u0440\u0438\u0432\u0435\u0442')

    # Arabic
    doc.add_paragraph('Arabic: \u0627\u062e\u062a\u0628\u0627\u0631 \u0627\u0644\u0646\u0635')

    # Mixed with simple symbols (avoiding emoji surrogates that lxml can't handle)
    doc.add_paragraph('Mixed: Hello World Test \u2713 Done \u2714')

    # Special characters
    doc.add_paragraph('Special: \xa9 \xae \u2122 \u20ac \xa3 \xa5 \u221e \u03c0')

    # Accented European
    doc.add_paragraph('European: caf\xe9 r\xe9sum\xe9 na\xefve \xfcber stra\xdfe \xe6\xf8\xe5')

    # Math symbols
    doc.add_paragraph('Math: \u221a \u00b1 \u2264 \u2265 \u2260 \u2248 \u03b1 \u03b2 \u03b3')

    doc.save(output_dir / 'unicode_intensive_doc.docx')
    print("Created: unicode_intensive_doc.docx")


# ============================================================================
# Main execution
# ============================================================================

def main():
    """Create all specialized test documents."""
    output_dir = Path('tests/test_documents')
    output_dir.mkdir(parents=True, exist_ok=True)

    print("Creating specialized test documents...\n")

    # Style Enforcer test documents
    print("=== Style Enforcer Documents ===")
    create_mixed_fonts_doc(output_dir)
    create_potential_headings_doc(output_dir)
    create_inconsistent_spacing_doc(output_dir)
    create_style_enforcer_complex_doc(output_dir)

    # Validator test documents
    print("\n=== Validator Documents ===")
    create_heading_hierarchy_issues_doc(output_dir)
    create_placeholder_doc(output_dir)
    create_empty_paragraphs_doc(output_dir)
    create_whitespace_issues_doc(output_dir)
    create_validator_complex_doc(output_dir)
    create_clean_validation_doc(output_dir)

    # Table Formatter test documents
    print("\n=== Table Formatter Documents ===")
    create_unformatted_table_doc(output_dir)
    create_multiple_tables_doc(output_dir)

    # Edge case documents
    print("\n=== Edge Case Documents ===")
    create_empty_document(output_dir)
    create_only_headings_doc(output_dir)
    create_unicode_intensive_doc(output_dir)

    print("\n" + "=" * 50)
    print("All specialized test documents created successfully!")
    print(f"Location: {output_dir.resolve()}")


if __name__ == '__main__':
    main()
