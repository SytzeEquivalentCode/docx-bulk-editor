"""Table formatting processor for DOCX documents.

Applies consistent formatting to tables:
- Border standardization
- Header row background
- Zebra striping (alternating row colors)
- Cell alignment
- Cell padding
"""

import time
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.processors import ProcessorResult


def process_document(file_path: str, config: Dict[str, Any]) -> ProcessorResult:
    """Top-level function for multiprocessing (must be module-level).

    Args:
        file_path: Path to the DOCX file to process
        config: Configuration dictionary with table formatting options

    Returns:
        ProcessorResult with success status and tables formatted count
    """
    start_time = time.time()

    try:
        doc = Document(file_path)
        changes = perform_table_formatting(doc, config)

        if changes > 0:
            doc.save(file_path)

        duration = time.time() - start_time
        return ProcessorResult(
            success=True,
            file_path=file_path,
            changes_made=changes,
            duration_seconds=duration
        )
    except Exception as e:
        duration = time.time() - start_time
        return ProcessorResult(
            success=False,
            file_path=file_path,
            changes_made=0,
            error_message=str(e),
            duration_seconds=duration
        )


def perform_table_formatting(doc: Document, config: Dict[str, Any]) -> int:
    """Apply formatting to all tables in the document.

    Args:
        doc: python-docx Document object
        config: Configuration dictionary containing:
            - table_options: Dict with formatting options
                - standardize_borders: Apply consistent borders
                - header_background: Add gray background to first row
                - zebra_striping: Alternate row colors
                - alignment: Cell text alignment ('left', 'center', 'right')
                - padding_pt: Cell padding in points

    Returns:
        Number of tables formatted
    """
    table_options = config.get('table_options', {})
    tables_formatted = 0

    for table in doc.tables:
        table_changed = False

        # Standardize borders
        if table_options.get('standardize_borders', True):
            _apply_table_borders(table)
            table_changed = True

        # Header background
        if table_options.get('header_background', True) and len(table.rows) > 0:
            _apply_header_background(table.rows[0])
            table_changed = True

        # Zebra striping
        if table_options.get('zebra_striping', False):
            _apply_zebra_striping(table)
            table_changed = True

        # Cell alignment
        alignment = table_options.get('alignment', 'left')
        if alignment:
            _apply_cell_alignment(table, alignment)
            table_changed = True

        # Cell padding
        padding_pt = table_options.get('padding_pt', 5)
        if padding_pt > 0:
            _apply_cell_padding(table, padding_pt)
            table_changed = True

        if table_changed:
            tables_formatted += 1

    return tables_formatted


def _apply_table_borders(table):
    """Apply consistent black borders to all table cells."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # Create borders element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    # Remove existing borders if any
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    tblPr.append(tblBorders)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _apply_header_background(row):
    """Apply light gray background to header row."""
    for cell in row.cells:
        _set_cell_shading(cell, 'D9D9D9')  # Light gray


def _apply_zebra_striping(table):
    """Apply alternating row colors (skip header row)."""
    for idx, row in enumerate(table.rows):
        if idx == 0:
            continue  # Skip header
        if idx % 2 == 0:
            # Even rows - light gray
            for cell in row.cells:
                _set_cell_shading(cell, 'F2F2F2')
        else:
            # Odd rows - white (clear shading)
            for cell in row.cells:
                _set_cell_shading(cell, 'FFFFFF')


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing shading
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)

    # Create new shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def _apply_cell_alignment(table, alignment: str):
    """Apply text alignment to all cells."""
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT
    }
    align_value = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = align_value


def _apply_cell_padding(table, padding_pt: int):
    """Apply cell padding (margins) to all cells."""
    padding_twips = padding_pt * 20  # Convert points to twips (1 pt = 20 twips)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Create or get margins element
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                tcMar = OxmlElement('w:tcMar')
                tcPr.append(tcMar)
            else:
                # Clear existing margins
                for child in list(tcMar):
                    tcMar.remove(child)

            # Set all margins
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), str(padding_twips))
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
