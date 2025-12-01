"""Document validation processor for DOCX documents.

Validates documents against configurable rules:
- Heading hierarchy (no skipped levels)
- Empty paragraphs detection
- Placeholder text detection ({{var}}, [TODO], etc.)
- Whitespace issues (trailing spaces, multiple spaces)
"""

import re
import time
from typing import Dict, Any, List
from docx import Document

from src.processors import ProcessorResult


def process_document(file_path: str, config: Dict[str, Any]) -> ProcessorResult:
    """Top-level function for multiprocessing (must be module-level).

    Args:
        file_path: Path to the DOCX file to validate
        config: Configuration dictionary with validation options

    Returns:
        ProcessorResult with success status and issues count (as changes_made)
        Note: For validation, changes_made represents issues found
    """
    start_time = time.time()

    try:
        doc = Document(file_path)
        issues = perform_validation(doc, config)

        duration = time.time() - start_time

        # For validation, we report issues found but consider it "successful"
        # if the document was readable and validatable
        # Use error_message to store the validation report
        report = _format_validation_report(issues)

        return ProcessorResult(
            success=True,
            file_path=file_path,
            changes_made=len(issues),  # Number of issues found
            error_message=report if issues else None,
            duration_seconds=duration
        )
    except Exception as e:
        duration = time.time() - start_time
        return ProcessorResult(
            success=False,
            file_path=file_path,
            changes_made=0,
            error_message=f"Validation error: {str(e)}",
            duration_seconds=duration
        )


def perform_validation(doc: Document, config: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Validate document against configured rules.

    Args:
        doc: python-docx Document object
        config: Configuration dictionary containing:
            - validation_rules: Dict with boolean flags for each rule type
                - check_heading_hierarchy: Check for skipped heading levels
                - check_empty_paragraphs: Find empty paragraphs
                - check_placeholders: Find placeholder text like {{var}}, [TODO]
                - check_whitespace: Find trailing/multiple spaces

    Returns:
        List of issues found, each a dict with 'type', 'message', 'location'
    """
    issues = []
    validation_rules = config.get('validation_rules', {})

    # Check heading hierarchy
    if validation_rules.get('check_heading_hierarchy', True):
        issues.extend(_check_heading_hierarchy(doc))

    # Check empty paragraphs
    if validation_rules.get('check_empty_paragraphs', True):
        issues.extend(_check_empty_paragraphs(doc))

    # Check placeholders
    if validation_rules.get('check_placeholders', True):
        issues.extend(_check_placeholders(doc))

    # Check whitespace issues
    if validation_rules.get('check_whitespace', False):
        issues.extend(_check_whitespace(doc))

    return issues


def _check_heading_hierarchy(doc: Document) -> List[Dict[str, Any]]:
    """Check for skipped heading levels.

    Valid: H1 -> H2 -> H3
    Invalid: H1 -> H3 (skipped H2)

    Returns:
        List of heading hierarchy issues
    """
    issues = []
    last_level = 0
    para_index = 0

    for para in doc.paragraphs:
        para_index += 1
        style_name = para.style.name if para.style else ''

        # Check for heading styles (Heading 1, Heading 2, etc.)
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading ', ''))

                # Check if level is skipped
                if last_level > 0 and level > last_level + 1:
                    preview = para.text[:50] + '...' if len(para.text) > 50 else para.text
                    issues.append({
                        'type': 'heading_hierarchy',
                        'severity': 'warning',
                        'message': f"Skipped heading level: {style_name} after Heading {last_level}",
                        'location': f"Paragraph {para_index}",
                        'preview': preview
                    })

                last_level = level
            except ValueError:
                # Not a numbered heading style
                pass

    return issues


def _check_empty_paragraphs(doc: Document) -> List[Dict[str, Any]]:
    """Check for empty paragraphs.

    Returns:
        List of empty paragraph issues
    """
    issues = []
    consecutive_empty = 0
    para_index = 0

    for para in doc.paragraphs:
        para_index += 1
        text = para.text.strip()

        if not text:
            consecutive_empty += 1

            # Report clusters of empty paragraphs (more than 2 in a row)
            if consecutive_empty == 3:
                issues.append({
                    'type': 'empty_paragraph',
                    'severity': 'info',
                    'message': f"Multiple consecutive empty paragraphs",
                    'location': f"Around paragraph {para_index - 2}",
                    'preview': ''
                })
        else:
            consecutive_empty = 0

    return issues


def _check_placeholders(doc: Document) -> List[Dict[str, Any]]:
    """Check for placeholder text patterns.

    Patterns detected:
    - {{variable}}
    - [TODO]
    - [TBD]
    - [INSERT]
    - [PLACEHOLDER]
    - <placeholder>
    - ____ (underscores for fill-in)

    Returns:
        List of placeholder issues
    """
    issues = []

    # Placeholder patterns
    patterns = [
        (r'\{\{[^}]+\}\}', 'Template variable'),
        (r'\[TODO\]', 'TODO marker'),
        (r'\[TBD\]', 'TBD marker'),
        (r'\[INSERT[^\]]*\]', 'INSERT placeholder'),
        (r'\[PLACEHOLDER\]', 'PLACEHOLDER marker'),
        (r'<[^>]+placeholder[^>]*>', 'XML-style placeholder'),
        (r'_{4,}', 'Fill-in blanks'),
        (r'\[\s*\]', 'Empty brackets'),
    ]

    para_index = 0
    for para in doc.paragraphs:
        para_index += 1
        text = para.text

        for pattern, pattern_name in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                preview = text[:60] + '...' if len(text) > 60 else text
                issues.append({
                    'type': 'placeholder',
                    'severity': 'warning',
                    'message': f"{pattern_name} found: '{match}'",
                    'location': f"Paragraph {para_index}",
                    'preview': preview
                })

    # Also check tables
    table_index = 0
    for table in doc.tables:
        table_index += 1
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text

                for pattern, pattern_name in patterns:
                    matches = re.findall(pattern, text, re.IGNORECASE)
                    for match in matches:
                        issues.append({
                            'type': 'placeholder',
                            'severity': 'warning',
                            'message': f"{pattern_name} found in table: '{match}'",
                            'location': f"Table {table_index}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                            'preview': text[:40]
                        })

    return issues


def _check_whitespace(doc: Document) -> List[Dict[str, Any]]:
    """Check for whitespace issues.

    Issues detected:
    - Trailing whitespace
    - Multiple consecutive spaces
    - Tab characters

    Returns:
        List of whitespace issues
    """
    issues = []
    para_index = 0

    for para in doc.paragraphs:
        para_index += 1
        text = para.text

        # Skip empty paragraphs
        if not text.strip():
            continue

        # Check for trailing whitespace
        if text != text.rstrip():
            issues.append({
                'type': 'whitespace',
                'severity': 'info',
                'message': "Trailing whitespace",
                'location': f"Paragraph {para_index}",
                'preview': repr(text[-20:]) if len(text) > 20 else repr(text)
            })

        # Check for multiple consecutive spaces
        if '  ' in text:
            issues.append({
                'type': 'whitespace',
                'severity': 'info',
                'message': "Multiple consecutive spaces",
                'location': f"Paragraph {para_index}",
                'preview': text[:50] + '...' if len(text) > 50 else text
            })

    return issues


def _format_validation_report(issues: List[Dict[str, Any]]) -> str:
    """Format issues into a readable validation report.

    Args:
        issues: List of issue dictionaries

    Returns:
        Formatted report string
    """
    if not issues:
        return ""

    lines = [f"Found {len(issues)} issue(s):"]

    # Group by type
    by_type = {}
    for issue in issues:
        issue_type = issue['type']
        if issue_type not in by_type:
            by_type[issue_type] = []
        by_type[issue_type].append(issue)

    for issue_type, type_issues in by_type.items():
        type_name = issue_type.replace('_', ' ').title()
        lines.append(f"\n{type_name} ({len(type_issues)}):")
        for issue in type_issues[:5]:  # Show first 5 of each type
            lines.append(f"  - {issue['message']} at {issue['location']}")
        if len(type_issues) > 5:
            lines.append(f"  ... and {len(type_issues) - 5} more")

    return '\n'.join(lines)
