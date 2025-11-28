"""Search and replace processor for DOCX documents."""

import re
import time
from typing import Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from src.processors import ProcessorResult


def process_document(file_path: str, config: Dict[str, Any]) -> ProcessorResult:
    """Top-level function for multiprocessing (must be module-level).

    Args:
        file_path: Path to the DOCX file to process
        config: Configuration dictionary with search/replace parameters

    Returns:
        ProcessorResult with success status and change count
    """
    start_time = time.time()

    try:
        doc = Document(file_path)
        changes = perform_search_replace(doc, config)

        if changes > 0:
            doc.save(file_path)

        duration = time.time() - start_time
        return ProcessorResult(
            success=True,
            file_path=file_path,
            changes_made=changes,
            duration_seconds=duration
        )
    except Exception as e:
        duration = time.time() - start_time
        return ProcessorResult(
            success=False,
            file_path=file_path,
            changes_made=0,
            error_message=str(e),
            duration_seconds=duration
        )


def perform_search_replace(doc: Document, config: Dict[str, Any]) -> int:
    """Perform search and replace operations on a document.

    Args:
        doc: python-docx Document object
        config: Configuration dictionary containing:
            - search_term: Text or regex pattern to search for
            - replace_term: Replacement text
            - use_regex: If True, treat search_term as regex pattern
            - case_sensitive: If True, perform case-sensitive matching
            - whole_words: If True, match whole words only (literal mode)
            - search_body: If True, search in body paragraphs
            - search_tables: If True, search in table cells
            - search_headers: If True, search in headers
            - search_footers: If True, search in footers

    Returns:
        Total number of replacements made across all document sections
    """
    search_term = config['search_term']
    replace_term = config['replace_term']
    use_regex = config.get('use_regex', False)
    case_sensitive = config.get('case_sensitive', False)
    whole_words = config.get('whole_words', False)
    search_body = config.get('search_body', True)
    search_tables = config.get('search_tables', True)
    search_headers = config.get('search_headers', True)
    search_footers = config.get('search_footers', True)

    total_replacements = 0

    # Compile regex pattern
    pattern = _compile_pattern(search_term, use_regex, case_sensitive, whole_words)

    # Search body paragraphs
    if search_body:
        for paragraph in doc.paragraphs:
            total_replacements += _replace_in_paragraph(paragraph, pattern, replace_term)

    # Search tables
    if search_tables:
        for table in doc.tables:
            total_replacements += _replace_in_table(table, pattern, replace_term)

    # Search headers
    if search_headers:
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                total_replacements += _replace_in_paragraph(paragraph, pattern, replace_term)

    # Search footers
    if search_footers:
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                total_replacements += _replace_in_paragraph(paragraph, pattern, replace_term)

    return total_replacements


def _compile_pattern(search_term: str, use_regex: bool, case_sensitive: bool,
                     whole_words: bool) -> re.Pattern:
    """Compile a regex pattern from search configuration.

    Args:
        search_term: Text or regex pattern to search for
        use_regex: If True, treat search_term as regex pattern
        case_sensitive: If True, perform case-sensitive matching
        whole_words: If True, match whole words only (literal mode only)

    Returns:
        Compiled regex Pattern object
    """
    if use_regex:
        # Use search_term as-is for regex mode
        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(search_term, flags)
    else:
        # Escape special regex characters for literal search
        escaped = re.escape(search_term)

        # Add word boundaries if whole_words is enabled
        if whole_words:
            escaped = r'\b' + escaped + r'\b'

        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(escaped, flags)

    return pattern


def _replace_in_paragraph(paragraph: Paragraph, pattern: re.Pattern,
                          replace_term: str) -> int:
    """Replace matches in a paragraph.

    This function handles text that may be split across multiple runs
    in the Word document XML. When a match spans multiple runs, the
    paragraph is reconstructed with a single run containing the replaced text.

    Note: This approach preserves paragraph-level formatting but may lose
    run-level formatting (e.g., bold/italic on individual words) when
    matches span multiple runs.

    Args:
        paragraph: python-docx Paragraph object
        pattern: Compiled regex pattern to search for
        replace_term: Replacement text

    Returns:
        Number of replacements made in this paragraph
    """
    if not paragraph.text:
        return 0

    # Get the full paragraph text (concatenates all runs)
    original_text = paragraph.text
    new_text, count = pattern.subn(replace_term, original_text)

    if count > 0:
        # Text was modified - we need to reconstruct the paragraph
        # First, try run-by-run replacement (preserves formatting if match is within one run)
        runs_modified = False
        for run in paragraph.runs:
            if run.text and pattern.search(run.text):
                run.text = pattern.sub(replace_term, run.text)
                runs_modified = True

        # If run-by-run didn't work (match spans multiple runs),
        # replace the entire paragraph text
        if not runs_modified:
            # Save formatting from first run (if exists)
            first_run_style = None
            if paragraph.runs:
                first_run = paragraph.runs[0]
                first_run_style = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'font_name': first_run.font.name if first_run.font.name else None,
                    'font_size': first_run.font.size
                }

            # Clear all runs
            for run in paragraph.runs:
                run.text = ''

            # Remove empty runs
            for run in list(paragraph.runs):
                if not run.text:
                    run._element.getparent().remove(run._element)

            # Add new run with replaced text
            new_run = paragraph.add_run(new_text)

            # Apply saved formatting if available
            if first_run_style:
                if first_run_style['bold'] is not None:
                    new_run.bold = first_run_style['bold']
                if first_run_style['italic'] is not None:
                    new_run.italic = first_run_style['italic']
                if first_run_style['underline'] is not None:
                    new_run.underline = first_run_style['underline']
                if first_run_style['font_name']:
                    new_run.font.name = first_run_style['font_name']
                if first_run_style['font_size']:
                    new_run.font.size = first_run_style['font_size']

    return count


def _replace_in_table(table: Table, pattern: re.Pattern, replace_term: str) -> int:
    """Replace matches in all table cells.

    Args:
        table: python-docx Table object
        pattern: Compiled regex pattern to search for
        replace_term: Replacement text

    Returns:
        Total number of replacements made in the table
    """
    total = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                total += _replace_in_paragraph(paragraph, pattern, replace_term)
    return total
