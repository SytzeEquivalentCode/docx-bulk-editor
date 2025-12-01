"""Style enforcement processor for DOCX documents.

Applies consistent styling across documents:
- Font standardization (name and size)
- Heading style enforcement
- Paragraph spacing normalization
"""

import time
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.processors import ProcessorResult


def process_document(file_path: str, config: Dict[str, Any]) -> ProcessorResult:
    """Top-level function for multiprocessing (must be module-level).

    Args:
        file_path: Path to the DOCX file to process
        config: Configuration dictionary with style enforcement options

    Returns:
        ProcessorResult with success status and elements styled count
    """
    start_time = time.time()

    try:
        doc = Document(file_path)
        changes = perform_style_enforcement(doc, config)

        if changes > 0:
            doc.save(file_path)

        duration = time.time() - start_time
        return ProcessorResult(
            success=True,
            file_path=file_path,
            changes_made=changes,
            duration_seconds=duration
        )
    except Exception as e:
        duration = time.time() - start_time
        return ProcessorResult(
            success=False,
            file_path=file_path,
            changes_made=0,
            error_message=str(e),
            duration_seconds=duration
        )


def perform_style_enforcement(doc: Document, config: Dict[str, Any]) -> int:
    """Apply style enforcement to the document.

    Args:
        doc: python-docx Document object
        config: Configuration dictionary containing:
            - style_options: Dict with style options
                - standardize_font: Apply consistent font
                - font_name: Font name (e.g., 'Calibri')
                - font_size_pt: Font size in points
                - enforce_headings: Ensure headings use proper styles
                - normalize_spacing: Standardize paragraph spacing

    Returns:
        Number of elements modified
    """
    style_options = config.get('style_options', {})
    changes = 0

    font_name = style_options.get('font_name', 'Calibri')
    font_size_pt = style_options.get('font_size_pt', 11)

    # Standardize font in body text
    if style_options.get('standardize_font', True):
        changes += _standardize_fonts(doc, font_name, font_size_pt)

    # Enforce heading styles
    if style_options.get('enforce_headings', True):
        changes += _enforce_heading_styles(doc)

    # Normalize paragraph spacing
    if style_options.get('normalize_spacing', False):
        changes += _normalize_spacing(doc)

    return changes


def _standardize_fonts(doc: Document, font_name: str, font_size_pt: int) -> int:
    """Apply consistent font to all body text runs.

    Args:
        doc: python-docx Document object
        font_name: Target font name
        font_size_pt: Target font size in points

    Returns:
        Number of runs modified
    """
    changes = 0
    target_size = Pt(font_size_pt)

    for para in doc.paragraphs:
        # Skip heading paragraphs (they have their own styles)
        if para.style.name.startswith('Heading'):
            continue

        for run in para.runs:
            modified = False

            # Set font name
            if run.font.name != font_name:
                run.font.name = font_name
                modified = True

            # Set font size (only for non-heading text)
            if run.font.size != target_size:
                run.font.size = target_size
                modified = True

            if modified:
                changes += 1

    # Also process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        modified = False

                        if run.font.name != font_name:
                            run.font.name = font_name
                            modified = True

                        if run.font.size != target_size:
                            run.font.size = target_size
                            modified = True

                        if modified:
                            changes += 1

    return changes


def _enforce_heading_styles(doc: Document) -> int:
    """Ensure headings use proper Word heading styles.

    Detects paragraphs that look like headings (short, bold, etc.)
    and applies appropriate Heading styles.

    Returns:
        Number of paragraphs modified
    """
    changes = 0

    # Define heading indicators
    heading_sizes = {
        'Heading 1': Pt(16),
        'Heading 2': Pt(14),
        'Heading 3': Pt(12)
    }

    for para in doc.paragraphs:
        # Skip already-styled headings
        if para.style.name.startswith('Heading'):
            continue

        # Skip empty paragraphs
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph looks like a heading
        # Criteria: short text (< 100 chars), all runs are bold, or large font
        is_short = len(text) < 100
        is_all_bold = all(run.bold for run in para.runs if run.text.strip())
        has_large_font = any(
            run.font.size and run.font.size >= Pt(14)
            for run in para.runs if run.text.strip()
        )

        # Detect potential headings
        if is_short and (is_all_bold or has_large_font):
            # Determine heading level based on font size
            max_font_size = max(
                (run.font.size for run in para.runs if run.font.size),
                default=Pt(11)
            )

            if max_font_size >= Pt(16):
                target_style = 'Heading 1'
            elif max_font_size >= Pt(14):
                target_style = 'Heading 2'
            else:
                target_style = 'Heading 3'

            # Only apply if paragraph is short enough for a heading
            if len(text) < 80:
                try:
                    para.style = target_style
                    changes += 1
                except KeyError:
                    # Style not found in document, skip
                    pass

    return changes


def _normalize_spacing(doc: Document) -> int:
    """Normalize paragraph spacing throughout document.

    Applies consistent before/after spacing and line spacing.

    Returns:
        Number of paragraphs modified
    """
    changes = 0

    # Standard spacing values
    space_before = Pt(0)
    space_after = Pt(8)  # 8pt after each paragraph
    line_spacing = 1.15  # Standard line spacing

    for para in doc.paragraphs:
        modified = False
        pf = para.paragraph_format

        # Skip headings (they have their own spacing)
        if para.style.name.startswith('Heading'):
            continue

        # Set spacing before
        if pf.space_before != space_before:
            pf.space_before = space_before
            modified = True

        # Set spacing after
        if pf.space_after != space_after:
            pf.space_after = space_after
            modified = True

        # Set line spacing
        if pf.line_spacing != line_spacing:
            pf.line_spacing = line_spacing
            modified = True

        if modified:
            changes += 1

    return changes
